from __future__ import annotations

import io
import json
import os
import re
import tempfile
import unicodedata
import httpx
import uuid
from datetime import datetime
from pathlib import Path
from typing import Any

import streamlit as st
import streamlit.components.v1 as components
from openai import OpenAI

try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    pd = None
    PANDAS_AVAILABLE = False

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt
    DOCX_AVAILABLE = True
except ImportError:
    Document = Any
    WD_ALIGN_PARAGRAPH = None
    Inches = None
    Pt = None
    DOCX_AVAILABLE = False

try:
    import xlsxwriter
    XLSX_AVAILABLE = True
except ImportError:
    xlsxwriter = None
    XLSX_AVAILABLE = False


# =========================================================
# CẤU HÌNH ỨNG DỤNG
# =========================================================
APP_TITLE = "Trợ lý CCTL_QNG"
APP_ICON = "💧"

BASE_DIR = Path(__file__).resolve().parent
DATA_DIR = BASE_DIR / "data"
DATA_FILE = DATA_DIR / "conversations.json"
TABLE_DATA_DIR = DATA_DIR / "table_data"

DEFAULT_MODEL = os.getenv("OPENAI_MODEL", "gpt-5-mini")
FAST_MODEL = os.getenv("OPENAI_FAST_MODEL", "gpt-5-nano")
SEARCH_MODEL = os.getenv("OPENAI_SEARCH_MODEL", "gpt-5-mini")
DEEP_MODEL = os.getenv("OPENAI_DEEP_MODEL", DEFAULT_MODEL)
DEFAULT_VECTOR_STORE_ID = os.getenv("OPENAI_VECTOR_STORE_ID", "").strip()

SYSTEM_INSTRUCTIONS = """
Bạn là Trợ lý CCTL_QNG, hỗ trợ công tác tham mưu của Chi cục Thủy lợi tỉnh Quảng Ngãi.

NGUYÊN TẮC TRẢ LỜI:
- Trả lời bằng tiếng Việt, đúng trọng tâm câu hỏi.
- Ưu tiên câu trả lời trực tiếp ngay ở câu đầu tiên.
- Không nhắc lại câu hỏi.
- Không mở đầu bằng các câu xã giao như “Kính anh/chị”, “Cảm ơn anh/chị”.
- Không trình bày quá trình suy luận hoặc mô tả dài dòng cách đã tìm kiếm.
- Không tự bổ sung thông tin ngoài tài liệu khi câu hỏi yêu cầu tra cứu kho.
- Không lặp lại cùng một ý ở nhiều đoạn.
- Không đưa ra các đề nghị hỗ trợ thêm ở cuối câu trả lời, trừ khi người dùng yêu cầu.
- Chỉ nêu cảnh báo hoặc mức độ không chắc chắn khi thực sự cần thiết.
- Với dữ liệu lấy từ bảng, biểu mẫu hoặc danh sách nhiều cột, phải kiểm tra đúng tiêu đề cột trước khi kết luận.
- Đặc biệt thận trọng với CCCD/CMND, số điện thoại, mã số BHXH, số thẻ BHYT, ngày sinh, số tài khoản và các chuỗi số dài.
- Không được suy ra loại dữ liệu chỉ dựa vào độ dài chuỗi số.
- Nếu đoạn trích bị mất hàng/cột, bị tách bảng, thiếu tiêu đề cột hoặc không đủ ngữ cảnh để ghép đúng người với đúng giá trị, không được khẳng định chắc chắn.
- Khi có khả năng nhầm cột hoặc nhầm dòng, phải nêu cảnh báo ngắn ngay sau câu trả lời.
- Chỉ xác nhận số định danh cá nhân khi đồng thời xác định được: đúng người, đúng cột dữ liệu và đúng dòng tương ứng.

ĐỘ DÀI MẶC ĐỊNH:
- Câu hỏi hỏi một thông tin cụ thể: trả lời trong 1–3 câu.
- Câu hỏi cần liệt kê: dùng danh sách ngắn, tối đa khoảng 5–7 ý.
- Câu hỏi cần phân tích hoặc tham mưu: ưu tiên 3 phần ngắn:
  1. Kết luận;
  2. Căn cứ hoặc nhận xét chính;
  3. Đề xuất xử lý.
- Chỉ trả lời dài khi người dùng yêu cầu giải thích chi tiết, soạn thảo văn bản hoặc phân tích chuyên sâu.

TRẢ LỜI TỪ KHO TÀI LIỆU:
- Nếu đã tìm thấy thông tin rõ ràng, nêu ngay kết quả.
- Không chép lại toàn bộ đoạn trích nếu người dùng không yêu cầu.
- Không mô tả dài về cấu trúc file, quá trình ghép bảng hoặc rủi ro kỹ thuật.
- Nếu tài liệu chưa đủ căn cứ, nói ngắn gọn: “Chưa đủ căn cứ trong tài liệu đã tra cứu để kết luận.”
- Không suy đoán ngoài nội dung tài liệu.
- Cuối câu trả lời luôn giữ một mục nguồn riêng, theo đúng mẫu:

**Nguồn văn bản trong kho:**
- [Tên tài liệu 1]
- [Tên tài liệu 2]

- Chỉ liệt kê những tài liệu thực sự đã được dùng để trả lời.
- Nếu xác định được trang, điều, khoản hoặc bảng thì ghi sau tên tài liệu.
- Không dùng tên file tạm kiểu tmp... nếu đã có tên file gốc.
- Không lặp lại tên nguồn ở phần nội dung chính.
- Nếu dữ liệu có rủi ro nhầm cột, thêm ngay trước mục nguồn một đoạn ngắn theo mẫu:

**Khuyến cáo kiểm tra:** Kết quả trên được đọc từ bảng/đoạn trích có thể bị tách cột hoặc thiếu tiêu đề; cần đối chiếu trực tiếp file gốc trước khi sử dụng chính thức.

- Không dùng cảnh báo này khi dữ liệu đã được xác định chắc chắn từ đúng cột, đúng dòng và đủ ngữ cảnh.

ĐỊNH DẠNG:
- Dùng bảng Markdown khi người dùng yêu cầu bảng, so sánh hoặc xuất Excel.
- Với yêu cầu soạn thảo văn bản hành chính, trình bày chặt chẽ, đúng thể thức diễn đạt.
- Chỉ dùng sơ đồ hoặc biểu đồ khi người dùng yêu cầu hoặc khi thực sự cần thiết.
- Phải kết thúc trọn câu, trọn ý; không dừng giữa câu hoặc giữa danh sách.
""".strip()




# =========================================================
# XUẤT WORD VÀ EXCEL
# =========================================================
def clean_export_text(content: str) -> str:
    """Loại bỏ các khối biểu đồ/sơ đồ không phù hợp khi xuất văn bản."""
    cleaned = re.sub(
        r"```(?:mermaid|chart)\s*\n.*?```",
        "",
        content,
        flags=re.IGNORECASE | re.DOTALL,
    )
    return cleaned.strip()


def extract_markdown_tables(content: str) -> list[dict[str, Any]]:
    """Tách các bảng Markdown hợp lệ trong câu trả lời."""
    lines = content.splitlines()
    tables: list[dict[str, Any]] = []
    index = 0

    def split_row(line: str) -> list[str]:
        stripped = line.strip().strip("|")
        return [cell.strip() for cell in stripped.split("|")]

    def is_separator(line: str) -> bool:
        cells = split_row(line)
        return bool(cells) and all(
            re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) is not None
            for cell in cells
        )

    while index < len(lines) - 1:
        if "|" in lines[index] and is_separator(lines[index + 1]):
            headers = split_row(lines[index])
            rows: list[list[str]] = []
            index += 2

            while index < len(lines) and "|" in lines[index]:
                row = split_row(lines[index])
                if len(row) == len(headers):
                    rows.append(row)
                index += 1

            if headers:
                tables.append(
                    {
                        "headers": headers,
                        "rows": rows,
                    }
                )
            continue

        index += 1

    return tables


def add_markdown_to_docx(document: Document, content: str) -> None:
    """Chuyển phần Markdown thường gặp thành nội dung Word."""
    lines = clean_export_text(content).splitlines()
    index = 0

    while index < len(lines):
        line = lines[index].rstrip()

        # Bảng Markdown.
        if (
            index + 1 < len(lines)
            and "|" in line
            and re.search(r"\|?\s*:?-{3,}:?\s*\|", lines[index + 1])
        ):
            table_lines = [line, lines[index + 1]]
            index += 2
            while index < len(lines) and "|" in lines[index]:
                table_lines.append(lines[index])
                index += 1

            tables = extract_markdown_tables("\n".join(table_lines))
            if tables:
                table_data = tables[0]
                headers = table_data["headers"]
                rows = table_data["rows"]
                word_table = document.add_table(
                    rows=1,
                    cols=len(headers),
                )
                word_table.style = "Table Grid"

                for column_index, header in enumerate(headers):
                    word_table.rows[0].cells[column_index].text = header

                for row in rows:
                    cells = word_table.add_row().cells
                    for column_index, value in enumerate(row):
                        cells[column_index].text = value

                document.add_paragraph()
            continue

        stripped = line.strip()
        if not stripped:
            document.add_paragraph()
            index += 1
            continue

        heading_match = re.match(r"^(#{1,4})\s+(.+)$", stripped)
        if heading_match:
            level = min(len(heading_match.group(1)), 4)
            document.add_heading(
                heading_match.group(2),
                level=level,
            )
            index += 1
            continue

        numbered_match = re.match(r"^\d+[.)]\s+(.+)$", stripped)
        bullet_match = re.match(r"^[-*•]\s+(.+)$", stripped)

        if numbered_match:
            paragraph = document.add_paragraph(
                style="List Number",
            )
            paragraph.add_run(numbered_match.group(1))
        elif bullet_match:
            paragraph = document.add_paragraph(
                style="List Bullet",
            )
            paragraph.add_run(bullet_match.group(1))
        else:
            paragraph = document.add_paragraph()
            paragraph.add_run(
                re.sub(r"\*\*(.*?)\*\*", r"\1", stripped)
            )

        index += 1


def build_word_bytes(content: str) -> bytes | None:
    """Tạo file Word từ câu trả lời."""
    if not DOCX_AVAILABLE:
        return None

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.75)

    styles = document.styles
    normal_style = styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style.font.size = Pt(13)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("NỘI DUNG TRẢ LỜI CỦA TRỢ LÝ CCTL_QNG")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)

    document.add_paragraph()
    add_markdown_to_docx(document, content)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def build_excel_bytes(content: str) -> bytes | None:
    """Tạo Excel từ tất cả bảng Markdown trong câu trả lời."""
    if not XLSX_AVAILABLE:
        return None

    tables = extract_markdown_tables(content)
    if not tables:
        return None

    buffer = io.BytesIO()
    workbook = xlsxwriter.Workbook(
        buffer,
        {"in_memory": True},
    )

    header_format = workbook.add_format(
        {
            "bold": True,
            "align": "center",
            "valign": "vcenter",
            "text_wrap": True,
            "border": 1,
            "bg_color": "#D9EAF7",
        }
    )
    cell_format = workbook.add_format(
        {
            "valign": "top",
            "text_wrap": True,
            "border": 1,
        }
    )
    title_format = workbook.add_format(
        {
            "bold": True,
            "font_size": 14,
            "align": "center",
            "valign": "vcenter",
        }
    )

    for table_index, table_data in enumerate(tables, start=1):
        sheet_name = f"Bang_{table_index}"[:31]
        worksheet = workbook.add_worksheet(sheet_name)
        headers = table_data["headers"]
        rows = table_data["rows"]

        if headers:
            worksheet.merge_range(
                0,
                0,
                0,
                max(len(headers) - 1, 0),
                f"BẢNG {table_index}",
                title_format,
            )

        for column_index, header in enumerate(headers):
            worksheet.write(
                2,
                column_index,
                header,
                header_format,
            )

        for row_index, row in enumerate(rows, start=3):
            for column_index, value in enumerate(row):
                worksheet.write(
                    row_index,
                    column_index,
                    value,
                    cell_format,
                )

        worksheet.freeze_panes(3, 0)
        worksheet.autofilter(
            2,
            0,
            max(2, len(rows) + 2),
            max(0, len(headers) - 1),
        )

        for column_index, header in enumerate(headers):
            longest = len(str(header))
            for row in rows[:200]:
                if column_index < len(row):
                    longest = max(
                        longest,
                        len(str(row[column_index])),
                    )
            worksheet.set_column(
                column_index,
                column_index,
                min(max(longest + 2, 12), 40),
            )

    workbook.close()
    buffer.seek(0)
    return buffer.getvalue()


def render_export_buttons(
    content: str,
    *,
    key_prefix: str,
) -> None:
    """Hiển thị nút tải Word và Excel dưới mỗi câu trả lời."""
    word_bytes = build_word_bytes(content)
    excel_bytes = build_excel_bytes(content)

    col_word, col_excel, col_note = st.columns([1.25, 1.25, 3.5])

    with col_word:
        if word_bytes:
            st.download_button(
                "📄 Tải Word",
                data=word_bytes,
                file_name=f"Tra_loi_CCTL_QNG_{datetime.now():%Y%m%d_%H%M%S}.docx",
                mime=(
                    "application/vnd.openxmlformats-officedocument."
                    "wordprocessingml.document"
                ),
                key=f"{key_prefix}_word",
                use_container_width=True,
            )
        else:
            st.button(
                "📄 Chưa sẵn sàng",
                key=f"{key_prefix}_no_word",
                disabled=True,
                use_container_width=True,
            )

    with col_excel:
        if excel_bytes:
            st.download_button(
                "📊 Tải Excel",
                data=excel_bytes,
                file_name=f"Bang_CCTL_QNG_{datetime.now():%Y%m%d_%H%M%S}.xlsx",
                mime=(
                    "application/vnd.openxmlformats-officedocument."
                    "spreadsheetml.sheet"
                ),
                key=f"{key_prefix}_excel",
                use_container_width=True,
            )
        else:
            st.button(
                "📊 Chưa có bảng",
                key=f"{key_prefix}_no_excel",
                disabled=True,
                use_container_width=True,
            )

    with col_note:
        if not DOCX_AVAILABLE:
            st.caption("Chưa cài python-docx nên tạm thời chưa xuất Word.")
        elif not XLSX_AVAILABLE:
            st.caption("Chưa cài XlsxWriter nên tạm thời chưa xuất Excel.")
        elif not excel_bytes:
            st.caption("Excel chỉ xuất khi câu trả lời có bảng Markdown.")


# =========================================================
# HIỂN THỊ NỘI DUNG TRỰC QUAN
# =========================================================
def render_mermaid_diagram(code: str, *, height: int = 460) -> None:
    """Hiển thị sơ đồ Mermaid trong một khung nhúng."""
    diagram_json = json.dumps(code)
    components.html(
        f"""
        <div class="mermaid-wrap">
            <div id="diagram" class="mermaid"></div>
        </div>
        <script type="module">
            import mermaid from "https://cdn.jsdelivr.net/npm/mermaid@11/dist/mermaid.esm.min.mjs";
            mermaid.initialize({{
                startOnLoad: false,
                securityLevel: "loose",
                theme: "neutral",
                flowchart: {{ useMaxWidth: true, htmlLabels: true }}
            }});
            const source = {diagram_json};
            const container = document.getElementById("diagram");
            try {{
                const result = await mermaid.render(
                    "mermaid-" + Math.random().toString(36).slice(2),
                    source
                );
                container.innerHTML = result.svg;
            }} catch (error) {{
                container.innerHTML =
                    "<pre style='white-space:pre-wrap;color:#b91c1c'>" +
                    "Không dựng được sơ đồ Mermaid: " +
                    String(error) +
                    "</pre>";
            }}
        </script>
        <style>
            body {{
                margin: 0;
                font-family: Arial, sans-serif;
                background: transparent;
            }}
            .mermaid-wrap {{
                border: 1px solid #e6e8eb;
                border-radius: 14px;
                padding: 14px;
                background: #ffffff;
                overflow: auto;
            }}
            .mermaid-wrap svg {{
                max-width: 100%;
                height: auto;
            }}
        </style>
        """,
        height=height,
        scrolling=True,
    )


def render_chart_block(raw_json: str) -> None:
    """Hiển thị biểu đồ bar, line hoặc pie từ khối JSON do mô hình tạo."""
    try:
        chart = json.loads(raw_json)
    except json.JSONDecodeError as error:
        st.warning(f"Không đọc được dữ liệu biểu đồ: {error}")
        st.code(raw_json, language="json")
        return

    chart_type = str(chart.get("type", "bar")).lower()
    title = str(chart.get("title", "Biểu đồ"))
    categories = list(chart.get("categories", []) or [])
    values = list(chart.get("values", []) or [])
    unit = str(chart.get("unit", "") or "")

    if not categories or not values or len(categories) != len(values):
        st.warning("Dữ liệu biểu đồ chưa hợp lệ.")
        st.code(raw_json, language="json")
        return

    data = [
        {"category": str(category), "value": value}
        for category, value in zip(categories, values)
    ]

    if chart_type == "pie":
        spec = {
            "title": title,
            "data": {"values": data},
            "mark": {"type": "arc", "tooltip": True},
            "encoding": {
                "theta": {"field": "value", "type": "quantitative"},
                "color": {"field": "category", "type": "nominal"},
                "tooltip": [
                    {"field": "category", "type": "nominal", "title": "Hạng mục"},
                    {
                        "field": "value",
                        "type": "quantitative",
                        "title": unit or "Giá trị",
                    },
                ],
            },
            "view": {"stroke": None},
        }
    else:
        mark_type = "line" if chart_type == "line" else "bar"
        spec = {
            "title": title,
            "data": {"values": data},
            "mark": {"type": mark_type, "point": chart_type == "line", "tooltip": True},
            "encoding": {
                "x": {
                    "field": "category",
                    "type": "nominal",
                    "title": None,
                    "sort": None,
                },
                "y": {
                    "field": "value",
                    "type": "quantitative",
                    "title": unit or "Giá trị",
                },
                "tooltip": [
                    {"field": "category", "type": "nominal", "title": "Hạng mục"},
                    {
                        "field": "value",
                        "type": "quantitative",
                        "title": unit or "Giá trị",
                    },
                ],
            },
        }

    st.vega_lite_chart(spec, use_container_width=True)


def render_assistant_content(content: str) -> None:
    """
    Hiển thị nội dung trợ lý:
    - Markdown và bảng Markdown;
    - sơ đồ Mermaid;
    - biểu đồ JSON;
    - giữ nguyên các khối mã khác.
    """
    block_pattern = re.compile(
        r"```(?P<language>mermaid|chart)\s*\n(?P<body>.*?)```",
        re.IGNORECASE | re.DOTALL,
    )

    cursor = 0
    found_visual = False

    for match in block_pattern.finditer(content):
        markdown_part = content[cursor:match.start()].strip()
        if markdown_part:
            st.markdown(markdown_part)

        language = match.group("language").lower()
        body = match.group("body").strip()

        if language == "mermaid":
            render_mermaid_diagram(body)
        else:
            render_chart_block(body)

        found_visual = True
        cursor = match.end()

    remaining = content[cursor:].strip()
    if remaining:
        st.markdown(remaining)
    elif not found_visual and content.strip():
        st.markdown(content)


# =========================================================
# THIẾT LẬP TRANG
# =========================================================
st.set_page_config(
    page_title=APP_TITLE,
    page_icon=APP_ICON,
    layout="wide",
    initial_sidebar_state="collapsed",
)

st.markdown(
    """
    <style>
    :root {
        --border: #e6e8eb;
        --muted: #6b7280;
        --panel: #f8fafc;
        --blue-soft: #eaf4ff;
    }

    #MainMenu, footer {visibility: hidden;}

    [data-testid="stAppViewContainer"] {
        background: #ffffff;
    }

    [data-testid="stHeader"] {
        background: rgba(255,255,255,0.94);
        backdrop-filter: blur(12px);
        border-bottom: 1px solid rgba(230,232,235,0.7);
    }

    [data-testid="stSidebar"] {
        border-right: 1px solid var(--border);
        background: #fbfbfc;
    }

    [data-testid="stSidebar"] .block-container {
        padding-top: 1rem;
        padding-bottom: 1rem;
    }

    .block-container {
        max-width: 1050px;
        padding-top: 0.5rem;
        padding-bottom: 3rem;
    }

    /* Thông tin thương hiệu đặt trong sidebar trái */
    .sidebar-subtitle {
        margin-top: 0.15rem;
        font-size: 0.78rem;
        color: #374151;
        line-height: 1.4;
    }

    .sidebar-author {
        margin-top: -0.35rem;
        margin-bottom: 0.55rem;
        font-size: 0.70rem;
        color: var(--muted);
        line-height: 1.25;
    }

    [data-testid="stMainBlockContainer"] {
        padding-top: 1.2rem;
        padding-bottom: 8.5rem;
    }

    /* CHAT */
    div[data-testid="stChatMessage"] {
        border: 1px solid #eef0f2;
        border-radius: 20px;
        padding: 0.65rem 0.8rem;
        margin-bottom: 0.55rem;
        background: #fff;
    }

    div[data-testid="stChatMessage"]:has([data-testid="chatAvatarIcon-user"]) {
        background: var(--blue-soft);
    }

    /* Khung nhập dưới cùng giống ChatGPT */
    div[data-testid="stChatInput"] {
        max-width: 860px;
        margin: 0 auto;
    }

    div[data-testid="stChatInput"] > div {
        min-height: 58px;
        border-radius: 28px !important;
        border: 1px solid #d7dce1 !important;
        background: #fff !important;
        box-shadow: 0 8px 28px rgba(15,23,42,0.12);
    }

    div[data-testid="stChatInput"] textarea {
        padding-top: 0.95rem !important;
        padding-bottom: 0.85rem !important;
    }

    /* SIDEBAR */
    .sidebar-brand {
        font-size: 1.05rem;
        font-weight: 800;
        padding: 0.2rem 0 0.75rem 0;
    }

    [data-testid="stSidebar"] .stButton > button {
        border-radius: 12px;
        min-height: 2.6rem;
        text-align: left;
    }

    [data-testid="stSidebar"] .stButton > button[kind="primary"] {
        background: #111827;
        border-color: #111827;
    }

    /* Nút xóa cuộc trò chuyện bên cạnh tiêu đề */
    [data-testid="stSidebar"] [data-testid="stHorizontalBlock"] {
        gap: 0.35rem;
        align-items: center;
    }

    [data-testid="stSidebar"] [data-testid="column"]:last-child .stButton > button {
        min-height: 2.6rem;
        padding-left: 0.35rem;
        padding-right: 0.35rem;
        border-radius: 10px;
    }

    /* PANEL PHẢI */
    .right-panel-title {
        font-size: 0.95rem;
        font-weight: 750;
        margin-bottom: 0.6rem;
    }

    .info-card {
        border: 1px solid var(--border);
        border-radius: 16px;
        background: var(--panel);
        padding: 0.85rem;
        margin-bottom: 0.7rem;
    }

    .info-label {
        color: var(--muted);
        font-size: 0.74rem;
        margin-bottom: 0.12rem;
    }

    .info-value {
        font-size: 0.86rem;
        font-weight: 650;
        word-break: break-word;
    }

    .status-dot {
        display: inline-block;
        width: 8px;
        height: 8px;
        border-radius: 50%;
        background: #22c55e;
        margin-right: 6px;
    }

    .small-note {
        color: var(--muted);
        font-size: 0.74rem;
        line-height: 1.45;
    }

    /* GIAO DIỆN HỘI THOẠI THEO MẪU ĐÃ THỐNG NHẤT */
    div[data-testid="stChatMessage"] {
        border: 1px solid #e4e8ee !important;
        border-radius: 18px !important;
        padding: 0.78rem 0.9rem !important;
        margin-bottom: 0.72rem !important;
        background: #ffffff !important;
        box-shadow: 0 2px 8px rgba(15, 23, 42, 0.035);
    }

    div[data-testid="stChatMessage"]:has([data-testid="chatAvatarIcon-user"]) {
        background: #ffffff !important;
        border-color: #ead9ab !important;
    }

    div[data-testid="stChatMessage"] [data-testid="stChatMessageAvatarUser"],
    div[data-testid="stChatMessage"] [data-testid="stChatMessageAvatarAssistant"] {
        width: 2.35rem !important;
        height: 2.35rem !important;
        min-width: 2.35rem !important;
        border-radius: 12px !important;
        border: 1px solid #dde3ea;
        background: #f8fafc !important;
        display: flex !important;
        align-items: center !important;
        justify-content: center !important;
        font-size: 1.35rem !important;
        box-shadow: 0 1px 4px rgba(15, 23, 42, 0.05);
    }

    div[data-testid="stChatMessage"] [data-testid="stChatMessageAvatarUser"] {
        background: #fff8dc !important;
        border-color: #f0c85a !important;
    }

    div[data-testid="stChatMessage"] [data-testid="stChatMessageAvatarAssistant"] {
        background: #eef7ff !important;
        border-color: #cfe6fb !important;
    }

    div[data-testid="stChatMessage"] p,
    div[data-testid="stChatMessage"] li {
        font-size: 0.97rem;
        line-height: 1.58;
    }

    div[data-testid="stChatInput"] > div {
        border-radius: 24px !important;
        border: 1px solid #d8dee6 !important;
        box-shadow: 0 6px 22px rgba(15, 23, 42, 0.08) !important;
    }

    @media (max-width: 900px) {
        [data-testid="stMainBlockContainer"] {
            padding-top: 1rem;
            padding-bottom: 8rem;
        }

        .block-container {
            padding-left: 0.75rem;
            padding-right: 0.75rem;
        }
    }


    /* Màn hình chào theo mẫu điện thoại */
    .welcome-shell {
        width: min(100%, 680px);
        margin: 0 auto;
        text-align: center;
        padding: 2.9rem 0.55rem 0.35rem;
        display: flex;
        flex-direction: column;
        align-items: center;
    }

    .welcome-logo {
        width: 84px;
        height: 84px;
        object-fit: contain;
        display: block;
        margin: 0 auto 0.45rem;
        align-self: center;
        transform: translateX(-1px);
        filter: drop-shadow(0 2px 5px rgba(0, 91, 150, 0.12));
    }

    .welcome-shell .welcome-title {
        display: block;
        width: 100%;
        text-align: center !important;
        font-size: clamp(1.26rem, 4vw, 1.9rem);
        line-height: 1.02;
        font-weight: 800;
        letter-spacing: -0.02em;
        color: #005B96 !important;
        -webkit-text-fill-color: #005B96 !important;
        margin: 0 auto 0.52rem;
    }

    .welcome-card {
        width: min(100%, 610px);
        margin: 0 auto 0.42rem;
        padding: 0.62rem 0.82rem;
        border: 1px solid #e3e7eb;
        border-radius: 18px;
        background: rgba(255,255,255,0.99);
        box-shadow: 0 2px 8px rgba(15, 23, 42, 0.035);
        color: #222222;
        font-size: clamp(0.78rem, 2.7vw, 0.94rem);
        line-height: 1.2;
        text-align: center;
    }

    .welcome-card-large {
        padding-top: 0.66rem;
        padding-bottom: 0.66rem;
    }

    /* Khi chưa có hội thoại, tạo khoảng thở như mẫu */
    .welcome-spacer {
        height: 0;
    }

    @media (max-width: 700px) {
        [data-testid="stMainBlockContainer"] {
            padding-top: 0.05rem;
            padding-left: 0.55rem;
            padding-right: 0.55rem;
            padding-bottom: 6.1rem;
        }

        .welcome-shell {
            padding-top: 3.15rem;
            padding-left: 0.18rem;
            padding-right: 0.18rem;
            padding-bottom: 0.18rem;
            display: flex;
            flex-direction: column;
            align-items: center;
        }

        .welcome-logo {
            width: 78px;
            height: 78px;
            margin: 0 auto 0.38rem;
            align-self: center;
            transform: translateX(-1px);
        }

        .welcome-shell .welcome-title {
            display: block;
            width: 100%;
            text-align: center !important;
            font-size: 1.22rem;
            margin: 0 auto 0.46rem;
            color: #005B96 !important;
            -webkit-text-fill-color: #005B96 !important;
        }

        .welcome-card {
            width: 100%;
            border-radius: 17px;
            padding: 0.56rem 0.68rem;
            margin-bottom: 0.38rem;
            font-size: 0.78rem;
            line-height: 1.18;
            text-align: center;
        }

        .welcome-card-large {
            padding-top: 0.62rem;
            padding-bottom: 0.62rem;
        }

        div[data-testid="stChatInput"] {
            width: calc(100% - 1rem);
            bottom: 0.35rem !important;
        }

        div[data-testid="stChatInput"] > div {
            min-height: 50px;
            border-radius: 25px !important;
            box-shadow: none;
        }

        [data-testid="stChatInputTextArea"] {
            font-size: 0.84rem !important;
            line-height: 1.12 !important;
        }
    }

    </style>
    """,
    unsafe_allow_html=True,
)


# =========================================================
# HÀM DỮ LIỆU
# =========================================================
def now_text() -> str:
    return datetime.now().isoformat(timespec="seconds")


def new_database() -> dict[str, Any]:
    return {
        "active_id": None,
        "vector_store_id": DEFAULT_VECTOR_STORE_ID,
        "uploaded_files": [],
        "table_files": [],
        "conversations": {},
    }


def save_database(database: dict[str, Any]) -> None:
    DATA_DIR.mkdir(parents=True, exist_ok=True)
    temp_path = DATA_FILE.with_suffix(".tmp")
    temp_path.write_text(
        json.dumps(database, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    temp_path.replace(DATA_FILE)


def load_database() -> dict[str, Any]:
    DATA_DIR.mkdir(parents=True, exist_ok=True)

    if not DATA_FILE.exists():
        database = new_database()
        save_database(database)
        return database

    try:
        database = json.loads(DATA_FILE.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError):
        database = new_database()

    database.setdefault("active_id", None)
    database.setdefault("vector_store_id", DEFAULT_VECTOR_STORE_ID)
    database.setdefault("uploaded_files", [])
    database.setdefault("table_files", [])
    database.setdefault("conversations", {})

    configured_vector_store_id = get_configured_vector_store_id()
    if configured_vector_store_id:
        database["vector_store_id"] = configured_vector_store_id

    return database


def create_conversation(
    database: dict[str, Any],
    title: str = "Yêu cầu hỗ trợ mới",
) -> str:
    conversation_id = uuid.uuid4().hex
    database["conversations"][conversation_id] = {
        "id": conversation_id,
        "title": title,
        "created_at": now_text(),
        "updated_at": now_text(),
        "messages": [],
    }
    database["active_id"] = conversation_id
    save_database(database)
    return conversation_id


def get_active_conversation(database: dict[str, Any]) -> dict[str, Any]:
    active_id = database.get("active_id")
    if not active_id or active_id not in database["conversations"]:
        active_id = create_conversation(database)
    return database["conversations"][active_id]


def shorten_title(text: str, max_length: int = 38) -> str:
    text = " ".join(text.split())
    return text if len(text) <= max_length else text[: max_length - 1].rstrip() + "…"


def append_message(
    database: dict[str, Any],
    conversation_id: str,
    role: str,
    content: str,
) -> None:
    conversation = database["conversations"][conversation_id]
    conversation["messages"].append(
        {
            "role": role,
            "content": content,
            "created_at": now_text(),
        }
    )
    conversation["updated_at"] = now_text()

    # Cập nhật tên trong mục "Trò chuyện gần đây" theo câu hỏi mới nhất
    if role == "user":
        conversation["title"] = shorten_title(content)

    save_database(database)



# =========================================================
# KHO BẢNG DỮ LIỆU CÓ CẤU TRÚC
# =========================================================
TABLE_COLUMN_ALIASES: dict[str, list[str]] = {
    "Họ và tên": ["ho va ten", "ho ten", "họ và tên", "họ tên", "ho_ten"],
    "Ngày sinh": ["ngay sinh", "ngày sinh", "ngay thang nam sinh"],
    "Chức vụ": ["chuc vu", "chức vụ", "chuc danh", "chức danh"],
    "Quê quán": ["que quan", "quê quán"],
    "Trình độ chuyên môn": ["trinh do chuyen mon", "trình độ chuyên môn"],
    "Trình độ LLCT": ["trinh do llct", "trình độ llct", "ly luan chinh tri"],
    "Ngày vào Đảng": ["ngay vao dang", "ngày vào đảng"],
    "CCCD": ["cccd", "cmnd", "can cuoc", "căn cước"],
    "Ngày cấp CCCD": ["ngay cap", "ngày cấp", "ngay thang nam cap"],
    "Số sổ BHXH": [
        "so so bhxh", "số sổ bhxh", "bhxh",
        "so so bao hiem xa hoi", "số sổ bảo hiểm xã hội",
        "so bao hiem xa hoi", "số bảo hiểm xã hội",
        "ma so bao hiem xa hoi", "mã số bảo hiểm xã hội",
        "bao hiem xa hoi", "bảo hiểm xã hội",
    ],
    "Số thẻ BHYT": [
        "so the bhyt", "số thẻ bhyt", "bhyt",
        "so the bao hiem y te", "số thẻ bảo hiểm y tế",
        "bao hiem y te", "bảo hiểm y tế",
    ],
    "Số điện thoại": ["dien thoai", "điện thoại", "so dien thoai", "sdt"],
    "STT": ["stt", "tt", "số thứ tự", "so thu tu"],
    "Tên công trình": [
        "tên công trình", "ten cong trinh", "tên ct", "ten ct",
        "công trình", "cong trinh", "tên hồ", "ten ho",
        "tên đập", "ten dap", "tên kênh", "ten kenh",
        "tên trạm bơm", "ten tram bom", "tên kè", "ten ke",
        "tên đê", "ten de", "tên mỏ hàn", "ten mo han",
    ],
    "Địa điểm": [
        "địa điểm", "dia diem", "xã phường đặc khu",
        "xa phuong dac khu", "vị trí", "vi tri",
    ],
    "Loại công trình": [
        "loại công trình", "loai cong trinh", "loại", "loai",
        "nhóm công trình", "nhom cong trinh",
    ],
}


# Các loại công trình được hiểu là thành phần của khái niệm chung “công trình”.
STRUCTURE_TYPE_ALIASES: dict[str, list[str]] = {
    "Hồ chứa": ["hồ chứa", "ho chua", "hồ", "ho", "hồ chứa nước", "ho chua nuoc"],
    "Đập": ["đập", "dap", "đập dâng", "dap dang", "đập ngăn mặn", "dap ngan man"],
    "Đê": ["đê", "de", "đê sông", "de song", "đê biển", "de bien"],
    "Kè": ["kè", "ke", "kè sông", "ke song", "kè biển", "ke bien"],
    "Mỏ hàn": ["mỏ hàn", "mo han", "mỏ hàn chỉnh trị", "mo han chinh tri"],
    "Kênh": ["kênh", "kenh", "kênh tưới", "kenh tuoi", "kênh tiêu", "kenh tieu", "mương", "muong"],
    "Trạm bơm": ["trạm bơm", "tram bom", "nhà máy bơm", "nha may bom"],
    "Cống": ["cống", "cong", "cống lấy nước", "cong lay nuoc", "cống tiêu", "cong tieu"],
    "Đập ngăn mặn": ["đập ngăn mặn", "dap ngan man", "ngăn mặn", "ngan man"],
    "Công trình cấp nước": ["cấp nước", "cap nuoc", "công trình cấp nước", "cong trinh cap nuoc"],
}

# Từ khóa nhận diện loại công trình theo tên sheet/phụ lục.
STRUCTURE_SHEET_HINTS: dict[str, list[str]] = {
    "Hồ chứa": ["hcn", "ho chua", "ho chua nuoc", "pl2"],
    "Đập": ["dap", "pl3"],
    "Đập ngăn mặn": ["dnm", "dap ngan man", "pl4"],
    "Trạm bơm": ["tram bom", "tb", "pl5"],
    "Kè": ["ke", "pl6"],
    "Mỏ hàn": ["mhan", "mo han", "pl7"],
    "Đê": ["de", "pl8"],
    "Kênh": ["kenh", "pl9"],
    "Cống": ["cong", "pl10"],
}


def normalize_table_text(value: Any) -> str:
    value_text = "" if value is None else str(value)
    value_text = unicodedata.normalize("NFD", value_text)
    value_text = "".join(
        char for char in value_text
        if unicodedata.category(char) != "Mn"
    )
    value_text = value_text.lower().strip()
    value_text = re.sub(r"[^a-z0-9]+", " ", value_text)
    return " ".join(value_text.split())


def canonical_column_name(column_name: Any) -> str:
    normalized = normalize_table_text(column_name)
    for canonical, aliases in TABLE_COLUMN_ALIASES.items():
        alias_set = {
            normalize_table_text(item)
            for item in aliases + [canonical]
        }
        if normalized in alias_set:
            return canonical
    return str(column_name).strip()


def clean_table_dataframe(dataframe: Any) -> Any:
    """Chuẩn hóa bảng CSV/Excel, kể cả bảng có tiêu đề nhiều tầng và ô gộp."""
    dataframe = dataframe.copy()

    # Chuẩn hóa ô trước khi dò tiêu đề; giữ nguyên số 0 ở đầu mã/số điện thoại.
    dataframe = dataframe.fillna("")
    for column in dataframe.columns:
        dataframe[column] = dataframe[column].astype(str).map(str.strip)

    # Bỏ hàng/cột hoàn toàn trống.
    dataframe = dataframe.loc[
        ~dataframe.apply(
            lambda row: all(not str(value).strip() for value in row),
            axis=1,
        )
    ]
    dataframe = dataframe.loc[
        :,
        ~dataframe.apply(
            lambda column: all(not str(value).strip() for value in column),
            axis=0,
        ),
    ]

    if dataframe.empty:
        return dataframe.reset_index(drop=True)

    # CSV hoặc bảng đã có tên cột hợp lệ: chỉ chuẩn hóa tên cột như trước.
    existing_columns = [canonical_column_name(column) for column in dataframe.columns]
    existing_canonical = set(existing_columns)
    if "Họ và tên" in existing_canonical or "Tên công trình" in existing_canonical:
        dataframe.columns = existing_columns
    else:
        # Excel được đọc với header=None. Tìm hàng tiêu đề chính trong 20 hàng đầu.
        header_terms = {
            "ho va ten",
            "chuc vu",
            "que quan",
            "trinh do chuyen mon",
            "trinh do llct",
            "ngay vao dang",
            "cccd",
            "so so bhxh",
            "dien thoai",
            "ten cong trinh",
            "dia diem",
            "loai cong trinh",
            "nam ht",
            "wtb",
            "f tuoi",
            "tram bom",
            "kenh",
            "ke",
            "de",
        }
        header_index: int | None = None
        best_score = 0

        for position in range(min(len(dataframe), 20)):
            values = [
                normalize_table_text(value)
                for value in dataframe.iloc[position].tolist()
            ]
            score = sum(
                1
                for term in header_terms
                if any(term == value or term in value for value in values)
            )
            if score > best_score:
                best_score = score
                header_index = position

        if header_index is None or best_score < 2:
            # Không nhận diện chắc chắn được tiêu đề thì không đoán cấu trúc bảng.
            dataframe.columns = existing_columns
        else:
            parent_headers = [
                str(value).strip()
                for value in dataframe.iloc[header_index].tolist()
            ]

            # Ô gộp theo chiều ngang tạo các ô trống; điền tên nhóm sang phải.
            last_parent = ""
            for index, value in enumerate(parent_headers):
                if value:
                    last_parent = value
                elif last_parent:
                    parent_headers[index] = last_parent

            child_headers = [""] * len(parent_headers)
            data_start_index = header_index + 1

            if header_index + 1 < len(dataframe):
                candidate_children = [
                    str(value).strip()
                    for value in dataframe.iloc[header_index + 1].tolist()
                ]
                child_tokens = {
                    normalize_table_text(value)
                    for value in candidate_children
                    if str(value).strip()
                }
                # Bảng nhiều tầng: ngoài Nam/Nữ còn có TK/TT, thông số đập...
                generic_child_tokens = {
                    "nam", "nu", "tk", "tt", "hmax", "cao trinh",
                    "tcs", "don vi", "km", "ha", "m", "m3",
                }
                non_empty_children = [token for token in child_tokens if token]
                if child_tokens.intersection(generic_child_tokens) or len(non_empty_children) >= 2:
                    child_headers = candidate_children
                    data_start_index = header_index + 2

            normalized_columns: list[str] = []
            for index, parent in enumerate(parent_headers):
                child = child_headers[index] if index < len(child_headers) else ""
                parent = str(parent).strip()
                child = str(child).strip()

                if child and normalize_table_text(child) not in {
                    normalize_table_text(parent),
                    "stt",
                }:
                    column_name = f"{parent} - {child}" if parent else child
                else:
                    column_name = parent or child or f"Cột {index + 1}"

                normalized = normalize_table_text(column_name)

                # Chuẩn hóa chính xác các trường nhạy cảm, không suy đoán theo vị trí cột.
                if "ho va ten" in normalized or normalized == "ho ten":
                    canonical = "Họ và tên"
                elif normalized == "cccd" or normalized == "cmnd" or "can cuoc" in normalized:
                    canonical = "CCCD"
                elif (
                    "so so bhxh" in normalized
                    or normalized == "bhxh"
                    or "bao hiem xa hoi" in normalized
                ):
                    canonical = "Số sổ BHXH"
                elif (
                    "so the bhyt" in normalized
                    or normalized == "bhyt"
                    or "bao hiem y te" in normalized
                ):
                    canonical = "Số thẻ BHYT"
                elif "dien thoai" in normalized or normalized == "sdt":
                    canonical = "Số điện thoại"
                elif "ngay thang nam cap" in normalized or normalized == "ngay cap":
                    canonical = "Ngày cấp CCCD"
                elif "ngay vao dang" in normalized:
                    canonical = "Ngày vào Đảng"
                elif "trinh do chuyen mon" in normalized:
                    canonical = "Trình độ chuyên môn"
                elif "trinh do llct" in normalized or "ly luan chinh tri" in normalized:
                    canonical = "Trình độ LLCT"
                elif "chuc vu" in normalized or "chuc danh" in normalized:
                    canonical = "Chức vụ"
                elif "que quan" in normalized:
                    canonical = "Quê quán"
                elif "ngay thang nam sinh" in normalized or normalized == "ngay sinh":
                    canonical = "Ngày sinh"
                elif "ten cong trinh" in normalized or normalized in {"ten ct", "cong trinh"}:
                    canonical = "Tên công trình"
                elif "dia diem" in normalized or "xa phuong dac khu" in normalized:
                    canonical = "Địa điểm"
                elif "loai cong trinh" in normalized or "nhom cong trinh" in normalized:
                    canonical = "Loại công trình"
                elif normalized in {"stt", "tt", "so thu tu"}:
                    canonical = "STT"
                else:
                    canonical = canonical_column_name(column_name)

                normalized_columns.append(canonical)

            dataframe = dataframe.iloc[data_start_index:].copy()
            dataframe.columns = normalized_columns

    # Tạo tên duy nhất nếu có cột trùng nhau.
    seen: dict[str, int] = {}
    unique_columns: list[str] = []
    for column in dataframe.columns:
        column = str(column).strip() or "Cột"
        count = seen.get(column, 0)
        seen[column] = count + 1
        unique_columns.append(column if count == 0 else f"{column}_{count + 1}")
    dataframe.columns = unique_columns

    for column in dataframe.columns:
        dataframe[column] = dataframe[column].fillna("").astype(str).map(str.strip)

    # Loại các dòng phân nhóm, tiêu đề lặp và dòng trống; không loại dữ liệu thật.
    if "Họ và tên" in dataframe.columns:
        excluded_names = {
            "",
            "ho va ten",
            "cong chuc",
            "vien chuc",
            "hop dong",
            "nguoi lao dong",
        }
        dataframe = dataframe[
            ~dataframe["Họ và tên"].map(normalize_table_text).isin(excluded_names)
        ]

    return dataframe.reset_index(drop=True)


def read_table_file(file_path: Path) -> dict[str, Any]:
    if not PANDAS_AVAILABLE:
        raise RuntimeError("Chưa cài pandas/openpyxl.")

    if file_path.suffix.lower() == ".csv":
        last_error: Exception | None = None
        for encoding in ("utf-8-sig", "utf-8", "cp1258", "latin1"):
            try:
                dataframe = pd.read_csv(
                    file_path,
                    dtype=str,
                    keep_default_na=False,
                    encoding=encoding,
                )
                return {"CSV": clean_table_dataframe(dataframe)}
            except Exception as error:
                last_error = error
        raise RuntimeError(f"Không đọc được CSV: {last_error}")

    if file_path.suffix.lower() == ".xlsx":
        sheets = pd.read_excel(
            file_path,
            sheet_name=None,
            header=None,
            dtype=str,
            keep_default_na=False,
            engine="openpyxl",
        )
        return {
            str(name): clean_table_dataframe(frame)
            for name, frame in sheets.items()
        }

    raise RuntimeError("Chỉ nhận file .xlsx hoặc .csv.")


def save_table_file(
    database: dict[str, Any],
    uploaded_file: Any,
) -> dict[str, Any]:
    original_name = Path(uploaded_file.name).name
    suffix = Path(original_name).suffix.lower()
    if suffix not in {".xlsx", ".csv"}:
        raise RuntimeError("Chỉ nhận file Excel .xlsx hoặc CSV.")

    TABLE_DATA_DIR.mkdir(parents=True, exist_ok=True)
    stored_name = (
        f"{uuid.uuid4().hex[:10]}_"
        + re.sub(r'[<>:"/\\\\|?*]+', "_", original_name)
    )
    stored_path = TABLE_DATA_DIR / stored_name
    stored_path.write_bytes(uploaded_file.getbuffer())

    sheets = read_table_file(stored_path)
    sheet_stats = {
        str(name): {
            "rows": int(len(frame)),
            "columns": [str(column) for column in frame.columns],
        }
        for name, frame in sheets.items()
    }
    non_empty_sheets = [
        name for name, frame in sheets.items()
        if frame is not None and not getattr(frame, "empty", True)
    ]
    if not non_empty_sheets:
        stored_path.unlink(missing_ok=True)
        raise RuntimeError(
            "File đã mở được nhưng không nhận diện được dòng dữ liệu. "
            "Hãy kiểm tra hàng tiêu đề hoặc gửi file để chỉnh bộ đọc bảng."
        )

    # Cùng một tên file: xóa bản cũ trước khi lưu bản mới để tránh Agent đọc nhầm
    # bản đã nạp trước đó (ví dụ còn sheet SGV trong khi file mới là Sheet1).
    old_items = [
        item for item in database.get("table_files", [])
        if str(item.get("name", "")).casefold() == original_name.casefold()
    ]
    for old_item in old_items:
        try:
            resolve_table_path(old_item).unlink(missing_ok=True)
        except Exception:
            pass
    database["table_files"] = [
        item for item in database.get("table_files", [])
        if str(item.get("name", "")).casefold() != original_name.casefold()
    ]

    metadata = {
        "id": uuid.uuid4().hex,
        "name": original_name,
        "stored_name": stored_name,
        "uploaded_at": now_text(),
        "sheet_names": list(sheets.keys()),
        "sheet_stats": sheet_stats,
    }
    database.setdefault("table_files", []).append(metadata)
    save_database(database)
    return metadata


def resolve_table_path(file_info: dict[str, Any]) -> Path:
    return TABLE_DATA_DIR / str(file_info.get("stored_name", ""))


def delete_table_file(database: dict[str, Any], table_file_id: str) -> None:
    target = next(
        (
            item for item in database.get("table_files", [])
            if str(item.get("id", "")) == table_file_id
        ),
        None,
    )
    if target:
        resolve_table_path(target).unlink(missing_ok=True)

    database["table_files"] = [
        item for item in database.get("table_files", [])
        if str(item.get("id", "")) != table_file_id
    ]
    save_database(database)



def detect_requested_structure_types(question: str) -> list[str]:
    """Nhận diện loại công trình; các loại đều thuộc khái niệm chung công trình."""
    normalized = normalize_table_text(question)
    detected: list[str] = []
    # Ưu tiên loại cụ thể dài hơn để “đập ngăn mặn” không bị rút thành “đập”.
    candidates: list[tuple[str, str]] = []
    for structure_type, aliases in STRUCTURE_TYPE_ALIASES.items():
        for alias in aliases:
            candidates.append((structure_type, normalize_table_text(alias)))
    for structure_type, alias in sorted(candidates, key=lambda item: len(item[1]), reverse=True):
        if alias and re.search(rf"(^|\s){re.escape(alias)}(?=\s|$)", normalized):
            if structure_type not in detected:
                detected.append(structure_type)

    # Loại tổng quát: “công trình” mà không nêu loại cụ thể nghĩa là toàn bộ công trình.
    if not detected and any(term in normalized for term in ("cong trinh", "cong trinh thuy loi")):
        return ["Công trình"]
    return detected


def classify_structure_sheet(sheet_name: str, dataframe: Any) -> str:
    """Xác định loại công trình bằng tên sheet, cột Loại công trình và bộ cột đặc trưng."""
    normalized_sheet = normalize_table_text(sheet_name)
    for structure_type, hints in STRUCTURE_SHEET_HINTS.items():
        for hint in sorted(hints, key=len, reverse=True):
            norm_hint = normalize_table_text(hint)
            if re.search(rf"(^|\s){re.escape(norm_hint)}(?=\s|$)", normalized_sheet):
                return structure_type

    columns = {normalize_table_text(column) for column in getattr(dataframe, "columns", [])}
    joined = " ".join(sorted(columns))
    if any(term in joined for term in ("wtb", "f tuoi", "dung tich", "hmax", "cao trinh dinh dap")):
        return "Hồ chứa"
    if "cong suat may" in joined or "luu luong bom" in joined:
        return "Trạm bơm"
    if "chieu dai kenh" in joined or "luu luong thiet ke" in joined:
        return "Kênh"
    return "Công trình"


def _resolve_structure_name_column(dataframe: Any) -> str | None:
    column = _resolve_structured_column(dataframe, "Tên công trình")
    if column:
        return column
    # Dự phòng an toàn cho bảng cũ chưa chuẩn hóa tên cột.
    for candidate in dataframe.columns:
        normalized = normalize_table_text(candidate)
        if "ten cong trinh" in normalized or normalized in {"ten ct", "cong trinh"}:
            return str(candidate)
    return None


def _is_valid_structure_name(value: Any) -> bool:
    normalized = normalize_table_text(value)
    if not normalized:
        return False
    excluded = {
        "ten cong trinh", "cong trinh", "tong cong", "tong so", "cong",
        "stt", "tt", "ghi chu", "danh muc", "loai cong trinh",
    }
    if normalized in excluded or normalized.startswith("tong "):
        return False
    return True


def lookup_infrastructure_table(database: dict[str, Any], question: str) -> str | None:
    """Tra cứu chung hồ chứa, đập, đê, kè, mỏ hàn, kênh, trạm bơm..."""
    requested_types = detect_requested_structure_types(question)
    if not requested_types or not database.get("table_files"):
        return None

    normalized_question = normalize_table_text(question)
    records: list[dict[str, Any]] = []

    for file_info in database.get("table_files", []):
        file_path = resolve_table_path(file_info)
        if not file_path.exists():
            continue
        try:
            sheets = read_table_file(file_path)
        except Exception:
            continue

        for sheet_name, dataframe in sheets.items():
            if dataframe is None or getattr(dataframe, "empty", True):
                continue
            inferred_type = classify_structure_sheet(str(sheet_name), dataframe)
            type_column = _resolve_structured_column(dataframe, "Loại công trình")
            name_column = _resolve_structure_name_column(dataframe)
            if not name_column:
                continue

            for row_index, raw_name in dataframe[name_column].items():
                name = str(raw_name).strip()
                if not _is_valid_structure_name(name):
                    continue
                row_type = inferred_type
                if type_column:
                    raw_type = str(dataframe.at[row_index, type_column]).strip()
                    detected_row_types = detect_requested_structure_types(raw_type)
                    if detected_row_types:
                        row_type = detected_row_types[0]

                if requested_types != ["Công trình"] and row_type not in requested_types:
                    continue

                row_values = {
                    str(column): str(dataframe.at[row_index, column]).strip()
                    for column in dataframe.columns
                    if str(dataframe.at[row_index, column]).strip()
                    and normalize_table_text(dataframe.at[row_index, column]) not in {"nan", "none"}
                }
                records.append({
                    "file": str(file_info.get("name", file_path.name)),
                    "sheet": str(sheet_name),
                    "row": int(row_index) + 1,
                    "type": row_type,
                    "name": name,
                    "values": row_values,
                })

    if not records:
        type_text = ", ".join(requested_types).lower()
        return (
            f"Chưa tìm thấy dữ liệu **{type_text}** trong kho bảng dữ liệu.\n\n"
            "**Khuyến cáo kiểm tra:** Chọn đúng file/sheet và kiểm tra bảng có cột "
            "`Tên công trình` hoặc tên sheet thể hiện loại công trình."
        )

    # Loại trùng cùng tên, cùng loại trong cùng nguồn.
    unique: list[dict[str, Any]] = []
    seen: set[tuple[str, str, str, str]] = set()
    for item in records:
        key = (
            item["file"], item["sheet"], item["type"], normalize_table_text(item["name"])
        )
        if key not in seen:
            seen.add(key)
            unique.append(item)
    records = unique

    # Nếu câu hỏi chứa chính xác tên công trình, trả thông số của dòng đó.
    named_matches = [
        item for item in records
        if normalize_table_text(item["name"]) in normalized_question
    ]
    if named_matches:
        item = max(named_matches, key=lambda x: len(normalize_table_text(x["name"])))
        ignored = {"STT", "Tên công trình", "Cột 1", "Cột 2"}
        rows = []
        for field, value in item["values"].items():
            if field in ignored or normalize_table_text(field).startswith("cot "):
                continue
            rows.append(f"| {field} | {value} |")
        details = "\n".join(rows[:24]) or "| Loại công trình | " + item["type"] + " |"
        return (
            f"**{item['name']}** là **{item['type'].lower()}** trong kho bảng dữ liệu.\n\n"
            "| Thông tin | Giá trị |\n|---|---|\n"
            f"| Loại công trình | {item['type']} |\n{details}\n\n"
            "**Nguồn bảng dữ liệu:**\n"
            f"- `{item['file']}` — sheet `{item['sheet']}`, dòng dữ liệu {item['row']}."
        )

    asks_count = any(term in normalized_question for term in (
        "bao nhieu", "tong so", "so luong", "dem", "co may"
    ))
    asks_list = any(term in normalized_question for term in (
        "danh sach", "liet ke", "gom nhung", "nhung cong trinh nao"
    ))

    type_label = "công trình" if requested_types == ["Công trình"] else ", ".join(requested_types).lower()
    sources = []
    for item in records:
        source = f"`{item['file']}` — sheet `{item['sheet']}`"
        if source not in sources:
            sources.append(source)

    if asks_count and not asks_list:
        return (
            f"Có **{len(records)} {type_label}** trong kho bảng dữ liệu.\n\n"
            "**Nguồn bảng dữ liệu:**\n- " + "\n- ".join(sources)
        )

    table_rows = "\n".join(
        f"| {index} | {item['name']} | {item['type']} | {item['sheet']} |"
        for index, item in enumerate(records, start=1)
    )
    return (
        f"Tìm thấy **{len(records)} {type_label}**.\n\n"
        "| STT | Tên công trình | Loại công trình | Sheet nguồn |\n"
        "|---:|---|---|---|\n"
        f"{table_rows}\n\n"
        "**Nguồn bảng dữ liệu:**\n- " + "\n- ".join(sources)
    )


def detect_requested_table_field(question: str) -> str:
    """Nhận diện trường cần tra, kể cả cách viết đầy đủ bằng tiếng Việt."""
    normalized = normalize_table_text(question)
    mapping = {
        "CCCD": [
            "cccd", "cmnd", "can cuoc", "can cuoc cong dan",
            "so can cuoc", "so can cuoc cong dan",
        ],
        "Số sổ BHXH": [
            "bhxh", "so so bhxh", "so bhxh", "ma so bhxh",
            "so so bao hiem xa hoi", "so bao hiem xa hoi",
            "ma so bao hiem xa hoi", "bao hiem xa hoi",
        ],
        "Số thẻ BHYT": [
            "bhyt", "so the bhyt", "so bhyt", "ma the bhyt",
            "so the bao hiem y te", "so bao hiem y te",
            "bao hiem y te",
        ],
        "Số điện thoại": [
            "so dien thoai", "dien thoai", "sdt", "so lien lac",
        ],
        "Ngày sinh": ["ngay sinh", "sinh ngay", "nam sinh"],
        "Chức vụ": ["chuc vu", "chuc danh"],
        "Quê quán": ["que quan"],
        "Trình độ chuyên môn": ["trinh do chuyen mon", "chuyen mon"],
        "Trình độ LLCT": [
            "trinh do llct", "ly luan chinh tri", "llct",
        ],
        "Ngày vào Đảng": ["ngay vao dang", "vao dang"],
    }
    # Ưu tiên cụm dài để tránh từ ngắn khớp nhầm.
    for field, terms in mapping.items():
        if any(term in normalized for term in sorted(terms, key=len, reverse=True)):
            return field
    return ""


def _resolve_structured_column(dataframe: Any, requested_field: str) -> str | None:
    """Tìm đúng cột cần tra bằng tên chuẩn và các bí danh, không dựa vào vị trí cột."""
    if requested_field in dataframe.columns:
        return requested_field

    aliases = TABLE_COLUMN_ALIASES.get(requested_field, []) + [requested_field]
    normalized_aliases = {normalize_table_text(value) for value in aliases}

    for column in dataframe.columns:
        normalized_column = normalize_table_text(column)
        if normalized_column in normalized_aliases:
            return str(column)

    # Chấp nhận tên cột có thêm chú thích, nhưng chỉ khi bí danh đủ rõ.
    for column in dataframe.columns:
        normalized_column = normalize_table_text(column)
        for alias in normalized_aliases:
            if len(alias) >= 4 and (alias in normalized_column or normalized_column in alias):
                return str(column)

    return None


def _normalize_person_text(value: Any) -> str:
    """Chuẩn hóa tên nhưng GIỮ dấu tiếng Việt để không nhầm Dung/Dũng."""
    text = unicodedata.normalize("NFC", str(value or "")).casefold()
    text = re.sub(r"[^0-9a-zà-ỹđ\s]", " ", text)
    return " ".join(text.split())


def _strip_person_honorifics(text: str) -> str:
    """Bỏ cách xưng hô, không bỏ các thành phần họ tên."""
    honorifics = {
        "ông", "bà", "anh", "chị", "cô", "chú", "bác", "em",
        "đồng chí", "ông ấy", "bà ấy",
    }
    normalized = _normalize_person_text(text)
    for phrase in sorted(honorifics, key=len, reverse=True):
        normalized = re.sub(
            rf"(^|\s){re.escape(phrase)}(?=\s|$)", " ", normalized
        )
    return " ".join(normalized.split())


def _person_match_score(question: str, person_name: Any) -> int:
    """
    So khớp tên theo nguyên tắc an toàn:
    - Giữ dấu tiếng Việt: Dung khác Dũng.
    - Họ tên đầy đủ/cụm từ 2 tiếng được ưu tiên.
    - Một tiếng cuối chỉ khớp khi đúng dấu; nếu nhiều người cùng tên,
      lookup_structured_table sẽ yêu cầu làm rõ, không tự chọn.
    - Chỉ cho phép so khớp không dấu khi người dùng nhập ít nhất 2 tiếng tên.
    """
    question_with_accents = _strip_person_honorifics(question)
    name_with_accents = _normalize_person_text(person_name)
    if not name_with_accents:
        return 0

    if name_with_accents in question_with_accents:
        return 3000 + len(name_with_accents)

    name_tokens = [token for token in name_with_accents.split() if len(token) >= 2]
    question_tokens = set(question_with_accents.split())
    if not name_tokens:
        return 0

    exact_tokens = [token for token in name_tokens if token in question_tokens]

    # Khớp chính xác từ 2 thành phần tên trở lên, vẫn giữ dấu.
    if len(exact_tokens) >= 2:
        tail_two = " ".join(name_tokens[-2:])
        tail_bonus = 500 if tail_two in question_with_accents else 0
        coverage = int(200 * len(exact_tokens) / len(name_tokens))
        return 1200 + coverage + tail_bonus + len(exact_tokens) * 20

    # Chỉ một tiếng: bắt buộc là tiếng cuối và phải đúng dấu tuyệt đối.
    if len(exact_tokens) == 1 and exact_tokens[0] == name_tokens[-1]:
        return 300

    # Hỗ trợ câu hỏi không dấu, nhưng phải có ít nhất 2 tiếng tên để an toàn.
    accentless_question = normalize_table_text(question)
    accentless_name = normalize_table_text(person_name)
    accentless_name_tokens = [t for t in accentless_name.split() if len(t) >= 2]
    accentless_question_tokens = set(accentless_question.split())
    accentless_matches = [
        t for t in accentless_name_tokens if t in accentless_question_tokens
    ]
    if len(accentless_matches) >= 2:
        coverage = int(100 * len(accentless_matches) / len(accentless_name_tokens))
        return 700 + coverage + len(accentless_matches) * 10

    return 0

def lookup_structured_table(
    database: dict[str, Any],
    question: str,
) -> str | None:
    """Ưu tiên tra cứu dữ liệu có cấu trúc trước khi gọi Vector Store/OpenAI."""
    infrastructure_answer = lookup_infrastructure_table(database, question)
    if infrastructure_answer:
        return infrastructure_answer

    requested_field = detect_requested_table_field(question)
    if not requested_field or not database.get("table_files"):
        return None

    matches: list[dict[str, Any]] = []

    for file_info in database.get("table_files", []):
        file_path = resolve_table_path(file_info)
        if not file_path.exists():
            continue

        try:
            sheets = read_table_file(file_path)
        except Exception:
            continue

        for sheet_name, dataframe in sheets.items():
            if dataframe is None or getattr(dataframe, "empty", True):
                continue

            name_column = _resolve_structured_column(dataframe, "Họ và tên")
            value_column = _resolve_structured_column(dataframe, requested_field)
            if not name_column or not value_column:
                continue

            for row_index, raw_name in dataframe[name_column].items():
                score = _person_match_score(question, raw_name)
                if score <= 0:
                    continue

                value = str(dataframe.at[row_index, value_column]).strip()
                if not value or normalize_table_text(value) in {"nan", "none"}:
                    continue

                matches.append(
                    {
                        "file": str(file_info.get("name", file_path.name)),
                        "sheet": str(sheet_name),
                        "row": int(row_index) + 1,
                        "person": str(raw_name).strip(),
                        "field": requested_field,
                        "column": str(value_column),
                        "value": value,
                        "score": score,
                    }
                )

    if not matches:
        sensitive_fields = {
            "CCCD", "Số sổ BHXH", "Số thẻ BHYT", "Số điện thoại"
        }
        if requested_field in sensitive_fields and database.get("table_files"):
            return (
                f"Chưa tìm thấy **{requested_field}** tương ứng trong kho bảng dữ liệu.\n\n"
                "**Khuyến cáo kiểm tra:** Mở mục *Kho bảng dữ liệu*, chọn đúng "
                "file/sheet và kiểm tra tên người, tên cột. Hệ thống không chuyển "
                "sang suy đoán từ PDF đối với trường thông tin nhạy cảm này."
            )
        return None

    # Chỉ giữ nhóm khớp tên tốt nhất để tránh chọn nhầm người có tên gần giống.
    best_score = max(item["score"] for item in matches)
    matches = [item for item in matches if item["score"] == best_score]

    unique: list[dict[str, Any]] = []
    seen: set[tuple[str, str, str, str]] = set()
    for item in matches:
        key = (
            item["file"],
            item["sheet"],
            normalize_table_text(item["person"]),
            item["value"],
        )
        if key not in seen:
            seen.add(key)
            unique.append(item)

    if len(unique) > 1:
        people = {normalize_table_text(item["person"]) for item in unique}
        values = {item["value"] for item in unique}

        # Cùng một người, cùng một giá trị xuất hiện ở nhiều sheet/file: trả một kết quả.
        if len(people) == 1 and len(values) == 1:
            unique = [unique[0]]
        else:
            options = "\n".join(
                f"- {item['person']} — `{item['file']}`, "
                f"sheet `{item['sheet']}`, cột `{item['column']}`"
                for item in unique[:8]
            )
            return (
                "Tìm thấy nhiều bản ghi phù hợp nên chưa thể tự chọn.\n\n"
                f"{options}\n\n"
                "**Khuyến cáo kiểm tra:** Bổ sung ngày sinh, chức vụ hoặc đơn vị "
                "để xác định đúng người."
            )

    item = unique[0]
    return (
        f"**{item['field']} của {item['person']}:** {item['value']}\n\n"
        "**Nguồn bảng dữ liệu:**\n"
        f"- `{item['file']}` — sheet `{item['sheet']}`, "
        f"cột `{item['column']}`."
    )


# =========================================================
# OPENAI
# =========================================================
def get_api_key() -> str:
    try:
        secret_key = st.secrets.get("OPENAI_API_KEY", "")
    except Exception:
        secret_key = ""
    return secret_key or os.getenv("OPENAI_API_KEY", "")


def get_configured_vector_store_id() -> str:
    """Đọc ID kho tài liệu cố định từ Streamlit Secrets hoặc biến môi trường."""
    try:
        secret_id = st.secrets.get("OPENAI_VECTOR_STORE_ID", "")
    except Exception:
        secret_id = ""

    return (
        str(secret_id).strip()
        or os.getenv("OPENAI_VECTOR_STORE_ID", "").strip()
        or DEFAULT_VECTOR_STORE_ID
    )


@st.cache_resource(show_spinner=False)
def create_openai_client(api_key: str) -> OpenAI:
    return OpenAI(
        api_key=api_key,
        timeout=120.0,
        max_retries=2,
    )


def get_client() -> OpenAI:
    api_key = get_api_key()
    if not api_key:
        raise RuntimeError(
            "Chưa khai báo OPENAI_API_KEY trong biến môi trường "
            "hoặc file .streamlit/secrets.toml."
        )
    return create_openai_client(api_key)


def ensure_vector_store(client: OpenAI, database: dict[str, Any]) -> str:
    configured_vector_store_id = get_configured_vector_store_id()
    if configured_vector_store_id:
        database["vector_store_id"] = configured_vector_store_id
        save_database(database)
        return configured_vector_store_id

    vector_store_id = database.get("vector_store_id", "").strip()
    if vector_store_id:
        return vector_store_id

    vector_store = client.vector_stores.create(
        name="Kho tài liệu CCTL_QNG",
        description="Kho tài liệu phục vụ công tác tham mưu Chi cục Thủy lợi.",
    )
    database["vector_store_id"] = vector_store.id
    save_database(database)
    return vector_store.id


def upload_document(
    client: OpenAI,
    database: dict[str, Any],
    uploaded_file: Any,
) -> None:
    """
    Tải tài liệu lên OpenAI và giữ nguyên tên file gốc.

    Trước đây NamedTemporaryFile tạo tên dạng tmpxxxx.doc, nên OpenAI lưu luôn
    tên tạm đó. Bản này tạo một thư mục tạm nhưng tên file bên trong vẫn là
    tên người dùng đã tải lên.
    """
    vector_store_id = ensure_vector_store(client, database)

    original_name = Path(uploaded_file.name).name
    safe_name = re.sub(r'[<>:"/\\\\|?*]+', "_", original_name).strip()
    if not safe_name:
        safe_name = f"tai_lieu_{uuid.uuid4().hex[:8]}.bin"

    with tempfile.TemporaryDirectory() as temp_dir:
        temp_path = Path(temp_dir) / safe_name
        temp_path.write_bytes(uploaded_file.getbuffer())

        with temp_path.open("rb") as file_handle:
            openai_file = client.files.create(
                file=file_handle,
                purpose="assistants",
            )

        client.vector_stores.files.create_and_poll(
            vector_store_id=vector_store_id,
            file_id=openai_file.id,
        )

    database["uploaded_files"].append(
        {
            "name": original_name,
            "openai_file_id": openai_file.id,
            "uploaded_at": now_text(),
        }
    )
    save_database(database)


def build_api_input(messages: list[dict[str, Any]]) -> list[dict[str, str]]:
    return [
        {
            "role": message["role"],
            "content": message["content"],
        }
        for message in messages[-12:]
        if message["role"] in {"user", "assistant"}
    ]


def question_requests_documents(question: str) -> bool:
    """Tự bật tra cứu kho khi câu hỏi rõ ràng yêu cầu dùng tài liệu."""
    normalized = " ".join(question.lower().split())
    terms = [
        "theo tài liệu", "tài liệu dùng chung", "trong kho", "theo hồ sơ",
        "theo văn bản", "quyết định", "quy trình", "quy định", "biên bản",
        "báo cáo", "phụ lục", "file pdf", "file word", "tài liệu đã nạp",
        "căn cứ tài liệu", "tra cứu", "liên hồ",
    ]
    return any(term in normalized for term in terms)


def search_vector_store_context(
    client: OpenAI,
    vector_store_id: str,
    question: str,
    filename_map: dict[str, str] | None = None,
) -> tuple[str, list[str], list[dict[str, Any]]]:
    """
    Tìm trực tiếp trong kho bằng nhiều cách diễn đạt.
    Trả về:
    - các đoạn liên quan;
    - tên tài liệu;
    - danh sách file ứng viên để có thể đọc toàn bộ khi cần.
    """
    expanded_queries = [
        question,
        (
            question
            + " họ tên tên cá nhân chức danh chức vụ lãnh đạo "
              "Chi cục trưởng Phó Chi cục trưởng trưởng phòng phó trưởng phòng"
        ),
        (
            "danh sách cán bộ công chức viên chức; họ và tên; chức vụ; "
            "chức danh; đơn vị công tác; cơ cấu tổ chức Chi cục Thủy lợi"
        ),
    ]

    collected: list[tuple[float, str, str, str]] = []
    seen_chunks: set[str] = set()
    search_errors: list[str] = []

    for query in expanded_queries:
        try:
            page = client.vector_stores.search(
                vector_store_id=vector_store_id,
                query=query,
                max_num_results=20,
                rewrite_query=True,
            )
        except Exception as error:
            search_errors.append(str(error))
            continue

        for result in getattr(page, "data", []) or []:
            file_id = getattr(result, "file_id", "") or ""
            raw_filename = (
                getattr(result, "filename", "")
                or "Tài liệu không rõ tên"
            )
            filename = (
                (filename_map or {}).get(file_id)
                or raw_filename
            )
            score = float(getattr(result, "score", 0.0) or 0.0)

            parts: list[str] = []
            for item in getattr(result, "content", []) or []:
                item_text = getattr(item, "text", "") or ""
                if item_text.strip():
                    parts.append(item_text.strip())

            chunk_text = "\n".join(parts).strip()
            if not chunk_text:
                continue

            fingerprint = f"{file_id}|{chunk_text[:300]}"
            if fingerprint in seen_chunks:
                continue

            seen_chunks.add(fingerprint)
            collected.append((score, file_id, filename, chunk_text))

    collected.sort(key=lambda row: row[0], reverse=True)
    selected = collected[:20]

    if not selected:
        if search_errors:
            unique_errors = list(dict.fromkeys(search_errors))
            raise RuntimeError(
                "Không gọi được chức năng tìm kiếm Vector Store: "
                + " | ".join(unique_errors[:2])
            )
        return "", [], []

    filenames: list[str] = []
    candidates_by_file: dict[str, dict[str, Any]] = {}
    context_blocks: list[str] = []

    for index, (score, file_id, filename, chunk_text) in enumerate(selected, start=1):
        if filename not in filenames:
            filenames.append(filename)

        if file_id:
            current = candidates_by_file.get(file_id)
            if current is None or score > current["score"]:
                candidates_by_file[file_id] = {
                    "file_id": file_id,
                    "filename": filename,
                    "score": score,
                }

        context_blocks.append(
            f"[Kết quả {index} | Tài liệu: {filename} | Điểm: {score:.3f}]\n"
            f"{chunk_text}"
        )

    candidates = sorted(
        candidates_by_file.values(),
        key=lambda item: item["score"],
        reverse=True,
    )

    return "\n\n".join(context_blocks), filenames, candidates


def should_read_full_document(question: str) -> bool:
    """
    Các câu hỏi cần rà soát đầy đủ thường dễ bị bỏ sót nếu chỉ dùng vài đoạn RAG.
    """
    normalized = " ".join(question.lower().split())
    trigger_terms = [
        "ai ",
        "là ai",
        "họ tên",
        "tên và chức danh",
        "chức danh",
        "chức vụ",
        "danh sách",
        "liệt kê",
        "từng cá nhân",
        "từng người",
        "bao nhiêu",
        "mấy ",
        "toàn bộ",
        "đầy đủ",
        "cuối tài liệu",
        "phía sau",
        "trong bảng",
        "phụ lục",
        "chức năng nhiệm vụ",
        "chức năng, nhiệm vụ",
        "cơ cấu tổ chức",
        "các phòng",
        "các đơn vị",
        "theo quy định",
        "theo tài liệu",
        "theo văn bản",
        "quy trình",
        "gồm những",
        "có những",
        "cccd",
        "cmnd",
        "căn cước",
        "số điện thoại",
        "bhxh",
        "bhyt",
        "ngày sinh",
        "mã số",
        "số tài khoản",
    ]
    return any(term in normalized for term in trigger_terms)


def fetch_parsed_vector_file_content(
    api_key: str,
    vector_store_id: str,
    file_id: str,
) -> str:
    """
    Lấy toàn bộ phần nội dung đã được OpenAI phân tích của một file trong Vector Store.
    Đây là lớp dự phòng khi tìm kiếm theo đoạn có nguy cơ bỏ sót bảng hoặc danh sách.
    """
    url = (
        f"https://api.openai.com/v1/vector_stores/"
        f"{vector_store_id}/files/{file_id}/content"
    )

    response = httpx.get(
        url,
        headers={"Authorization": f"Bearer {api_key}"},
        timeout=60.0,
    )
    response.raise_for_status()
    payload = response.json()

    parts: list[str] = []
    for item in payload.get("content", []) or []:
        item_text = item.get("text", "")
        if item_text and item_text.strip():
            parts.append(item_text.strip())

    return "\n\n".join(parts)


def build_full_document_context(
    api_key: str,
    vector_store_id: str,
    candidates: list[dict[str, Any]],
    *,
    max_files: int = 2,
    max_total_chars: int = 120_000,
) -> tuple[str, list[str], list[str]]:
    """
    Đọc toàn bộ nội dung đã phân tích của tối đa vài file phù hợp nhất.
    Trả thêm danh sách lỗi để có thể chẩn đoán thay vì bỏ qua âm thầm.
    """
    blocks: list[str] = []
    filenames: list[str] = []
    errors: list[str] = []
    total_chars = 0

    for candidate in candidates[:max_files]:
        file_id = candidate.get("file_id", "")
        filename = candidate.get("filename", "Tài liệu không rõ tên")
        if not file_id:
            errors.append(f"{filename}: không có file_id.")
            continue

        try:
            full_text = fetch_parsed_vector_file_content(
                api_key,
                vector_store_id,
                file_id,
            )
        except Exception as error:
            errors.append(f"{filename}: {error}")
            continue

        if not full_text.strip():
            errors.append(
                f"{filename}: OpenAI không trả về phần chữ đã phân tích."
            )
            continue

        remaining = max_total_chars - total_chars
        if remaining <= 0:
            break

        clipped = full_text[:remaining]
        blocks.append(
            f"[TOÀN VĂN ĐÃ PHÂN TÍCH | Tài liệu: {filename}]\\n{clipped}"
        )
        filenames.append(filename)
        total_chars += len(clipped)

    return "\\n\\n".join(blocks), filenames, errors



SENSITIVE_FIELD_PATTERNS: dict[str, list[str]] = {
    "CCCD": [r"\bcccd\b", r"\bcăn cước\b", r"\bcmnd\b"],
    "Số sổ BHXH": [r"\bbhxh\b", r"\bsổ bảo hiểm xã hội\b"],
    "Số thẻ BHYT": [r"\bbhyt\b", r"\bthẻ bảo hiểm y tế\b"],
    "Số điện thoại": [r"\bsố điện thoại\b", r"\bđiện thoại\b", r"\bsđt\b"],
    "Số tài khoản": [r"\bsố tài khoản\b", r"\btài khoản ngân hàng\b"],
}


def detect_sensitive_field(question: str) -> str:
    normalized = question.lower()
    for field_name, patterns in SENSITIVE_FIELD_PATTERNS.items():
        if any(re.search(pattern, normalized, flags=re.IGNORECASE) for pattern in patterns):
            return field_name
    return ""


def has_explicit_labeled_value(context: str, field_name: str) -> bool:
    if not context or not field_name:
        return False

    label_patterns = {
        "CCCD": [r"(?:CCCD|CMND|Căn cước)\s*[:\-]?\s*(\d{9,12})"],
        "Số sổ BHXH": [r"(?:Số\s*sổ\s*BHXH|BHXH)\s*[:\-]?\s*(\d{8,15})"],
        "Số thẻ BHYT": [r"(?:Số\s*thẻ\s*BHYT|BHYT)\s*[:\-]?\s*([A-Z0-9]{8,20})"],
        "Số điện thoại": [r"(?:Số\s*điện\s*thoại|Điện\s*thoại|SĐT)\s*[:\-]?\s*(0\d{8,10})"],
        "Số tài khoản": [r"(?:Số\s*tài\s*khoản|Tài\s*khoản)\s*[:\-]?\s*(\d{6,20})"],
    }

    return any(
        re.search(pattern, context, flags=re.IGNORECASE)
        for pattern in label_patterns.get(field_name, [])
    )


def build_sensitive_field_warning(field_name: str, source_files: list[str]) -> str:
    source_lines = "\n".join(f"- {name}" for name in source_files) or "- Chưa xác định"
    return (
        f"Chưa thể xác nhận chính xác **{field_name}** từ tài liệu đã tra cứu. "
        "Bảng trong PDF có dấu hiệu bị tách hàng/cột khi trích xuất, nên việc "
        "gán một chuỗi số vào đúng cột có nguy cơ sai lệch.\n\n"
        "**Khuyến cáo kiểm tra:** Cần đối chiếu trực tiếp hàng của người được hỏi "
        f"với tiêu đề cột **{field_name}** trong file gốc trước khi sử dụng chính thức.\n\n"
        "**Nguồn văn bản trong kho:**\n"
        f"{source_lines}"
    )


def stream_openai_answer(
    client: OpenAI,
    database: dict[str, Any],
    messages: list[dict[str, Any]],
    *,
    use_file_search: bool,
    fast_mode: bool,
    deep_mode: bool = False,
):
    if deep_mode:
        model = DEEP_MODEL
    elif fast_mode:
        model = FAST_MODEL
    else:
        model = SEARCH_MODEL

    api_input = build_api_input(messages)
    instructions = SYSTEM_INSTRUCTIONS
    vector_store_id = ""
    diagnostics: dict[str, Any] = {
        "enabled": use_file_search,
        "model": model,
        "vector_store_id": "",
        "question": "",
        "manual_search_files": [],
        "candidate_files": [],
        "retrieved_context_chars": 0,
        "full_document_files": [],
        "full_document_chars": 0,
        "full_document_errors": [],
        "native_file_search_enabled": False,
    }

    if use_file_search:
        vector_store_id = (
            get_configured_vector_store_id()
            or database.get("vector_store_id", "").strip()
        )
        if not vector_store_id:
            raise RuntimeError(
                "Chưa xác định được OPENAI_VECTOR_STORE_ID cho kho tài liệu."
            )

        diagnostics["vector_store_id"] = vector_store_id

        latest_question = ""
        for message in reversed(messages):
            if message.get("role") == "user":
                latest_question = str(message.get("content", "")).strip()
                break

        diagnostics["question"] = latest_question

        original_filename_map = {
            str(item.get("openai_file_id", "")).strip(): str(
                item.get("name", "")
            ).strip()
            for item in database.get("uploaded_files", [])
            if item.get("openai_file_id") and item.get("name")
        }

        (
            retrieved_context,
            source_files,
            candidates,
        ) = search_vector_store_context(
            client,
            vector_store_id,
            latest_question,
            filename_map=original_filename_map,
        )

        diagnostics["manual_search_files"] = source_files
        diagnostics["candidate_files"] = [
            {
                "filename": item.get("filename", ""),
                "score": round(float(item.get("score", 0.0) or 0.0), 3),
            }
            for item in candidates
        ]
        diagnostics["retrieved_context_chars"] = len(retrieved_context)

        full_document_context = ""
        full_document_files: list[str] = []
        full_document_errors: list[str] = []

        sensitive_field = detect_sensitive_field(latest_question)

        if candidates and (deep_mode or sensitive_field):
            # Với trường dữ liệu nhạy cảm, luôn đọc toàn văn ít nhất 1 file.
            full_file_limit = (
                2 if deep_mode and should_read_full_document(latest_question) else 1
            )
            full_char_limit = (
                120_000 if full_file_limit == 2 else 80_000
            )

            (
                full_document_context,
                full_document_files,
                full_document_errors,
            ) = build_full_document_context(
                get_api_key(),
                vector_store_id,
                candidates,
                max_files=full_file_limit,
                max_total_chars=full_char_limit,
            )

        diagnostics["full_document_files"] = full_document_files
        diagnostics["full_document_chars"] = len(full_document_context)
        diagnostics["full_document_errors"] = full_document_errors

        combined_context_parts: list[str] = []
        if retrieved_context:
            combined_context_parts.append(
                "CÁC ĐOẠN TÌM KIẾM LIÊN QUAN:\\n" + retrieved_context
            )
        if full_document_context:
            combined_context_parts.append(
                "NỘI DUNG TOÀN FILE DỰ PHÒNG:\\n" + full_document_context
            )

        combined_context = "\\n\\n".join(combined_context_parts).strip()

        # Cơ chế khóa an toàn tuyệt đối cho dữ liệu nhạy cảm lấy từ PDF.
        # File Search thường làm phẳng bảng PDF nên có thể ghép sai cột dù tiêu đề
        # xuất hiện trong tài liệu. Vì vậy không cho phép trả số nhạy cảm từ PDF.
        all_sources: list[str] = []
        for name in source_files + full_document_files:
            if name and name not in all_sources:
                all_sources.append(name)

        has_pdf_source = any(
            str(name).lower().endswith(".pdf")
            for name in all_sources
        )

        if sensitive_field and has_pdf_source:
            st.session_state["rag_diagnostics"] = diagnostics
            yield (
                f"Chưa thể xác nhận chính xác **{sensitive_field}** từ file PDF này. "
                "Bảng PDF có thể bị làm phẳng khi trích xuất, dẫn đến ghép nhầm cột "
                "như CCCD, BHXH hoặc số điện thoại.\n\n"
                "**Khuyến cáo kiểm tra:** Không sử dụng số do Agent suy ra từ bảng PDF "
                "cho công việc chính thức. Hãy đối chiếu trực tiếp file gốc hoặc tải lên "
                "bản Excel/CSV có cấu trúc cột rõ ràng.\n\n"
                "**Nguồn văn bản trong kho:**\n"
                + ("\n".join(f"- {name}" for name in all_sources) or "- Chưa xác định")
            )
            return

        # Với nguồn không phải PDF, vẫn yêu cầu nhãn và giá trị nằm gần nhau.
        if sensitive_field and not has_explicit_labeled_value(
            combined_context,
            sensitive_field,
        ):
            st.session_state["rag_diagnostics"] = diagnostics
            yield build_sensitive_field_warning(
                sensitive_field,
                all_sources,
            )
            return

        if combined_context:
            api_input.append(
                {
                    "role": "user",
                    "content": (
                        "Dưới đây là kết quả tra cứu từ kho tài liệu. "
                        "Hãy ưu tiên dùng nội dung này để trả lời câu hỏi trước đó. "
                        "Kiểm tra kỹ bảng, danh sách, phần cuối tài liệu, họ tên, "
                        "chức danh, ngày tháng và các dòng liền kề.\\n\\n"
                        f"{combined_context}"
                    ),
                }
            )

            all_sources: list[str] = []
            for name in source_files + full_document_files:
                if name not in all_sources:
                    all_sources.append(name)

            source_text = ", ".join(all_sources)
            instructions += f"""
YÊU CẦU TRẢ LỜI THEO KHO TÀI LIỆU:
- Trả lời thẳng vào nội dung được hỏi, ưu tiên kết luận trước.
- Không kể lại quá trình tra cứu, không giải thích kỹ thuật tìm kiếm.
- Không chép dài nội dung tài liệu nếu người dùng không yêu cầu trích nguyên văn.
- Không suy đoán ngoài dữ liệu đã cung cấp.
- Nếu chưa đủ căn cứ, nêu đúng một câu ngắn về phần còn thiếu.
- Với dữ liệu bảng, phải kiểm tra đúng tiêu đề cột và đúng hàng của đối tượng trước khi trả lời.
- Tuyệt đối không suy giá trị CCCD, BHXH, BHYT, số điện thoại hoặc số tài khoản từ vị trí tương đối của các chuỗi số trong bảng PDF đã bị dàn phẳng.
- Nếu nhãn cột không nằm sát giá trị trong đoạn trích, phải từ chối xác nhận thay vì chọn một số gần đó.
- Với nguồn PDF dạng bảng, không được trả CCCD/CMND, BHXH, BHYT, số điện thoại hoặc số tài khoản; chỉ được hướng dẫn đối chiếu file gốc hoặc dùng Excel/CSV.
- Với CCCD/CMND, số điện thoại, BHXH, BHYT hoặc chuỗi số dài: nếu không thấy rõ tiêu đề cột thì không được khẳng định chắc chắn.
- Nếu có khả năng nhầm cột, phải thêm mục:

**Khuyến cáo kiểm tra:** Có khả năng dữ liệu bị lệch hoặc tách cột trong quá trình trích xuất; cần đối chiếu trực tiếp file gốc trước khi sử dụng chính thức.

- Cuối câu trả lời ghi đúng mục:

**Nguồn văn bản trong kho:**
{source_text}

- Chỉ giữ tên tài liệu thực sự đã dùng.
- Không tự đổi tên tài liệu và không dùng tên tạm nếu đã có tên gốc.
"""
        else:
            instructions += """

YÊU CẦU TRẢ LỜI THEO KHO TÀI LIỆU:
- Kết quả tìm kiếm thủ công chưa lấy được đoạn phù hợp.
- Bắt buộc sử dụng công cụ file_search để tìm trực tiếp trong Vector Store trước khi trả lời.
- Chỉ kết luận theo nội dung tìm được; không tự suy đoán.
"""

    if deep_mode:
        reasoning_effort = "medium"
        verbosity = "medium"
        max_output_tokens = 2800
    elif use_file_search:
        # File Search không tương thích với reasoning.effort="minimal"
        # trên một số model, nên dùng "low" để vừa nhanh vừa ổn định.
        reasoning_effort = "low"
        verbosity = "low"
        max_output_tokens = 1100
    elif fast_mode:
        reasoning_effort = "minimal"
        verbosity = "low"
        max_output_tokens = 700
    else:
        reasoning_effort = "low"
        verbosity = "low"
        max_output_tokens = 1100

    request: dict[str, Any] = {
        "model": model,
        "instructions": instructions,
        "input": api_input,
        "stream": True,
        "reasoning": {"effort": reasoning_effort},
        "text": {"verbosity": verbosity},
        "max_output_tokens": max_output_tokens,
    }

    if use_file_search and vector_store_id:
        request["tools"] = [
            {
                "type": "file_search",
                "vector_store_ids": [vector_store_id],
                "max_num_results": (
                    20 if deep_mode else (5 if fast_mode else 8)
                ),
            }
        ]
        diagnostics["native_file_search_enabled"] = True

    st.session_state["rag_diagnostics"] = diagnostics

    stream = client.responses.create(**request)

    for event in stream:
        event_type = getattr(event, "type", "")

        if event_type == "response.output_text.delta":
            yield event.delta

        elif event_type == "response.file_search_call.searching":
            diagnostics["native_file_search_status"] = "Đang tìm kiếm"

        elif event_type == "response.file_search_call.completed":
            diagnostics["native_file_search_status"] = "Hoàn tất"

        elif event_type == "response.file_search_call.failed":
            diagnostics["native_file_search_status"] = "Lỗi"

        st.session_state["rag_diagnostics"] = diagnostics


# =========================================================
# KIỂM TRA THƯ VIỆN XUẤT FILE
# =========================================================
if not DOCX_AVAILABLE or not XLSX_AVAILABLE:
    missing_packages = []
    if not DOCX_AVAILABLE:
        missing_packages.append("python-docx")
    if not XLSX_AVAILABLE:
        missing_packages.append("XlsxWriter")

    st.warning(
        "Ứng dụng vẫn hoạt động, nhưng chức năng xuất file đang chờ cài: "
        + ", ".join(missing_packages)
        + "."
    )


# =========================================================
# KHỞI TẠO SESSION
# =========================================================
if "database" not in st.session_state:
    st.session_state.database = load_database()

database = st.session_state.database
conversation = get_active_conversation(database)


# =========================================================
# SIDEBAR TRÁI
# =========================================================
with st.sidebar:
    st.markdown(
        """
        <div class="sidebar-brand">💧 Trợ lý CCTL_QNG</div>
        <div class="sidebar-author">
            Được xây dựng và phát triển bởi Hồ Hải Khôi Anh
        </div>
        """,
        unsafe_allow_html=True,
    )

    if st.button(
        "💧  Yêu cầu hỗ trợ tham mưu mới",
        use_container_width=True,
        type="primary",
    ):
        create_conversation(database)
        st.rerun()

    st.markdown("##### Trò chuyện gần đây")

    conversations = sorted(
        [
            item
            for item in database["conversations"].values()
            if item.get("messages")
        ],
        key=lambda item: item.get("updated_at", ""),
        reverse=True,
    )[:15]

    for item in conversations:
        selected = item["id"] == database["active_id"]
        col_open, col_delete = st.columns([5.8, 1])

        with col_open:
            if st.button(
                item["title"],
                key=f"open_{item['id']}",
                use_container_width=True,
                type="primary" if selected else "secondary",
            ):
                database["active_id"] = item["id"]
                save_database(database)
                st.rerun()

        with col_delete:
            if st.button(
                "🗑️",
                key=f"delete_{item['id']}",
                help="Xóa cuộc trò chuyện này",
                use_container_width=True,
            ):
                database["conversations"].pop(item["id"], None)

                if database.get("active_id") == item["id"]:
                    database["active_id"] = None

                save_database(database)
                st.rerun()

    st.divider()
    st.markdown("##### Chế độ trả lời")

    answer_mode = st.radio(
        "Chọn chế độ",
        options=["⚡ Nhanh", "🔎 Tra cứu nhanh", "📚 Chuyên sâu"],
        index=1,
        label_visibility="collapsed",
        help=(
            "Nhanh: trả lời hằng ngày bằng model gọn. "
            "Tra cứu nhanh: tìm trong kho nhưng không đọc toàn văn. "
            "Chuyên sâu: đọc kỹ hơn, có thể đọc toàn văn tài liệu."
        ),
    )

    fast_mode = answer_mode == "⚡ Nhanh"
    deep_mode = answer_mode == "📚 Chuyên sâu"
    use_file_search = answer_mode in {"🔎 Tra cứu nhanh", "📚 Chuyên sâu"}

    if fast_mode:
        st.caption(
            "Tốc độ cao nhất. Chỉ tự tra kho khi câu hỏi có dấu hiệu cần tài liệu."
        )
    elif deep_mode:
        st.caption(
            "Đọc kỹ, có thể đọc toàn văn 1–2 tài liệu; phù hợp câu hỏi khó hoặc cần độ chắc chắn cao."
        )
    else:
        st.caption(
            "Tra cứu kho nhanh, lấy các đoạn liên quan nhưng không đọc toàn văn."
        )

    st.divider()
    st.markdown("##### Tài liệu dùng chung")

    uploaded_files = st.file_uploader(
        "Tải tài liệu",
        type=["pdf", "docx", "doc", "txt", "md"],
        accept_multiple_files=True,
        label_visibility="collapsed",
    )

    if uploaded_files and st.button(
        "Đưa tài liệu vào kho",
        use_container_width=True,
    ):
        try:
            client = get_client()
            progress = st.progress(0)
            status_box = st.empty()

            for index, uploaded_file in enumerate(uploaded_files, start=1):
                status_box.write(f"Đang xử lý: {uploaded_file.name}")
                upload_document(client, database, uploaded_file)
                progress.progress(index / len(uploaded_files))

            status_box.success("Đã đưa tài liệu vào kho.")
            st.rerun()
        except Exception as error:
            st.error(f"Không thể tải tài liệu: {error}")

    if database["uploaded_files"]:
        with st.expander(
            f"{len(database['uploaded_files'])} tài liệu đã nạp",
            expanded=True,
        ):
            visible_files = list(reversed(database["uploaded_files"][-20:]))

            for index, file_info in enumerate(visible_files):
                file_id = str(file_info.get("openai_file_id", "")).strip()
                file_name = str(file_info.get("name", "Tài liệu không rõ tên"))

                col_name, col_delete = st.columns([8, 1])
                with col_name:
                    st.caption(f"• {file_name}")

                with col_delete:
                    if st.button(
                        "🗑️",
                        key=f"request_delete_document_{file_id}_{index}",
                        help=f"Yêu cầu xóa {file_name}",
                        use_container_width=True,
                    ):
                        st.session_state["pending_delete_document"] = {
                            "file_id": file_id,
                            "file_name": file_name,
                        }
                        st.rerun()

            pending_delete = st.session_state.get("pending_delete_document")
            if pending_delete:
                pending_file_id = str(
                    pending_delete.get("file_id", "")
                ).strip()
                pending_file_name = str(
                    pending_delete.get("file_name", "Tài liệu")
                )

                st.warning(
                    f"Xác nhận xóa **{pending_file_name}** khỏi kho tài liệu? "
                    "Thao tác này sẽ làm Agent không còn tra cứu file này."
                )
                confirm_col, cancel_col = st.columns(2)

                with confirm_col:
                    if st.button(
                        "Xóa tài liệu",
                        key="confirm_delete_document",
                        type="primary",
                        use_container_width=True,
                    ):
                        try:
                            if not pending_file_id:
                                raise RuntimeError(
                                    "Tài liệu không có OpenAI file ID để xóa."
                                )

                            client = get_client()
                            vector_store_id = ensure_vector_store(
                                client,
                                database,
                            )

                            # Gỡ khỏi Vector Store để Agent ngừng tra cứu.
                            client.vector_stores.files.delete(
                                vector_store_id=vector_store_id,
                                file_id=pending_file_id,
                            )

                            # Xóa file gốc khỏi OpenAI để tránh lưu thừa.
                            try:
                                client.files.delete(pending_file_id)
                            except Exception:
                                # File đã được gỡ khỏi kho là yêu cầu chính;
                                # lỗi xóa bản gốc không làm hỏng danh sách kho.
                                pass

                            database["uploaded_files"] = [
                                item
                                for item in database["uploaded_files"]
                                if str(item.get("openai_file_id", ""))
                                != pending_file_id
                            ]
                            save_database(database)
                            st.session_state.pop(
                                "pending_delete_document",
                                None,
                            )
                            st.success(
                                f"Đã xóa tài liệu: {pending_file_name}"
                            )
                            st.rerun()
                        except Exception as error:
                            st.error(
                                f"Không thể xóa {pending_file_name}: {error}"
                            )

                with cancel_col:
                    if st.button(
                        "Hủy",
                        key="cancel_delete_document",
                        use_container_width=True,
                    ):
                        st.session_state.pop(
                            "pending_delete_document",
                            None,
                        )
                        st.rerun()


    st.divider()
    st.markdown("##### Kho bảng dữ liệu")

    if not PANDAS_AVAILABLE:
        st.warning("Chưa cài pandas/openpyxl nên chưa dùng được kho bảng.")
    else:
        table_uploads = st.file_uploader(
            "Tải Excel hoặc CSV",
            type=["xlsx", "csv"],
            accept_multiple_files=True,
            key="structured_table_uploader",
            label_visibility="collapsed",
        )

        if table_uploads and st.button(
            "Đưa bảng dữ liệu vào hệ thống",
            key="save_structured_tables",
            use_container_width=True,
        ):
            try:
                for table_upload in table_uploads:
                    save_table_file(database, table_upload)
                st.success("Đã nạp bảng dữ liệu.")
                st.rerun()
            except Exception as error:
                st.error(f"Không thể nạp bảng: {error}")

        table_files = database.get("table_files", [])
        if table_files:
            with st.expander(
                f"🧮 {len(table_files)} bảng dữ liệu đã nạp",
                expanded=True,
            ):
                # Hiện nút xóa ngay cạnh từng bảng, không cần chọn bảng trước.
                for table_index, table_info in enumerate(list(table_files)):
                    table_id = str(table_info.get("id", ""))
                    table_name = str(table_info.get("name", "Bảng dữ liệu"))
                    table_name_col, table_delete_col = st.columns([8, 1])
                    with table_name_col:
                        st.caption(f"• {table_name}")
                    with table_delete_col:
                        if st.button(
                            "🗑️",
                            key=f"delete_table_direct_{table_id}_{table_index}",
                            help=f"Xóa bảng {table_name}",
                            use_container_width=True,
                        ):
                            delete_table_file(database, table_id)
                            st.rerun()

                selected_id = st.selectbox(
                    "Chọn bảng",
                    options=[str(item.get("id", "")) for item in table_files],
                    format_func=lambda value: next(
                        (
                            str(item.get("name", "Bảng dữ liệu"))
                            for item in table_files
                            if str(item.get("id", "")) == value
                        ),
                        "Bảng dữ liệu",
                    ),
                    key="selected_structured_table",
                )
                selected_info = next(
                    (
                        item for item in table_files
                        if str(item.get("id", "")) == selected_id
                    ),
                    None,
                )

                if selected_info:
                    selected_path = resolve_table_path(selected_info)
                    if selected_path.exists():
                        sheets = read_table_file(selected_path)
                        selected_sheet = st.selectbox(
                            "Sheet",
                            options=list(sheets.keys()),
                            key=f"selected_table_sheet_{selected_id}",
                        )
                        dataframe = sheets[selected_sheet]
                        st.caption(
                            f"Đang đọc đúng file: {selected_info.get('name', '')} | "
                            f"Sheet: {selected_sheet} | {len(dataframe)} dòng | "
                            f"{len(dataframe.columns)} cột"
                        )

                        filter_column = st.selectbox(
                            "Lọc theo cột",
                            options=["— Không lọc —"] + list(dataframe.columns),
                            key=f"table_filter_column_{selected_id}_{selected_sheet}",
                        )
                        filter_text = st.text_input(
                            "Từ khóa lọc",
                            key=f"table_filter_text_{selected_id}_{selected_sheet}",
                            placeholder="Nhập họ tên, chức vụ, địa phương...",
                        )

                        filtered = dataframe
                        if (
                            filter_column != "— Không lọc —"
                            and filter_text.strip()
                        ):
                            filtered = dataframe[
                                dataframe[filter_column]
                                .fillna("")
                                .astype(str)
                                .str.contains(
                                    filter_text.strip(),
                                    case=False,
                                    regex=False,
                                )
                            ]

                        st.caption(
                            f"Hiển thị {len(filtered)} / {len(dataframe)} dòng"
                        )
                        st.dataframe(
                            filtered.head(200),
                            use_container_width=True,
                            hide_index=True,
                        )
                    else:
                        st.warning("File bảng đã mất sau khi máy chủ khởi động lại.")

                    if st.button(
                        "🗑️ Xóa bảng dữ liệu",
                        key=f"delete_table_{selected_id}",
                        use_container_width=True,
                    ):
                        delete_table_file(database, selected_id)
                        st.rerun()

    if st.button(
        "🔎 Kiểm tra kho tài liệu",
        use_container_width=True,
        help="Kiểm tra trực tiếp trạng thái kho tài liệu trên OpenAI.",
    ):
        try:
            client = get_client()
            vector_store_id = ensure_vector_store(client, database)
            vector_store = client.vector_stores.retrieve(
                vector_store_id=vector_store_id
            )
            counts = vector_store.file_counts
            st.success(
                "Kho đang hoạt động — "
                f"Hoàn tất: {counts.completed}; "
                f"Đang xử lý: {counts.in_progress}; "
                f"Lỗi: {counts.failed}; "
                f"Tổng: {counts.total}."
            )
            st.caption(f"Vector Store ID: {vector_store_id}")
        except Exception as error:
            st.error(f"Không kiểm tra được kho tài liệu: {error}")

    if use_file_search and st.session_state.get("rag_diagnostics"):
        diagnostic = st.session_state["rag_diagnostics"]
        with st.expander("🧪 Chẩn đoán tra cứu tài liệu", expanded=False):
            st.caption(
                f"Vector Store: {diagnostic.get('vector_store_id') or 'Chưa xác định'}"
            )
            st.caption(
                "File tìm thấy: "
                f"{len(diagnostic.get('manual_search_files', []))}; "
                "Ký tự từ tìm kiếm: "
                f"{diagnostic.get('retrieved_context_chars', 0):,}; "
                "Ký tự toàn file: "
                f"{diagnostic.get('full_document_chars', 0):,}."
            )

            candidate_files = diagnostic.get("candidate_files", [])
            if candidate_files:
                st.markdown("**File ứng viên:**")
                for item in candidate_files[:8]:
                    st.caption(
                        f"• {item.get('filename', 'Không rõ tên')} "
                        f"— điểm {item.get('score', 0)}"
                    )
            else:
                st.warning("Tìm kiếm thủ công chưa trả về file ứng viên.")

            full_errors = diagnostic.get("full_document_errors", [])
            if full_errors:
                st.markdown("**Lỗi khi đọc toàn file:**")
                for error_text in full_errors[:5]:
                    st.error(error_text)

            native_status = diagnostic.get(
                "native_file_search_status",
                "Đã bật, chờ câu hỏi",
            )
            st.caption(f"File Search trực tiếp: {native_status}")

            if (
                diagnostic.get("retrieved_context_chars", 0) > 0
                and diagnostic.get("full_document_chars", 0) == 0
                and not full_errors
            ):
                st.info(
                    "Đã tìm được đoạn liên quan. Câu hỏi trước có thể chưa kích hoạt "
                    "đọc toàn văn; bản cập nhật mới sẽ luôn đọc ít nhất file phù hợp nhất."
                )

    st.divider()
    if deep_mode:
        current_model = DEEP_MODEL
    elif fast_mode:
        current_model = FAST_MODEL
    else:
        current_model = SEARCH_MODEL
    st.caption(f"Model đang dùng: {current_model}")


# =========================================================
# KHU VỰC TRÒ CHUYỆN
# =========================================================
if not conversation["messages"]:
    st.markdown(
        """
        <section class="welcome-shell">
            <div id="welcome-top-anchor" style="height:1px; scroll-margin-top:132px;"></div>
            <img class="welcome-logo" src="data:image/png;base64,iVBORw0KGgoAAAANSUhEUgAAAIwAAACQCAIAAABVthSJAAABCGlDQ1BJQ0MgUHJvZmlsZQAAeJxjYGA8wQAELAYMDLl5JUVB7k4KEZFRCuwPGBiBEAwSk4sLGHADoKpv1yBqL+viUYcLcKakFicD6Q9ArFIEtBxopAiQLZIOYWuA2EkQtg2IXV5SUAJkB4DYRSFBzkB2CpCtkY7ETkJiJxcUgdT3ANk2uTmlyQh3M/Ck5oUGA2kOIJZhKGYIYnBncAL5H6IkfxEDg8VXBgbmCQixpJkMDNtbGRgkbiHEVBYwMPC3MDBsO48QQ4RJQWJRIliIBYiZ0tIYGD4tZ2DgjWRgEL7AwMAVDQsIHG5TALvNnSEfCNMZchhSgSKeDHkMyQx6QJYRgwGDIYMZAKbWPz9HbOBQAACIJUlEQVR4nLz9Z7AlV5ImiH3u55yIuOrd+7RMLZFAJpBAQosCSqBQqruqxUzLmenZmRWzbbZrS9LWjLbG/UMauWY02s6SQ+PszkyP7Onuqu6e0gKFQhVUAUgACaTW6mW+fFpeFRHnuPNH3JdIZKG6qwUZlgZ7uCJuxPE4Lj//nEII+OhBRLf/VtW/9PW73ire/dlXACgYAEHu+krxyl1nVnBgFoCLzygAsIIAIhSfUoDgqXiPyMMoeufl3mvFOT9ygwoBIAAR9z4G7V2DClB8WqCMj9wEgVh7l/NzV+Zn7/pjP3bX14v/LT5w1+qpqv3Ys//Nj7t+6fZ1YFNUUCGiO8WDu+9KLEShAANMSgBUoAAIqlCCAsoWECI1UCseAMBKrCABAkCAEhggwKiS5FAFxIDVRArqSagnnuIPgAoJkZINRMXjwArqifDDFb9LEn/xmnzsEt35eu9p/ujq8S/4A38rx198P3dIEVAmCEEZwhBWISiKbQRAobq5mXp7FL17ES3+qRYfxe2fZCg09CRRfAwQgoIUDCrkyCAGWYBvS0gBTwgoflT/Yi3yt7U4t6UF4BfaSXcKtrhKIrprk9556o+9biLiD0/yMT+h4OKLokIgwxbBwzBUETzAxAxiBIgImBikEAoCCDODrYglIjBUISIqgen2pQZV7okHRGRBnAupgRAIpGIAYwlEvetTEFFP4ZmPu+aft763b/9j9VuxbsVV/ey3PvbFX0hId53xb/7sEIHu0PciIsV1MwMgMlqskYl6G8FAJA/BF7fKhlVVghARkWEYDZR7VcMKUCjMBzFYVVUl5B4A9QRmiaiwWUGEwYUqZIYCXout17tD0zOBSkqq8rd3+3+1M/z/yib9ogcRADaOVENQ70WosOXwABOgUAYIwi5QIFYDyiRVVQYbMgauOA0pgoUAKuDeA2ugENHIRQEQKVSVKIFIjUjJARogKmR1U1oCqPZsGFRYfaEhiYziQ8fh/5/HX1NIH7uj8VEH4Rc7Q0/zihIIWhjSYtEJLCCGFvaAoGQywAPCZQVgEBTBQxWGwQ4eCOg5ZQQYArhnbQzguHAGGIXBYQY8NICY4FWNegogURhDhW8iSgCxKujnyuZOPfaLr9tf6Qy/qLr7S//+i1/cvBrcfnPzygiA90LMyqQKLwgKBawiRl4saa4kjnMgBVaBDWAjxVITK+vYaKGTFf408rQnVOHeK2zgCEN11BMMlTAQo2ZQNagQLCHAGmYVCSIheCKyxhlL+aYfHogCHAhEMAQImD7GvP8FLvjfilvxV9hJd13QXd4E/uqq9rYDY50RQBReIIX3a6BAUxwzPJACG8DMOs7P4Mrc6sxSa6mdL6y211rdLJegBsQErcbWQAMjV2SiQQRGGUoh70vsUDUa7StN9JUnBhtTQ+XBGiYHUGZOmA3DWIiH5sJZHltDrAQjRFJEVQQF3N+GH/eLK5sPv/KLBLPMfNv1vO3X/TWEJHc+WCofnpOtMhQICiEEIPdYVdwQ3FjHzBxmlzoLK9351e7SerfZyQBWAlOIrJZiVEu2krjYap81BBXlTLTj0fXa9ZIFamc+DZJlPs0yEThjkyQpRbxteHCkFm0fLe0cxZY6hhL0EUoKk+eGckssxKIUyAqRITj6ME76RW78Y5foLiH9IkHxX1lIzIzbjjizKgh6x7UKkVFVLULBzUhewQoKCrN5n6IUoCIQ1aCMCAKkwLpibhk3b7Yur7aOLfiba53u6oqFGlu2NqmU+2q1Wl856atioIG+EiqMyKJURjlCzQIKUXgg9+jkaKXoeGxkaOfY6GC1iZVmutpst9rdPM/bzRZLiIwMlOPtw409W4Z2TZjJOsbKKCtKgIGyBABCFoyINzMfmw9dEX0RFbeqpEAvPOe7fO7bf/91hCQid710lx37edtTQTkBgFOwqlKhFQTgoMYr2MACFLrigxqnFGUKR3AkIMnJpoV+YzChBdxMcWYe71zF+en28uJcM+10YIYqdrtd29Ywu7dNjg721ypmsI7IwOeoVBEDUYrEIWdo4R0AoXd5EKCbwUYIQEfBiqBYbWGjg9xjaV2nF9evLKxNL7Wn55ueauLiUn91cjB+cKL2+G5MlTDiMOQgqRdYE4F6+6h4tDQiR4LgYSyU4AGGN1AQqZISF/Hyxyz6HaroYwPQO2Whqn9LQkKACkihBGIvLEQADGBYAAHYBxYD8YgkhUFq4i4gQAbcauPN0/nRiyuXlrLptSz1eX/NNsp225axR/fHh0cwWULFwAEMhAyJRZqmJo4dIc5zQ5rDwZIjCJAKZNMXkxwAWqmPY5sYkMADZKFAG9gAWopbqzhzDe+da56bW7vcWifWStfv6q8c3D7w5D39D2xFv4UhlBgWCMELgjGGoaTEaglQ9PxPRsYkBCPqek8K7l7evwUh/SK77045AQAF1gDJoAx2gA0qbIiAPA+eHDMoBEXIArnEOfWZqDcuBa6t4thV/fax2XO3ms1uVo5szXT2jpc+cf/W+3aVxvpQAgigHGVCbOGBoJpInqdZzmUwJQ6kEjKNrDVG8sAdixQIQA5EACliggUgoAAiqEGuACMrEkEKS2gDZxbxo4urb52bXZirLq9JO2/VS/nj94x+6bnRw6OoA3UghnhREW+KEJytsXEhC1aAM1AqMNAyCZh66vFnUxJ37YSft+w9O3KnkP5KErrjEEJAL7kZKdhLcM5AJQh5YwAYL0TaDciJ1VEGzDTx7tnm0RMXz8+0b7QTW+7bMdZ/eP/gkZ3YO45BIAISIEsFji2DAA80AVHEiphgFFkGYyUyjNRbNpnIGycunF7IN6gebIVt0qiWYtK+Eo/30dYG+ktghXFoAwrEQKSIpRsyH0x1MUezgnngtbdx+gqml5dnV+Y73dXJwdKRXZNP3zP0wARGEjhAM19ywfTCAwNiUrAKUaqUQ51KCfh4IalqYdc/8rj/NYT0i0hICVLk7ovLQFBAiwQCGKoiEAIMqQpnmQTOTJTHuN7FG+fwysnZSzeX8zyPSB/Yu+XBPYP3b8dIDRWGB7pBY6CfKO9040rSBq60cXMNN5bgSkhiTDVwTxUlhWSIDJyqZ5rz+Bd//tZ3359uurG26c80KtkosppQ3q9rv/bYll97bjRv49z00ow29mwzD5SQ5OhqOxBD4m4gH6Odo5JgbgOXb3WOXrr57mw+syKhSRP1xlMPTDywA/dvwRChBDXSYTEgExhCwshZhQJDnXIkBO6F5n95ZefnvYWfV6r4K4U7QmCFgjysEghiVFDoaWYJIc1yYyMbJ53C9ryH189MX5jbmN/I4tg9dGjPU4fKe4ewpYIBwAAAOgBIY2YIokqyBrx+sf3G+fXppc7MzEyqWh4Y3jta//WHxh6ahGUYqDHkCaUY+/bvD4N7W+X+Vy/i9OVZD7NvYrhO3ebVxWazKRi9toav/eTkySX/mSce2fVozahqqZwDqx30l2EVCYO6OmX92PZ43/bdh1p48zROHru5sZ7++Uunjk30Pf3glifuw+4y1bkcKyTraFIkZYu0PEOUKGixMD//uG2Z/uJlJ6K/PJj9eWEzKQxUQZ7hixoaEAdmFdGuANYllowXow63gJMr+LO38cGllRvX5xqJPLRv/On7po5spck6LBABDjCbblgfsXjJvO8k0Q8u4Gs/mbu2vDw63Ni7cywLfPxm963F6f19emBsfIjF5m2KqixoGDx/b/0ZgznGzBJu2faB8frf+wy2xwmtPT5ZhwEuruEmjS77tWaO1KEdKANeOb1xa8Pt3ZUcGEAVcIJK5HIDIjwc4d4nMbN38oNT3TdPTp+Znj2x0n57cfeX7jVPb8OkAdSrOiI2MFSUQkiFpDDXdy33XX//goHt3yjBSipEKPKZfDspCWIX5VnKIE/IYiwDP76EP3x5+sSsdLrZ7qHqZx/e9Zkjbm8F/YAJHTElUTBhJcWJGx0TlXYNwSmXkmixhRePrX9ws3P//m2//pmhhwewmuKH7+Po0WNZZ1VknBnqBU4tSD36RCrEqSLrtDVP66a1ozZ4bw3lPljGguL6Oq6saBCzd09NDVY9Ftv451/9yXTTPn547+9/bsdAAwkobWuIiBPcWMJyR7aM8+efTQ7eu+e7b7V/eGHh9Q8uzV+L549s+6VDGC/XFHBQBhEpURZIBUXJ5O4a9F/v+IiQPrb+gZ+7mQQsCMHCWmJQgGgOzpU5N9ZYBVrA5Rxff8f/8U/OdbhWL/Nn7h378sHa47sRARGg6QbBawg2qnrg3avr/8O/fb0xvOX3Pn3fU/ugHvMrePfC5Xr/wPMHh54bQB3YFmPHEbywZX9fvVSJkHWYUSoxeZ+SWqsSKUOx3Eo1qVTKrhJBM8RGm51uGpeuzqWBSjvHS1uGkAF5jJPXsKDDITawUZJQlrWs2CiK1wVt4Ifv3/rJ8bPj4zueOrj9ub34x58u7xjd9uIHV69v2H/76o1zN2tffqJ+/wgGQLkizYittaaX5XLCRUELH+dh/yK67mOE9POOj1d3AEIAE0SRZ0XK2pOB4zygqWjnuLiMPz86++rZWy7ur1r59IPbn9mDp0ZQAtoBakBxTQAjwkAz1XK9L1SGzy5mb1/d2L61tr0PGx3pilRIxgutmCOxUjJcmShtFOUMRkKu026DJI5iw8YHaXU4zb0gVCKuxUgAD8rZza9idqXVTDXt6qXraPYjN/jxGSykth7Rge1TgyUg5U4S31rEzTXYARy4b3yu7a9emz/65sIzow/vGMKvPYzR+vbvnUqPX7n1zqW5ucWlX35457MH0UcoxxAxa+1OtVzKfDDmI27czy7j31Td/ULf1yLz79WAbOIDhxxC6BqsAa9cwJ+/cvn6UprElcOj+NIT2w+MY9CCPZppMBXTAs7OI23h4W3cBzjWPWP03JGdf/rKB8dvLhxYqQ32IReFzxlGE3ggdSBwO2A1x6VliMf+ETgHBxfFrptDDTLhjQ6875ZtGCq7GNhQJASU7cISNroerjTXav+b71wKzka1wZvLaXDl7UPh0V2oeqQaXyf8+/euXJ2R8cHGb3568O8/ueVEIx7vrw7XIJk0LH9uLw6Mxi+fGHz14ur5G2v/7rWZM7Njv/4E767BeCRRkgsSZyAfSXj+pcL4KwvpLz0UrByTeiImJhHOgBChzVgAvvuufPfti4trnYFS9NlH9/zqY9QPRHkuGxsZubhcI+DUJfzTb77rItv/lfsPjoMNW+CTjwy8caJ6bbX95nl9ZBs1BkytHDebzaNnsetJMNAF5gO+89qt19+/3Ejkf/+bTw71M4CVpQ3XVxOBWOQBkndLJow2ygawBAE6wPnp5srSwvbxXYe3jvlMz80tX5hb2pB64uJd43a8D1mGvMwvncG3P5jrrMjzhyuDDmNlDD86EhRRDAs2KSLBwToGHq71DdRyL5fm2j9+/7K26195anjXMFxGMWOj2a4nCf4y0dwZzP48Qf4164yq6kFdQkstKIYgT7ts0GKcbuMP3sQf/OjyzYXW9gb9k09O/VePUb9HGSg748qOnFVoSOHzzsXVsGBG3r22sQ5khG6OLQ3s3Tq1kSanriwvrWJ4AAf2bQuM778//72LuCq4HPDSTbx4w1/NKty/xVn23ZQZ1b4qGyhgGSENWWuVQ16v9TkgBhjoAAvLK1FoHZ5Kfu8Z/Pcv0H//lcHHtpqqk8jYw7u31hw6Jbw/hz9/ZWZ1Y2BXfejzh8dGHXJghfAff3z568fmF4EogvPeZpiK8UsH8N/82rbDU07z9M9+fOaffXP6g3nkEZjRSMq4w+r8TY6/aCd9bGD10U9AAQmBiU0SrQMnlvCN97pf/+llAT26ffjvPjn16R2oKTILAdbAsDW2UKBisGOyNDY5cXWlc3Q6HD5U21mGNagBh3ePv3zx+vxK6+yl/k89xF94tF+y9Pjl+W99rz09WVaWs4vZ9Ozi9uH+px7ePlhHEuJuNxXKjI2JgQDya+OJ5xhDFWsB7QZOzPo6Zm/eSHxzZz/GLUgwVkHNBd9Z2zEycP8+WGDO4ztv3Tx3fa2/2v/C4+MPbIMBPPDORbxy4saByer+iZHhQVRjC4akUo75QAO//5vb/v23Fo5dSt48O7O0Mv97v/zQU9vQFxAxO7q7yICP7puflxG/8/gbVewLTaIhFeIW4eQCvnm0/dK7F0sm/+SBod97buoTu2AVHQ8FmsCNHOeauLqKlRZCqvUIj903FaT7zlz63fOY9zCMRLFnDCM1zdrrl27MRsAnxvE79w98eZ9t+OnjFy4cPXU1XVl8ZFvjtx6bfG4PrKCdISdnktgLrIHV5pYqfuWxXV96cNuuAcQeVWucoL0OTTe2DZit/QiCdo7Lc5hbRUT5/XvqgyV44PhVfHBpo1rCY/cNPf8I6hECMJ3hR8dWWjI5MrBlsAGyyBxSQpdZgQiYivBffGX4mQenarXapdX+f/P9pW8fC+sOnu+26389y/SRndTzESFQVkKReCrAa0UZRW9DGwFWBIFhCJuM+WIbPzix+ubJGxSyx3YN/vZzw/vrqAFk0AaawCun8da5+Zlb8+NlfuHwjvquUrmMJ/fi+z8Nt1K8cXH1iZ2NuIKSQbmKoaHBpbX19bUlh4kB4FN7ol2j+6+t7T87n4mNGhVsaWDvMMoBGiAEl3DXwwMWIFuanKh+eRIZMAJwCh88RXbHKH7lU4/YKNk9BZOhkmB+Hks3rw/b0v1bkQAEfHCpvdTJpwaizz3qxkqQHBsOL3/QuTC90l+p3rdzeDCGAZoBZ26FG7cW924b3TaKMiCKX3s2LscHvvryjRMXbyW6XK/ueW4fGC5WFCDC4rHelFJvbYkI8pfoxLvjJIJAAlSJrRAEIAUrQQMYEA4GoOJFGIONdloux3OKP3h16aUPrjsvz947+XufGttVgwq8wgLB4N+/hW8fveRDWrJ04+psJeGDW/YPV7C/H198cPcfvHXrwo32n78Z/d1PlWvApVUsrncjxuHdUyWAPChgXwPjZTyyLcrRA7VGgLUQi9wjD1CLNMADlgwIFogBeCQRQFaBUeDT9zZSQANKgixHw7R31FoDfe6BQTAwn+NWs8s2fOLesSMjsBlshHfm8f0PVvKseejg6H1bUQdMDlh8/YP5t06cf+GB9m9+dkeNURJsUfz246iXpv7tN+ZOXsnph9dHqluPbIUDslxSRSliARx6j7souICWaSAUuwJ3ieNjhNQrfijj9ncIvcqjMIInpgLQYxSqaKWw5Xge+FffvfzK6WZA9Pi+/n/0pbGtDOdFwWwhwCunOl976XQaDzz7wL5tgyZv7tg7XM4dugE1i8/cXzp6MXr/ysy7HyyGbFe5XD57+fry3MqOofjIff1GkQUkEQRIIuRA2yM3aHXRbKOdodNFsxWazVarK13viNViIymRLcUlYycrtf4yogpKCRIHKJhQMmCCenziYHnvthdshuEaQopWCoVXTUdH6gboEDqCn3yA+XXZM1Z54aHKcAyfwRJuruBSxy64gUuLK6udHXEJVYt8rdufJF95CMY89C/+9P1jF6b/7Xfbg39v/84SjOESwysiDcpgMgRQsdgEkFENSmCYjzVLdwpJQEEURFHxQSEwcoAACyEwk+k9BaaAWUW4muKrb4Wv/nTelRuP7h7+R88P7jCIAQFbggdutPHi6+faK0tf+uJDf/8TGAdEy8VbG4rQxq5+/N4T4z8utU5cW3z1zfepOlApx4/tn3ru4NBUHxgIMZpAS7HYxWwTVxZwYQYLa/nNxfnF9Y1WO8tTVh8LtFQGc5dCMym5rtiQy1C5FLGp9fUND9SnRsqTg5hsYKIPIwkqERJgTx8cQMB6hqEK7t06dPL8+VdPzUbVsXtGcOYi3j4zp1nruf0jh0cRA2rRZNxKMdPJWxHGd41XKsgDWNFfTdbWUbH49CGstQ/86J1Lb063/umfLf4f/87QmAVlqEQg0uA9bJG4C4B6cgFEJIQClcQ/K6cPhaSqRAKwFJDEXpq9AL5bGIAZCII8KJTiLjCX43vH8j9741IoDT8wVv69Tw3ursFm3jjLjJABBksdXJtbHx0Z2TaAYSDpInFY9nhvBe+fmy6H1gtP73/6gNk5uu+N81tutmNfM5UKSinGBtABGGgCV+dweaF97vrayctLSxt+teU5ilwcR0l5pDIUJbUkiksxYodyCciaZFwqUbuVWg2dZmtpZX1mdum9k+2yCwOJjjUqu3dO7Zka3D2BGmMoRkyoJDCEh3fyxR2Dp86f/dr88lhf7dr1GeLK5x7c9bmHSnWAgYyRAZeWsNjKkvroxJ6RecB53NrAZAVJCd1mNlaPPnckStOpb7118/3Lc//+e/jdzw6NRgCgIsYagAGBeoCUbhv+n3vY2xICoDCKXv7WaLEhN0NlUzQWgLMU1qWEBcUrV/Dn71yf7+jBLZV//LmxA2UMADAWADFyiw4QGMEyM6vCAmSQC5Zz/OBk94fvzI6XOoef2G8J2wYxcKQMh5kuvvVy881TZ5fu2TX1hf4zN/D2mfa75+amVzo5ENt8fKD04N54tBpNDDSG+iqlEkwME8E5JIpqDN+tMoEIeZ5A0Ukra2sjN5azK0tLMyvrs0sb79/YeP3WFVe6tX1s+OF9ow9NYd8AJh3KwMEG+j5zz3fe5LMzzSvTa41y/PDW5CsPlSYTUJ7Dmky5yXj3zHSa8ejUyOlZXLiOi5dXE0rvHa994ZHynnoUsu7uKPn7T9dEdn7njRM/urhYnRj69YcgghIMh9wYR8pQU4CtCSC1XGiuO3I9t1Omd6o7VpViGzFACIVtK+oiROSVWOEIwqUV4Owa/vSNmRvLnYFE//4Xxu4ZxQgV2EbAGQDOIgX6y9g+2f/Ts4szy7i8ju198B7rBufm/UJaHq9H5QQlAF7KzAGoWawtzuWmzwz2/+lLrdPHP7i5rBwPTA4O7toxemAXphrYNogKIQGsgAFlBIYCDljdQJUxUEUCaAGABbIRdBCtyPhiOr64gSuLOD7TujK7NH1r4ea1669Q9+CWwU/cv/P+PUkjxu4h/Gdf3Dezgfk2SozdFUyWYRRwrhUCDI7P4sTV5RwDM9fWfnC521lbFqMmX7l1Od059OiefZUYeZQlA8CXH4tm5offm8u/9sbFscbuT+9CDAMvBdS9sP0kYMAUPvPP2VB3OQ4MBUOIPKmAjKoJBApKpnBwrWOTA2ea+PrR9Pj0Ui2S//ZLB58bBysyQDutUikBiZegqiVb2hLjuSM737nW/N5bZxvl/Y8fhCGcnsXMSqvkzD0Tg2VBRBAlsWgBbQszOLGy0n3xaNMsnRlJwpOHdzx6aGr/OEYrqAFZFz4DmQIjrgnDEKlHU3BL8d03Lo8N1p9+YHCUEQUoQAwOSAgNRdVi/yieHUfzYGU5rXxwKRy/OHNheuno9Y2fXDp+776tnz0ydngbRi0O1oAauhnKEdIM1iITeGNmAr71Vjrb6YMz3eWFikXD5MN1O1rpPzBgdyYsmbdg51AO2B7jn/za1v/u392cWU2/9c7N0YHJ+/sxaJ2EQMpCluhDdDkR/azPUGwmi48Gugyw9DBzAIT0tsZTgScISrcyvH0F33zlWKVaef6RfY/uYu5qnFCqqJQSQCDeGiOiCnHgw7tKzzxy4OVjl//ou69fmN7rKJy9Op1mpQOTA88/ONYgZDkQ0Y0uTl/HOyeunLm2VHJj9Ur56cMPP3oAEwOoAQNA1M1dyMtAyqXAFCwgyj5lBYslcusBLx+f2TGZ779ncCBGoipKPoclkKJMICtx8BR8H8cjsdl+wDy6d8uZ+S3HruHNs7dOXZ2dvXHt7D1Tz90/uW8UDUY9QieDtRCBWrSAHx7N3vzgKlNlex37t+3Y0jD7t2KkhuEEA4qaRcnBwIjCAHWLYeB3X5j8Dy8uHr909cVjje3PVRKgrMQaAlsAEYQgKFyyjztU1d6JawQCKZNaaATKAAF7wDKRChyjS1gO/Mql9I++fyFOKo9sK3/pkSSxABMBuVCbucJkEBByNiUPtoKxGL/0WE396IkLcvT4aaStwXr1iV0Tzz04cf84Io+lgHM38O23Vt45ccHm69vGa585aB7aN3DfGKKecwkjIMNC7EESUy4QD0eccAkCL0EMTl3G5UXlKlY9tATN1yw3wuYKEEnRJhFxZLSo+WDKYngChyfwxK7xYxcH3rt464fH5968qUf2TX1iDw6Oo+YQUkkcdwQvn8VL711fXV198ODwb396YO8wRhKQR0JAChY4AwFCgZa08IKS4pMTWN7d+PZ647V3r2yt7P7VR5OY2GomZEWUNXWWAXhRs6nY7kpM2DtfYtVew6MCxikJIwCSZ8ZFCIqUcKmDb787vd7RkWr8+SO7tieoWwhwaQmNfjTYqmo17xrnMo8shFJsFDhQR/UzY+9tH1tZz0vIEkP795YHyhBgvotvv7P2kxPTV+fTWjl+4oGDLzw2uncMVWgJOcMoDAFOAIUQQJR7KKFsYYFuMyuXo0CmQ3jlvZWOGby+jvNzOFyHWGsYESAEQ4HEEylIAVMEhDEjBjRHFTg8gt0j8f7d29+9tP0nxy7/6M3jVy6Vnr5/52P7zI4KZx4rq3j/+PnpK1d29Dd+/RMDh8cxaOA3fInZaWC2sFRADAsPK6TqIkoUCfDsPnt1fvjNM1e/e/TSnp33Pj4EI6QKZ8goiwRiw0WH1MfapA+Fprp5fgWTJ1IyBsIqUWR8AFksC759Em9eT6O0/YVPHXx2D2wXseL6Kv7lDy7Uh0a+/Gj9ngapqSH3MLYUAxnIy0CZKxG23otO7px31kIc1oEbKf7D95Z+fPKmB+0YrvzqZ/Y/dwA1oAZY9eKZYHqNQwIjAtKgPrEWTHERtkc+56jFOHoJ756/kbvh2U7n7UvZZ/ZGxlZrQGxACtIc5BWBwFACLBQhIM8DDMeGoCgRnhrB4QF8ZtvOP/nx6TdudY/94PSbM7t/91Olgw6NPuwbSzo73IP3TDw8hj5GyJAkNmZoJmx9cE4Uoi0oMcqRodDJKrFDoHuG8Nkj9YvX7Qez3a++1d72qfJWF4uqAxFx8MEy68/3wz90walnkgASsPEEAFYZKiB4oAOcmsHrp1dbIX58X/3zD4O7KAWJLc/NLb3x/mmtLJej+5NHk71lqFgJcAyyiEVImATGosQoJVjPMdvCpVX8x+/dPHbxVqnaeGDX6H/2hdr+GoygwYgQQp4bOFDRPwlG8ZiCyFhiJqhvkiUbxcvAzYBv/PTKmiZdU8qUjl+dO7eypdKPmgICIs+GRIjuTCgTDMOrGoAR4FMEqdjKoKHRCUz95oGRo3j51K3X3jl55Xz0X79w/3078MKzWx97eOtkCSUqcICBybQ7WTVWH1IvPoojA4AoDcgCqmUjoavCbOIHt+C5h++5+ONLP/ng6rM7D9T2oMqUpsE6smQAVkFRx/3ZvPhHs+Bqit6tXmM3oMQg5D4NFjdz/Pj48uJ6d6xe/q3PTE05VAiJZcMYHxt86rFDJom++uqlf/6djbOLCA5gtFsZSDi2omCLTgphLKZYd/jpDP5f35h/+YMrI0O1rzwx+d/9cu2BMvqzMMhAAIOciyyLZTEMQ1BWNeQpChx3QoEWjgCbwXWAd6/jnWvtrDTUgg3OzS43Xz22CsARKG8zSBEFruRcDuQUpojvVRE5E1vhkEbSiWMweUk9PEYI//AB/P6TA8/uG15cbf7fv3XqX74up1ZRqyBrwbfgMy3FhiQrx16kaY2WjOHcU/AEMQY2hkAzaGZjUVSBLzzKe6eGOl6/d+z6+Q14RlBDQlCj/i+qR3xoqVSZWIs+uY90DhDI2BQ4Pou3z9/iED//6O7DY4hTWItuAHtMDuG3P7djaALfeuXym2eum27yq8/semArytUoCEiCwmRd1BKsePgE33oXf/Taratzy7u3DP7G8/s+dS/qXdTVk3TFl53lLPMRK5iAQNTb38U+EsA5ZKEIxkyuCISjF3UF9TZImEAupdqrp+b+zicbY4CxjCJHTCAYq54gShCCz33RyeysCYE4zchExjqfaaw0FaFvb1wZ2m7j6htn5l566xT7qcYz/ftqsCnYUJp2Y3hmCDnrSqQuV3iOM2CmhRsL7Y3VW0fu31UnOMAKxgy+8kT/4vzc0UsL91zburUf4zHgBVARsOlZmzuD2Z6Qboe1RCSkRaucbMZVikzBxO5WEz883lnIcHBL8vkHkHjEBiuEFGgYmAx7Y8QH4Xjnt988+50LKwtm8SvZ0NO70c9IyFiSRELaYlcxf/g2/t1Pzt5o48CuoX/w8OCn9qOUor+MVtdmWVQynAAuKrIoWRGTqwLKrCFQ78osARIKDX1lCaeuLXVMyWepjdgDKfedW+NvvYOJ+7A1thBDhftRaGESMVZgXGQlBK/CZCQeyEKwwVhFHFOnuUHVKhm6fwBjvzI0/B37o7dPv/zmeZVDv/l4aWcNDQvptKMoyTKhuG8jBxQSYQM4v4ZX3g/vnr4eQivZhif6EXswo9vGF3fj3M76n5zUH55avX+4MTYFUA4V4UjpzkLGRw6+U2K9opFAey28EmAyxCuKS8v46cnLLko++dDUeAzO0QXePJf/wXfOff/Y4oogMpis4vOP4je+uH9krH7i2vV/+c13vvVuWAZaQKcrFLm0ZL5/zH/1xWMbGe8YrX7pmdFnDtmKoBahtZ46i3pfxIrgUUCiitJvQdmgxAImUiLN8mAZIMoDcuDVd27dmF/vBuoFEkGzYFpU//pPpq8swostJG0BQ0VK3xTQ6lwAYww5wKaqygYWXtS3Nkp9NW+oYE4ZAX7nM41fe/ZwOYp+8vaprx9NbypWPFxcywNFSbXD6DisRzjXxp8czf6XPzn1jTdOzbR4bi27cC14wBikrbTh0Bfw/MPjQ/3VU5dvnrjWXQU8MViIw19QUrrTcZCiEULZMSgoRFhN1AbOt/CHL5/tenlg78S9E6gKohLOL+APX7r81pn5amn9Ew/bf/grjT6HYeCFXdj1W3v+3X96/+IM/fOvv2Po0ecPo8F2PeCHM/h//+jUapv3TVZ/+ZMTh3ZgrdvDFlcqsclgAyLb6yQBQdUWXCWqJCAYayGOREgYBGVx8flbOHNlvtlOTNmS4ywQ1JBxkHR5de3kxfiegZHBGIa06HwVk6gqg6TIJwEOUKgNcBZE6PhOteQCYQ1YCzAelRjDFl95qrTe3fbisSt//PY5M3rot+5FX+QkmABsKK608doFvH58fvrWokgYriUT/faeyT33jhkGAknsYvFgYNcYHt3deOmdlW+9ceOerbs/vyWx3Q1Elc09JB8KZNMSfTQtJEKWoBSgjqnrVUHrwJvXcOzaWqM+8NCO0niEQYs0Q+IwNTZ8o12dWcu//f7MleX5zz2x95HdmKjgYB3/6EsPfOunG8ffP/f2G69/7sEn1wTXW/jq67fmMt4/3vjlxyee2IFuF3/04rmtW7c/eygeUIzGoAwAqAClWr69v28rAdUAFcsMoJuFLEEygLgxbue6UKNCztlcbFGdicu1wYkRThAApl73VAApyKgwAaDgQRYEcoQsC5ExSbW8kSEHTqzi5TeuR+3FX//cg1sqyFL80pMDS+3spavNr754am/lXh7DRJmXmnj5WvflM7eOXVz0Gm0fGd67dfjgDnf/LkwY1AAHpF4SYSY4DgnMw7tqZy6V55ajn5zEo1MY5FLE5s6u74/fSeh1QglRlBGCIEJuIB3EN9fw7bdns2Tonj57ZBKTESpApFJp8N9/fuDANE6u4uuv3HjzZrjwnblHtze+/ET89DY8NIDSA7WHKiMHJvtKihuEb7176/yNG6P1vt96dstTe1ByOLOKd2f89y5dmk0PfPIAXA1DMYwCqs4iy9uWLVQFBDKGDACCqCobFrCLyhmAGKY6SNGKwohuoubBbBy7JKmAuaBmICp0p1LRkEKaW1FrTJ5KmvtyuVyKjPdoemQRfnQZX3t3/cS5+X3V9AsAgBqpS+i3Pjm2+N3Zczc7f/7axcpnd5sIr5/d+L9953iTK/tHa4e3DDy9a+TQLlRt4SyIpQBy1li2CDmEMsele3fiwM7xy3NLr5649flD449OVC2gvs0uKhJzvSLRzwoJbCAFiBkQFQ1Kds3j2BV/sxnIxUd2Dx0YRprCMSjmpmCyD/ZezF7ASnO9VB5L1b1y+ubsLb/w2OiXn6gfmsSh0a2xxQrw2lW8euZ6xeKXHt/36H7YNpxDpQJva7NZ+JPXz8+vTJknynvraDBiUJqm5TiSEApHh1TJCFPR6QgfAhujjBQ4dglXF1obOTKCjU3uBSIcRSDuZnL2qj80ZOMYRXZfNkN3VeXC8CK4iJ2LQ/CebNdiDXjxFP7s6M03Z8S5sW600fRopxgqUztgPMZ//stj/+yPWpduLHz3zcGhT/aXB2t7dmwv1aqfOVB/bBt2JKgQWp2MNKtEHFTEOGXkQCCoahTykcg9uq/y2un1m0trr18Y3D0ZJUBECv34iHZTpRBBicmFnGIgVg2ZF2NvdnD00uJqK63HePSepAaYGIsOb69glrEOzHq8/OY8ZZ0Hx/QLD9ZGJ/oudfT/85Pr//q1tAM4iyXgVBf//s3pRfR95uDWz+wDB7gy0gyRRYngbGmRBr51evl//OMb37+MRWAxQ+aSLoyChRjKqsoSCiItYjbGKOCB1RQvvnnjwq21TJwaK2AYhmEQBUHHy5vvnrgygx7rE92+5wK9wQIbsgDNgK7PujnhquBPjuOf/fDq6ze0FY/kttJOxURQiywH5RgrYYLw9z6/q16qvP3+uQ8u4eAu/F9+c/L/9Jn6r+7D/hL6CE5RZq24OE+D5Fz06rYEKSOyro9R7+iDk7h/R5+n9k8vL1xoowMQJ5vb6G7E60dCKCKQKoXgyAPoAqcXcfZWU9LWM4entjbgMzQD/uDrx/6H/+fX/7eXlt9ewsvvdq5Oz+7aMvwbzw7+/gv4L//O0LZRk4XOmUtXVjeQAsvAf3pl+drM0pbhgeePjE8mKAtKDMdI2yjFSbOVdTVe5cbZZvI/ffXS/+Mb7Ysp1hhNpS45T5GQVWJVLRAyAAimCBLE4NrcRksTTspkTPDiXGwipwwQu1L11nJ7qeVDQT6wSewFgEFC8IE4LiNAPGulermDP3xN/+WL1y61yxuo5xynnjKybOEM8sw7Qr4RxiM8MIbnH9lXS+zXXjx64iamItxTQz9gMtUsEMHEcVOcr9RWXHkm4KbHLYNmgty4rNOucBhJ8PD+ylAVMyutE9fRBgIb/XkYhw8lRgqQtRlCF4CyXfZ47xpmW7q1YR/dgYZFBZhdw6XplTVp/PC9+Z+ebi0uryZ99c8/tfXRXUg8PjmIHb+6++13Lty/d9dADW3gm2/htaNX7x0uf3b/wFiEBoO03V5u1gZGoghCiYtcN+1G1ep6Xuki+tOTzeMz6195bPST91G/ogpYogimIAEqtIFXycGml6yKOogzL6o5mKDcQ9BYKMea1HJrA3qwJ0CZibVoTSdmDkBwlTZwZh3/8ZX1r59cXo+GERnKlRTGOHC83gZq6CuZNM8bZdfuSp/l5/dHczO1755u/+vvLxz8e8OJhc+RWIrIdAJaBrPAu2dwZqa5sJrBxUPDlYe248khNJIyrAVweCce2No4ebN7/kq6sD3uSxAr7Ed7mPRnqdS8iuOisZ28caspLsxupF73jtbvG0MZyFIMVvDMww+4a91zc9l6S61t5Jnkgvl1bCkjAfb2Y++n9jiGBy4t4a3ji1mXHtoy8Jl73KCFdNIy+7iWtLsbiasZ2G6366pJu5Oz6+NyaTXtO74wl71x+eZ8+beeGVdC1cIAVoudDgBMXJRfshy5KXUDCYGdI6N58BAmx6JodzOJyjlQdPMGwBZlzWIzqQT1XXIp490Z/JuXZ968mTeTrS1PnLdi5jxIp932kTEOCSPfWANbH3wlKnHAngE8++C+129cOn199sfv9X3poXggQkjRcmgSTs7hm2/efPfiraWOhhCnARnp2T2Dpacmn9nuAHCOXSUc2tZ/dubWrcWN+Y14LIGlO+VBt32HjwhJQB5OKcmExEXXrsqV6elqqfzIvol+QHL0RYgVnz888MhBnLmBoyfX3ru6er7Nf/zDWzcuNZ7YW3pkD8bKSARBMZ/jx+9mt2ZX9kyM/cqTIwOCSOBcDO8hvhTFNoWTEHPU8hYugSJLEQlyWz+7rvNHb4Wu//1f2uIDMq+RMcQEkiKE6DlACh9EiIkZEqRgULMMRgggm2RqMu0Rb912GaggFxQh8ezcyVX88x9ee3+pNGsGM2URJJobH2BJVIJSlgMKxxqgMAYGIhoTPbgbT943+dKb73//7RMP7j8yVIGBbIDPruN//fHS2+fX6uXK3p0jg6WklcnZtfTopav9Lpua2rmFMeDQBp66r/9H51ozi6vnb1TuHS55hQOYoEVGDCQ+sDEfERIT5Qo1ScpoAxdmVrLUT4xG+7cgBtQHdYYIscXWCIN78dje+psX6t8+haOnLr/x7q3LZ7rrT2/70pOTkiLqw+k5vHHuhkP2lefGRyLUqajtKsEWO4IMDMgqkRaETSAqjEeUc50pO3pp4eTNLUfGUXHEbDbjPNxBslmgmlAU+3t0j4DRHrFCDpvfvjvtwTo3rZJaa5cDTk7jxGx6y9e7cQwGE4iMihjAkLLk1COYtGBLxuReYseZR8Xh8w8lVy/GZ6/cev1UeujxGMpdxk8+WDl5faVcH/zKJ0afPYQ+wnKGn1yv/eDV7LXjZz//5M5tU0DwjnWs4fZsGXrr5OVrcxvLoVQ20NBzDu6kxvzQcaACzbOJ2Jtu4vjVJVa/a7xvez84VWPRAZaBRSAHasAg8Ok9+D9/Gf+7z+88vL2+vLby0tvHr6/BJ1jo4odn9cLy2n17+57ajxqjZOAYxBQo9qacciwOsBFUI5E4BCPCAlUEQAJSsdeW/ZmbQTYJH0PhNDB9GEQU6aOPNi/03laowoMKXWd+hplTYHO4DnB5wS+2BGyMwgSIwHM5pQiAQx6FthUEoGtK3jghQ2AHRPB1wqND+PTBbUll4O2zt240sSJoAacv3dLuxhePjP7aQWxdwy7FPREe34YD4wnAH5zpBiAIK0mNcGhH0mezy7Or0xsQIBS3X9CUihhj7s44EMQRZ0AOnF/ElYVWNeYDk9WSgkPGcXx5HW9fDzdmZg+O2KcOjjaMMHgtxXOHsWfXrj/7jmeb1xrwMU5e0vcvXW4MNz752JYKUGKIT8nFAhgDMHJBYLgoyvM1NpElEZD0OAFFFKKu6/qmV1NFOeTes7WGgaD6IQ0TCVTlNuvIJsuPAAyBsgTlXrJYC4a1Xk+kkgE4A7zBSmbEVZ1zwQdRw0BOVjRiBUOteFYIkBvLpAo4h+BhHTMgwHMPD7xyaeuFa9feuzT42MGaApbJMibHUWbUnGpGOWE0Ql/JJS6JS4kEOMMB1gA7RjBcj+c30uuLONIAGOF2soUITB+i7wEQPLRtuMWaZ8CJGSx5M1ov7RhEiWBsvJjjpQ9W/9V3T/yHH55968J6arDieRm4sIEzc9jfh//xN/b97tO7J2Msebx18Xy7NX/frpF9W2A9EiteWqoqAhYYBaswUG8ghMBasN0FaA6EAv/i1a6HeHatS4BjdVykSdlLj7KqyC6wBKgqseA2lkOgASoiEqBeCu3mi7LhpjfLgQCDVoaljSyTKOReso713hX8keQy5aDMzKb4YYaCAsETvIUnFqDV8SMlPHxwgkl/evLmWgABu8argezxq7iwjnaFlmJsRHj3Oq7PrIlmjnNLhRhMAIZqmBxstH1+dd5nRUpMb9MyFnkFfxekSwEKxEspTk4vd0xpaqg8XgUrKMLJK3j5vQvrXTp8+PC+PQMMJBZn59N/9sdHBxr9Y5/dt2/Q7h1NlhQ31/DBxdlGVHl0R6VBqFjkPnVRFCizsIaYiSxrAmybQMlokJATeUANBYYqSFnUeMQtzxaIWAvshWEY43outQACRYAU7F+FpVMSZSaoAEFERKDaq2sQM5EQGQFE0BXcWsTc4rqSQRDHpvhgIPQIeoEkds6Be2R3YpjTDC6CVxhCo2Q3gKfvxcnjQx9cmT03v//IJB7cP/XKhc6b75w13dF7d/a3NnRhtXl1eWNuduHQttFHdruEe66bBQYNJodr3QsLl+ZW1rLhoQhCCAViUANgRITvkA+DShkqOcytNVxbythWtg/Xh8rwgi5wcrZ7bWHp/u31//KXBj55CLFHyKGIy4N7Lyzwm2eX1oGgIML7F3BrId4/OPXAMOoBiWK9q2qrjkKMLhNAcAgRsHsCfSbnkGpApiYnq8RgDsRCNrhyICsAaQ4JPWDTpu3XHgm7ggR39SRoKOxYYYao6FxgJdbb8LRA6ApuLmKt2TLGqCpb55Vz7TlXBG9Zy0kcRyAgpiymtpVW2XciLbjvkCAtZdkOhwe3DSxlePtSMMCBKXruvskodF587/r/9Odn/uCduW+cXfrg2tzYUPUffHrfg8MoK6A99ooysGW0TOWBqwvZ3HJPcYv0in4/GycxQGmGEGGtjVY7d3E0VC/3JaAMAdho54G4rxxVE8RAxUKBWgOub7h5a+NWJzQJMbC4jrdOLuTU2D42OhyhDJDmzrmCNFrVExRBiJUZ/TVYbsPHQQUaQQnqC5MTlOB9yDTLAQJbU/hxUgBTUaQcCltjcAdsjQBWsARCIAFJz9PDR2luVUEWG51ulmXQkHWD5UgoYgPxYAeoY+ZSwklx+sL4iVBUhioT5T6HCRVrmXFo11DpvcbJSzPLj20ZjfClpyrO3fvSqZtNU0o1TG0Z3z22/XP34z6LEpADDFhAIDboaN2Uy9XFpbWb62iOoAwYCQzb0xRMd1ZmFZpbF20Aly4vW03HBhpj41xs/SAYLMcx0WIm8zmGE6RFJciCy9TSLKmP5YxmwPQKLs+tmWr/1FaKS/AKI6FkjPc+twk5JAp4NRQTimbBPIgq3aYSFqagYMtiNXTXVjuhX5MoFbHMphAeQSU4GAdIZgyXggQowEYhUCYoazAKQ1y4ixkzU2RBTIHRi7hIEVuNDCMV69gYUoEGMEMCwJyqcdpNqHgGDCgGK+CESAFrDcQqSBUDfdg2Onj60vTV+S3DU+jz+J0no0/cu6MTIIqBQdQIBhBgUUFArOhjuAAJqMfYMti4Obv07vTaQ3vrUx4xBwiRQolA7q6uCgrAaorZlY7m2WAl6q8gAqwgttg1GvWX6eLVa2+c3KG7MFpG3sXlVUzPzPaXMN5f6gNy4MzFpTTNd03Vt44gEqiEoCGKEgrie36Yh+QIJWthGcYCqVHYIshhCoYC1BCMCrc7+UYbaeKKAp0BDCggEITVCCHABbJATqoKT4ASi5KApTc/oacnlHreI2/qDcMIKpmoWvaphDzYGD70wt7AsHHibGeT6aWQE6P3VAsAJSr4qWsljPTXToVwYTrbPxX1G5gcWxpoAd5jo4mL61hvo9XOO6121ll/YOfY4ztc2bAErToMxoihcxvZco4pBvIAZxkUVPVOF1yJM3AGzKa4vpKS2olGPFlBBXCQPOedI3jukX0/PHnjxbfOvHp6vFZtGA3NteWFxdnnH9j++BauA9dbuHZ1rurbB4aSXTVUc40obBLusFNAA6gL7kKFUbECF4xqLDA9340DKxk1qiY39dW8tdZG6Efh2wXAMpxI8D1HKzVcDEsgFQqFKqcAKyR5T05FX6L0QIW9ZEWPYFJcxUeRUCnkXQQ4AAFsVUBQsI3YJHfUTAsoLOhDsfUCz2oJW8f7LOTEuUtP3X/PYAkrTfzwTH5lHak3S6udmaX14BV5NwotNG9FlvePTUZxcKz91k72oczZwtziWmtY6r2nikiLa/0wwSqgHMiA1Qwb3awW24kGGkAscBF8mk1Uo08f2dIV/tH7168tZ9WB0F1fHS6FJ/eM/vIjje1lKDDXwsxqty+xu4ZQVkQUnDGpIHgQgwWQHFZgTRAQYDyqpiBIR481sHiKYQoicI9SFnpFbqMFa7eqkhALIVCvUISi0oKAgh6QitEfCngLmAC7WSsjIu1xf7IHogTqXK4WURmBQgAkZwlkOahjzZlsAXCTgqfujvAfm6ISkYh5yzBqlejyzdnF1j27ymi19fV3z7w/00ziqokSCVyrxPWKGS71bb1n6J7to7VSAftUA9QqYM0XVjorTUgDHsYqEUNV6M6EXsFY5xlrLQTp9kVmrIQEMBAF+6CSYlcf/vNPT37iwOS7l5bnVlv9teGtg6WHd7vJEiLJW+yut3F5QwZq8Y5xFJ01HSGOI+/VgkhyQEEOKAWyAPoYu0bji+st1moQgQhDhWyP40rgjBWGEMjDMRA8lDxRl4vsFmLRAr6pTMqF69HLFJF6q6lFUSTVYHqOnoCpwJgIEgeyRjoKm4Ch3kdGSDMLGyQY6ZZjS4pAyGHj235HL7cmUAX1RmwM1zDaqJ6+vj7XBIZhSUuhu6WMXVOVnVtGJwarO8cwWsagAQQlRqRAkIxMG3AlUBxtrLSWmvCAkFMqGiXu6E/SzQq7AEtryHzeV9J6GapgUCaalGP28CnqFodGsGN0wJiBgseiAhjknazTTtzVNWwEu7tsdo4hsUDwmdeSK5EjFqCwG4gCjDIY6Ld4YM/4yxfmDImHV1Ei6Y13UajAQAxgAzgXRAWUiYXMJqYbBrmBD8SgzcecihjKW/IGnjdxuaSkgt7TCyHmiAHA2RiFsrXQrGuNCUKkcCqx8aP9ddeL/20vzaFFp11vF0GIiKFoxBgZqLx33d1agezAQJn/8a890g4YH0CZ4BR9hDJgggdbDSIhwLoCmluOUYnsUuCFNXQVfcXQDUIBu/1IMGsEQlhea3c1DLqkViugvaSKECAKa2CASFC3iIAqvIQ0I26r1bhvQXBtCSaJtg5Ek1HBa5snjqG5oaJQEilJChYgAgxQBQ5sqRqeU2JAQMIaABKUQIjZR9R1ecH+nEEJGgIRQFxgSxjEGWkOiu4oYCrBq3qCJ9ZirktA1NP0AWSESAnIgYXlDDmcsXk3ix1l2lXEQBQEjnyjbHdMJAUPrGUE3J5hgQ9hCCSsbAl9MQbqiedkelHbgUYTlAjGwG4OjXFF8p6tJ+SG2XCxKyww4jDA4ZrS0rpv5VZsT5MWBdqPaFhVeMVaGoLayLpGGQxo8JERx4gMLIM8EkadYbob6K4bhuOYjWNCu4v1ppLy5EAlBjRtMyS2rGlXfDFuqpgQsanQBTEwVINDZsgX5brCFzcIBiF2gYp4x0BcDMOeJEBZi8RSYcA9Ud7LpSgAMIIppoCQ8WR90QdITouoV1EETDljpoXzN1fXWqkhQDKjwbCoqrDz3kQkjYgmBxADxvd+EXditZWhohpCCEyoWPSVXDBubj3tdGEVVQ01CYmXGHC3Cag3QfcCeIUPaoHhCoYqLIqVtk/zXoUfm1nHHit7EduKQVex1kXb23KU9DmYDMaApOOoWeJurGmJfBSC9WnZlTQeaOUlD+5Ju4v2ervs4qnhYQNYjgkGgpithZVN/pbbPJERacj8YB92bR0Mfh1ZF67kvVNEFr5kfSdvopzMtLEGNC1tqPMuyYJnkoiRZyiVYCxp8JYJIpZQ1P6DT8EucJSSFYuOgBxCgDEQQeqx4U0beOsKjl7dyFD2HrFhVgleYVzu4RIr3c5Q2VUNEoC63SJ6ECq2cCGhXmrdWiMeCTDQV7JJZXGtlXsQQB4MA8OZ9oAomUBFnCIJUgpSEakZsjnKjFLshc16N88IuX7okuDOnVSIOfVopsEHRMQRgbnoZ4wAI+I1ZDBMrPAAmaDFeCmQwAKSQXJv2dXKzICQAccgC3YfskpCHMSgx9lhjcYGnzwysquR1V2WiGdXCiaOEtdsr5UbjWvLze+fmH/tFm4AK4wmWYpKeZ5CYSzSDLG1sTMacutYFQxPDGJL1oEjYwwIVQfJEELezREiNBldh0tdfOfdmdm8qnGNHHLhXGCiUvBa7+N0dX68lD24q38wQaSIYgvkH6aelDdp5Da1gsIA1oDJBGEvIIJhUPDwOcJmX1GxpHkXISPJNeTBeyq0onGBbDv1ub9dMBMmBoy9XUgnwAB5F51OGrxGJsRAZpAR5+BcnSM4qxKEYRHZzKsziDVEIVAQtonvotvtGmf76lAgNQA5BlB0sPUseDCqBqRqiADWMvCrD2OosfWPXp55b2ZpMQxH/aWNbh4Nj7R9QDLw6kzn1osLTx8c/uQh7DTIu1y1NvMpx7GkcGwsYAyFwq0jK6TCCljJ26XYJwIb4DSNK3G3oGDoYt3glbO4sBZ3bKNZrL1NMlXniLJmWFvYUtr49H2Tzx3AaAzkRUt9MfYHis0RXcIFzuq28Jwz1pkQgghAYOOJgwvMSmmwwUMgsAJ4GAIbHyg3NkOhz/uVpd31me+lXz+m+1wVxiAEdHOIGrtp1gpjEhfTjIRESRhEMI6gwRgPCaQSFJ1c0rRjo7hWhgIiCGYTPQog9ErCIA81SvBamKU86dDn9thGeeKf/unFC+1ssWttzADy4EvlxoaWjs2uLqZLc2sDn9pNT05B4Lz3Dogi9Ncr1E7jUm2jk1ESKRU4b0WADflwxY5XUQIg2UaILy/jz16/eHXDrlD94rJf9pUOTJFGIAtp55K16qbTbze+8PD2X3k43lNHLL10FbERFWhvyJMtUqFEgBZ9TwGInLFMEnJVKCH3PmYCq1FxRem5mB8IJlElAlvanD1pXWJclPng/W10ZCBEKnd6dwoAPiCICz0oYpGuViseCoQAEU3KKSBenaWggUiFKRUrBq0AH9IyQpnhggd8EQz0CLJRCNyqyQJ5hRMlUmckTJRts4P7h/Df/Mru/+VPz4dl9nag3XHlOAE0hNhL34Wl7vz69Ws3eebg+Gd32/GSDR4lh337xt5fvqEaoKRalMUAYidUYd05VNk9ioShLmk7nFzFjy5uXN2Iu1EUolqALUWoxFhfb9ugfSXEaA7S+q8/d88nDvB9DcQ5FCAL36ss9IajANDNbHyBMS/GWBjDjoSCLzR7oDJIoZmGDGyJOIikXhxHIQQEVmdYYQFHKDuxTEVtpbBH2pvWcjdaCKpQIQjx5tyWItxGt4tSPWPM5zAOfZZyBZMpPBZPLgBCrKoiQgqGsEHRcmvguFDfAcLITVHqDFArSg4W3bRkTL+zD0zh//pP9n71pfT1MzdnOtpGNdgqacVTEuJyy1TeX5hd+fGpWrjnc/dHkSIC7t+PH3xgZlbbzvXlCMVkRgQ10Cj4wUpSj9HqAtZd83jx7NL5VqXpBuNyX6vdLZfQXl+oV+JRF3x3udxtPbpr+O8+f+99I6gKkoDYeAnwYtSQKFQDE4XbXniR5qBilCMEUIIlMeoLmVlHAcRgjSoZqAt4NTEZAZitLRIIAiichWMPkgJdgzthHB/JOFCv4ZFDZr03oZc4ysGGXdu6FeDiCo5eC9WIn9tGU+VeL3wxpAs9hzjyntqClCJTTEFRTwA03kydwcMpglVPIoooZ3QptjGygCqjxPg/vBB/8cGd33ln+fXLzZmmpPBiK2sZtQPaqBhHPzozv2Ny6sAAHDAco0+7M+SU1SKHUhAhZUegELhook+wDHzrPXz/zPKC74Otlsga32XV4aoth2aDO/v21Z6+d9fDOzFeQh3I04yMUuy8SqZZpBFL2BzD85Gjl3zgXjLQFgl406vdKQCOusDVNi7OIvMYHcBkBYMx+hgclNU7chEAUVaw6RGE0G0umruC2UIxkRZTWk3xuARFFoLG7u0r+MNXz5+aTxPN5vf2/aMv7Ok3QDAwm8xqBCb1SpnC9wJ7EoFh7lk2AhO4aA+SoMErWWUOFpZgFOy131IAHhrBjs8PPDk38NppnLm2em113qbIoiiotLP8zGzn5MzU7mFEgF9BlLUT22j7zBpPSirEzNZYY4wznAOrwNHL+OrLHyz5vqRRy5rrZr21Jfb9pXjrcGPP8MTD+7BjGP0G/YAVaLdbiSwgrawL64yJJAiLgI1uZll7ORHYwn8o1lEAoWJmJ1hhGV4RgJstfO8D/603Tqw2OzvHG7/3woG9Q6hUAPWqGtlepK1MtsDD3GGCPlR3RGQKIhiADKux7VyLtKZFYJYUeOfCyvlFK307ltbmX7k0+9g8joyibiACSzBAhfNarOvtdgZ0QzE/ybQ9yrErhrGxBQhWRTRASEgZgYiZFYGqDgyizRseBj4xiiOjuL7ROHuz79ry+lw7v7qY3ZjvbHTC66enn7pnS90iJlTZ5K1Wqa+ep01YFgGTzRWBkcROgEXBt356bblryomJ8lsjff6Bseje8fqDO4cn+jBURwW9vEABV9LIpeoDlKwlsAQ1QkpOQEQFTLNHBSAKw0w5DCGyaG500+BrjRoZMEECnEFHceEWvvnWtbPNQW/d+uLGN964+vu/tL0tqBIZa1MPWBjj2BgAFQPKQY49yBGYN3kcCi88FI68IVENqrf3b0FvbYh9N+U4c+Kr1YSK9sEi0tegYBtDSNMgi6vYVYd6sJU4LmUh5F0tV2za6w5jpcgzcvQKBgpiLhJxm94nerrXAmM19O/nB9FYA+ZbWGxNvfb2aliaWZrDxFZM1bF/on5itbOWp8zOGAdj0mbLclSKdKLfNYBzC1i+dW1HbXRwanSsZh/Ybbb3Y2cDNaBvM9NT4CYKgKuH8TAFFKBnGkxPVUARMQyKDAWsi0yhdAQBaOXKxsXWxAZQxAaph1fkilYegimZpNZtNtc3Wp0ObBmu8KgMApAGr6rWsi3Yt4Dbw9fvdhzYwhmGepD2gJ9EgMaKI7vqZy9evTJ3bOdA+Vcf23v/OJIAsPiAnFWYbJ8LcaUt6cJqsNsMk5J6kCeyUTlKi75MoAl0A7qKbBOMFRvYAAaIEAhBkCuCggSdHBsBmSLkIEGjD4dH8OAXGxs3GzvrqAGVGp65b+Qn169upJlH0g0A2hSj05ofm7KP7UUtIL126fk9pf0P7RuaQsvDWmRAB8iBjTsUi9kMOQAEgQ+bnaBFTwqhlgAAA33AgI0i+J7dZ+8p2gAttrwq1ZKo5OAUCHlsXQUYrmHbgKzdmg/d5TI2Hti+YyBGDLCSDyGw7RKCMUpSSiK2t1PCRVboo01kIBiHUmyIMmwORAJY8jR27vGd6Pu1+2/OycQgPziOOMBqBiixBRkFSmVUS1WfL88vLSuGDSh4E1yUM1ZzfHA5vHF2eiGj1Va+nklHrJAhIkvQvGNJLFQBD86EssBBfcjbeR4yLRE5TTsOfmKosXOs9oVnxrYMoxahSNQe2G4O7ey/cXo1uBFkgEjCeWw2Pnn4nqkE5QxPPbLrEezaAK6v4Nwszt1Mp+fm1tsdZsuWmIUNmMSAvPchl+A55AgeQaCqQaVwXSNr+/vLjbo7tK3/hQOVSWsjFUuAwBta9lhc7+ad9nAtSRgGCmNERIT3jeFXn95fffvc+lr74J5dzz86MlSDeoUVMvABKaPYqH3VSlTUWzaFhDttUmH3mBE5WPhc86xQvgHWJQRwjnuGcN8wl4AkF0Np8TtMxioUKDuMVvvOpOnC/GqG4eARRWYNuNTCK+9lPzp2ZSbEK8G2cpOTDewUBiqkElNkyDsmY0wg62FzOBGJrOYhU0+WDJz6LLu+gOMr6ds3Lj15YORzD9d29WMA6C/hs08MvH3x8pXAsI2SxG5t8Yl95RceQLYCV8cycKGNn5zBK6cWzs7p0oYYcuVyrRtSZhALIFwYfFUIjDrvNXgVQVASqBCYEVm4FvENef/60tJS+OLhvoMNtoq0a6SEpTZWml6z5tahcgSAqJN1ojhKiEPAJ/fi8b37ul2UElQIDvCcBzCIjUGWo9PKKM9qsY1sb89wT93Khzup6NwmRSUx1oROnrU84CCgEMAGVdf7ZASQb8ExVBRaUH1aQUzYMojI54trnZZHwyAjXGrha6+3f3zi1pKvt+Ja05gMCGRRVO5FodJBMCTWsDGG2PTQtoS2IJAQcgqwzJpQYJshZHlr4b3Z6wtrX3ps6tFdKAH7R/Erz93/P794Gahx1tk3Vv/ckYEJh6SBWzlem8Y33ls+Mb92bZ03dDwaiDpNbGS5jfvJFD1nakgLTBErsjR4ksAqm89yMcSly8oBpVJyo9t86dzGyGDfVB9qOYThgYV1rLeyemS2jxb8nUjiKA0ZQl62Jc5RAshCcoBhLSyzFHaQuJVifb2reVaNYcyHfVQAQHcKScFAbFAv2zjirg/NFOpABDClAQV/vc/z2BnEJXS7cA5KhfthFFVg9wj6Yp1fzadXMTGEtQzffx/fO70yHwaSwf48LVIaOUtQ9VoU2giwUSimXxe8MsEb8SqUpwZsyRqS4KWrgGoi1q6hD31978wtXP/OlROHtj7/iLmnhGfuc9/8YOz8PFk0nzoy8eh+pAGXPV76QL99av3UbCpJOcRldKK8A9WUjfoMYEuGIkeWqJfKUi1YHbWg21KGgIiYKCKAhIG2RmfXSj86l+/td49uAQQtxcIK2s3OZH/floHCxQji84iYIxtUWELVWGeRZV1RF8EaIPggxMFgpYvVrjfq+2JEtuc1MQMaSMJH+mFiQsKol2w54lTR8bid6S0opLyKcy6T4NSSi8GCILTJEpEQJgdRLePmavvsLO4fwobgg2vL8yGyg/3z67kjtb7rQubgIwTqMQGQpEJsjGFTXFbI4dMgnEuUZgjqlZWZiFQ8fGop7tfVELkcZH/6wZW0OUSPNsaGsGO8fvnmfCWSxiBgMdPG116e/fHJWxs8UuurtdvzLmzUkWetLNhW8J7gSI1RtWwiY4k1qAo0MDIJvgC9KozCCDOzB2wSS244roak8c71+R+Vzb1bBiPGegfLLXS62d6hvpFKMecxGDZgChLSLC0nZckz6XTiOFYqLL4YtoatB1Y62EizmKUvQlSw3d4uWt3Jd8eKrkfZYetY5Ai52NUWqAEJwbIppsQzsQeIDQkMO9Wu3hF3WcJgH3buHDx/fOHkLf3ifbQW0LGGy5Fx4dB+N1xCw0VDZUxWMJKgEaFqEXOvcOkVPsAL8hxZhjzPO2lIhZtKnp2J4SysIAS0u+i2vMk6C0vrF+fkg/evzt0aOvL01MYqXLoWu2hD8G4bX/v6/IUby0nSGNb2SNTpH8yGyqWR2Ea2HEr9cSWOGBbgAASEEAp4Vw7T1mCiiCxIQKKRaqSaw/gKQgldxXLASyewOG3eunztK2Fw3KAdcGV2o5OHB3ZvrRMiBVTABIHCGhdnXg3YuhiqRAZEUI/gQY4JS2tYa6aR5FuHYQARWC5YT/zd/UmJQwYMlNBfTVa9rLYAgCmQGqYewiagp7wVIOYCTMdgw3DAcIL79k1+84O5mZX2XLey2MbcygYnY83W2iJXqRylDCmDW9AybBX1fvSX0F/6kNK8V3cBPJzANYFVxVqOVhsS0KhgOEECRLAxah61iwv4zhvNlz+4+sEfpyuZRJXK0mrz1bf1tRNy+ercnqnJfeMDj2zH9lEM1BADCSDAiqLl4Qg1iz7AAhamaMUNQAqTA10gL0r6nrI2Nrq4keLSLJbX9OpyPrsEFydIGnMtjPZhdR3L62tJEu+csGXAA8Yw9+CNRns23xWcYdpDo7Jy5BXtgMU1ZKn0J6Y/LjIzyrfR1Lhr8KLCEIarGO+vzc10lpvwgGGC5oyCUpJzOCnQVQCTZSIv4pgL0FqfwaGdleF6ZW5+8eS10uAY55QEsYHMSsdPz+YIZCRz8GWTD5R4y1BtcqA8OYiRPgw30F9FJUZiAKCbY24Dl2ZxbhYzS36j1YTm44O1baPJgSkcmsQEUAMeGsbws9W4et9/eO0i6gPLPiJbu3ijlcRzj+we+cK9A4/tRd0iAAvApQ1cu4aFVcyl2OiiVkJ/BZP92NqPrQMYML2hgu0cqymWWljuYHEdNxcwM7sxv95d7FDbM+faCZqVGux8dbi/3gcC1pYWms3lRn80OQ4GUoIiiuGL/CvhNmVIDFIhZIqIWAy6AetdzC1mPgtTI5XRGriYKFoQTyhwVzAbBDAYSjDRX33z6vriWujABDKs/jYA18KFgp2NRISZXC9friDvI4Opuj2yb/uxs9MLG1MT+zA8PLS8iCyFV89Rn/fwWs4kNEM6v5ZdbXZqN9u1SPsrrpFwYjQ2FFtAQifHcjPMrmYLbYiJrGOoXF1ePnk1O3YqvDNaev7hiSMjSAJGa/jMw7i2uuWP3r2q0UhS7W+t33x0b/8/+NzAoSo0hwA/vYZXzy5cmm8trOl6FnW5mivDpzFnVZsOlbF1MBmrlx2k28nnl9obHV5u02rK6xm1FYFYXNIOplyqcZ65UpQT2t2wmvLb56ATODe9GtKwc+dY3SAqRqQBBWy2yMrqZiSKggicitQGMmAjxdxa04vsGB8arsApmIIGLkrMuNMFB2nBhVS3GK7GkmN2tbXY7qtHXLJRpBk0QIRYjBLBQzQXuIiLYchFDidBGDb2U4cGThw/eW2uPbhcq1vYbpvFRUk5b3UoDxoCGWMNU2SJTIc089ho4Wan6DNUIsqC+tRrHlQibxLjIhhxpGSyDtx0x8/dzF+7dv7RXcO/8Uj/vn5sifHbT8ZvncYca6tza/8E/1cvDOxK0M0w3ca/+Enr5GxrY0OjeETZbFDWDaoqESdqypn6hY30zJoY6ogPWeYTrgBW1ClMRpQhKFgVcWwy3zQUyHK3q4mrnJ1eaK2sZp/Z/sa1DqT8wPZGH5CETXeZuICthKKBDTAKJhiAFUTIFEpYaGGu2c4tbZscqBcglDt5Ce/MghdkYwYoKUb7SpFLFlc7s+t9E2OFlrROfU9bAkqsCB+SJG8idSDiGPu3YOdo/Z0T50Z2HNk3hvffu16qjaTtUMm7Dp6MMLwJwgZGKQRNSpVArKCcWIk91IOINLEhMqrw4oPxIXHkSIv+oywzOeo/vbyxvrz6uQe3P7uPdozh6fv3/OkbZ4catft2T9Ydmk2cmMUfv3Lx+JrrUKVC1gWvvtmQnDgwc6uZErvcxEI2UJSxsTFHMSjLOWTGdzSIY1M1NlfOfCBmDiFJkuDbiSuVLKTdOfDEnuk2ZtZl50B07zhivd010StqcFHYLbLnFBQCtapUcO6lwLW57mo7rdZqW8asA5x4UBARw6yqgo+2YyoCApWYtwwm/f0Dq618ehn7x4pfYkcOokVRl2CVlG3RdMWF9K21nRACMFDCow/uOvfD04uL+Mwh3NyenGxl04tLJaQ7xvr3TA2P9kU1FxIjrOIVzVRWW+n8Rnd5o9NKfRqEIpOUdKwRDVXtQH9ftVIqRcSA74R2q3trqTW7EW62dGYNr62ahWxlvTPw8AO4714+eiaQad27nbWE7x/Hv/7h8RUfVTls7aepgfJkPd4yODxYIQ7SydKVTOY2/LXF9vTS6lIzBOJqOW6UeWrM1ZyvWo2YybhM3WorzLdwcX69KUxdTTOJEl8z+QP31D/3JL76EmxkH9pe2V4FCYIR4wAggJVgNTgJrofiCgDUwMN6UBdoCq7cmM2D7N2+bWwQCDCskMAQ9DBkHy36FUjpiDA+iOF6sjA9f2u1syGlUkF+oAYqvXk9BFVm9UqhGBHDlgrungioA/u2VUeGGu8cPf25ew/8gy9v//47reUJO1Sr3rMj2bMDDUKsJlJTlGwFWAulpQ5Wu2h2keUwBkmE8QFUHCpUdCKBgQgGqDTTys0VHJ/GG+dWPphNT12aQ2cNAzuGplCrVdop6sP405fxgx+/3TTVoWr8xI7Bh3f33bcDYzXUi5JEYDKlFrAmWNio31jCzBJSj+EBTA5hsIS+BGUuKkZQQjPHchdnboxcnJZOV8ma9QxDA/jEQ7hxE5cvXEpMOLi1PmhgRNkFBfleaH67SCi3eQqo0DkwGbASML3cTMB7x6MBB0qVYwaZAoNCRCC+gxGlJ/mQpelAHD+4s3H67IlrM6W59S1TDXAWCrhGSqqA9T6yzosAzASoMHJAEkYO38zsvRM4uHvypTfeefmN6f/iS1v+4ScqnFcswUUoELp9BAeEToicUULVYLyKUMVtA1d0HBnpTXIu0vaFrh6OMTGA+4fw7N7+//T2xrffmb00T//sP83/0hdG0mRwRUrfeQ/Xj90ImR7eEf3Kp3Y8PIURg4qFA0hABDHwQatMJcJIDfdUELYAgCUY04OaAh/2D447iMOBvejuZlU0PTKDECElfO2bM9ppHtw7fmBLbDOJIgkiubCzrmcBYFAUg4pqKylBnaohrAPnlzG9phM13dtAGUgsQgDBgklUiJjuwjgIxJCxMWXAPVsw2kjmV7sLG9AGEmskD8rEqrnkziZZN3WRgbL0YL89c+VAjQjdgMM7K6fOVt69ePP8xpYHKhiIETyg8AKR1BkpkYVjLhomA7QXpYjAiLcqm6pchUkhPcgEFCbzNWvAVK7it56tVfoe+N9euri0kf/hn93qeKSu8ua77dr6yjMHtv7el8e3VjGQo0ESgzVABGANJEoSw6oqRIt0NwBTPPxERCTEAlZVAzJQsNaM5JDc29EyryuWCa+ewPmr82UXHtgxMliCFfLBG2NsgVgEAHiComhchxAZgELGRFmac+xOXEdToh1xfu9EwVQvm5nVO/jaPyIkEUKB/MOuMWwdG7i52Ll8CyngVQIjqDdBYzEAlISgVGDjAUEEFARfxgrqwMNTOLht6/Vu/V+92GkzNnKIRUSoa1rXQHAejq0JDM8IBj3cHCyL7WEyCWazrYxYtUikMUxsYQgQEvQ7PPsgvvjk7j60W10KdsB3U9tdOLKz+l//6vh9NfQHVEzmjA+AJwSGEDEZW9BKgEyvfcEQW2Wn5ApyUylMAvXG4EALNgIyVLRxYLmJ186mC1m8Z6L/0T1wBCHKxRKMJSHJoXlRMZPeHiK9jYFVgbEbitPXOm2xU8P1qf5erY8+KqGPkG0AMMSEoCJ5in5g99Roq6tnpzuLbXSVAzFErYhjyvIsikuFwLXn/X/o5ZEiCtgS4+mDI4b44sza995DyyHLQZKzBmuNJZsrPBCYwEW4x6RckIwWJCmGpCchUqJexya4KAwGybuJQ9bCkMMvP4FDWxsWmqV51Ybtw8lvv7BjTxm2q0m2nmgGDQoIQwrJUK/AaqiXFestIpMHeTKeWDcRg8XUG4UBW5AzxO0Olrs4ejY/fWWmHEfPHdq5LYEGKGCtLWaygAhFcHnb6qtKD+sfgg/B0smruLXaieLkwI6JsvaYTu6S0Ed2EiksMSmEcmtDCTgwVa1V+s5euXV8WnMLr2SKHy44EgCQLTh4LYpCri+eFDJwLCXFoTE8f2hE27M/eO/c+TUEgDVvK2XkrMIKPOALFlsNpCBlQJSEtNhI0mOGJTDp7Vp2DrAxqgL1AzHKObY5/OozwyVpRtqm7sKR3cMPboHrYNhKw1gbFMI9YoFNvDzdAZsCivNL0ZWeEtKCsVWD08wgF0VG6ApyQRCIw7kVvHLyVt5cf2TPyON7UAXiXhamMBz27nmJxZYqzGqeC/Eq8PqZhZVmd2K4/8B2E8ttp/3u46PqLvSyESXLDtg2iG1jjbmF+ddOTS8DOUDKIELwpTjZ3Me+lykvOhZUCp/IGmXJ+xP8+rP94xW5utz55tGNdUBc2XOSCRH1rokLT1OBTWITKjYj6W0NUTivvNmMUDh7Jol9t8W+WWeUcjy6DZ9+eE+dmxO18Pi9XAH6HLhI3GoEjfhD2RRL5gmeVD5cB2Ul3GYILVQdEHp+s6KbQi1yi5WAH5/unpxe2D6U/NrTtYkIJm0VTerFvVCvJsTAZiVBFBoKFw/WSJScXMKxayu+s3FgqjxSQwTwJq74Q9ES/YyQCvAkMUEij7ESDm3pTyzevLz02g0EAoL0/INi+gNSaJeQCxBgA5EgQD3Bd3yWG2cJUyX85uce7Nrkjcurr1/FPKCGAArwoCwGoiJlCwMWUM4wpGwEkKAFZwuhINYoei8Z8ILMK0DWOUeaGF8JeUnwyUNImtMP7hwa64MFsiwPwcBVwHHRhsmAg2d00euNARdQUCKQ6XHnKqzCagEPKXwiGPVOETsIY83glXP6w5M3xJWeOTD68DAiEUu+t0HF93yE3neZCvsGIQgLVBQcLeX4yQVcW5PJGj2yA2WFgd5uwr5LTh+1SQVbCZx4iTTvM7hnDBON0vRKeOVU6Fp0lMGG2HY7HeZizhQXDOIM6O1GA5HYujRTAP3AYzvx7AN7Ws2NH7537ewquihyJKoSCmx9D3IEo1RM5RTAGy0ebRZw4csAUhBxOYaz5PMAMhxFmmUOeZ2xcwDb6ta2FmsWkqNcdjBUpGREC6hJTsg2s2k9B4BJC6JjFCoOGoU0lnSTOcyBHBBYJRCWMlzewA/eu7i0unZo75bnDjfKAEtKzBKCarEAHgGGoL00kHLRYynBIFfQhtjL6zh2vbORY99o+f4xRJsuw+0xzncNteodSvAaemBndY5DAtw3gUPbBoTKb5+ZPXYdVLWZGPE+KZV8gFKchwRCpEI+NSGQss+UvBrvS4SywLWzYcHvPmAeHrHnbm782dudmbyYRGctO4QcSMFQQh44DTYQQB5SoK9JyDA7sC3iZ7MJkCPA2EjJQIjYWBMhIA6oE2LJHRA75EXHgIGnogHCQz0KsykOWqByhKFsEAiBAMAhd9JhaUIlE2oG7pJBAYUwWPD4Ny8unZpevm9L/fmDtd2D4AAhiI2Ui6FvAeRJU0CEEKhIaCP4nK1xJHnQRcZrl+TczaWSM589PDVuUL5js3zoxWweH+2l5t5GZQDwTrrjJRzeOdwoxYvrnaMXNuYEGYiN8XkwlnMBWwYxcg8F1HoltS4NHIQlz5BLyXJVcaAPv/vc3pF68sHZy994bXYRaIOUbPA5ENR3mcAG1kCVtMDAI2jPY7wdcgi0t5l6GrfQTOSUuHhuDYJRb6UgiQwgFIGC9mzbXY8m3Q5KCsHbXoK/B4k2bMggUwSN14KdzvD9d9ZOXZ9PkviZA2NP7Aa68FkAmQDjgUwAFYQejVWBnmQu6JIYxCGEzNC5ZZy8udFtNQ/vGt0z3OPhMObnuA13t2MW/QIKNgBZFd9weHB3ct9kJd9Yee/K6rllpAAMszOhMOUAkAE5bJwa3vBoKlIXZRwJNM/TrnKaSznL7h/Fbzy32/q1Hx2b+0/vYpWwGoCkqmwUPktbeaZ5Ds3BcHBWHBc7TEGFcodyIacP74YgxMW/YOAZvggstFhu/6GPpQS1CqNatK8VUij8cceAVUQC1gAwTEkQZVleRDU+YJ2xZPHtU/jB+1eWVhafO7L38w9UJwkVhygyWnDcUtFbbYRtACvZAlGDYk2NDUI5XNfiras4dnWxivZz9/Ztb8ABmoefK6K7hNSbD9HzgSwrjIaJPnzy/rGBKLs23zpxEy1GDiuE1PeAn8XW8yHLBOKQOcwzbhqsl5PVamk+5m7CphxVIzy3G7/8zP2drnz3tcvfegdziiaQwqhLyDpjyBlYggiEKN/sLycqmrfMh2D520n3zT6DolLsGZ57jxopGNRDqBc7r5htBSp8u1CIvxAigaEGniV4UFui3FW9iXojjSyue3z3FL791tnljn/43p0vPFqeKCFvddkgI2RiciXfy1pYsM3Zpneur7FKruW57dzVVRy9sLa82jy8rf/h7YiBYprPz5fRXaUKFQOTE5hg1JJJEJAQPnEP3t7Z/+q5tbdPrz29u16twOSIXG8gfIF6zvOWSaJAmO7g/CpeP9m5Nb9sjBkeGnjmUPTUCGpAHfjykcpGc9933zj/5z+9rsnW5/diPLJNgRUtFaBnVS9ZZmJmc3sKigKkhNtN0dJjZOi1HRBIILwpJO7h5SF8u5kABO3V/KGqwgEw2otpilDfQwIIijgzxIBQ1A0wBqvA65fwH390em6jc/+uLb/z6dE9VZCCK8mNVSQlVGJEgBHkouoEbD1xEVo4KgJSyomyCAuKl0+lF67eGOmLPnt4x44yDEBsTZG0wcdvpzuEBDEQaF6wWkBhyILhFOMJvvjUgXNzp0+cv/r2+XsnD9sKfYj8D+AgbKOoRfjgJl48ufT9dy4utCVJkiRJ2udnz5wvNT+x94V7aRAYBn796XLa3fPKqVtf+8EZpDufuT/ut6gy5V3A57GFMaSb2WJzG/hAPf8YKNDPBVsuDPUwwMWeIzIKW+wtUA9LXfyXi+SZ6iYrda+RpaBkYxj0SFTAhG7vAcR0htdO+++8d2N2pXlw39YvPTV6aACRoCM4faVz9Ozi0MjoU4ejkQgNB80pD0SbvX/wgZ2BihcmgxQ4MYeXj1/rdruP7N15ZBdqgGTQCKSk8nNnZN7Z6SdADjFinAJQb5ih7BilgIf34Zkje7730+M/OXZz/9Ztjw2CJTBAbNKMiWts6OQM/uNbqz8+PRdS99C+qR11E0Xu3Hz73PkL3307GhrY8YlxJIIx4O8+UdGs9tq5G//uxxduZve9cB/211CLIT4LpsSOyXuBmAIhD6NUzM8tzIgUJQQi2WQiVEJgWFYwDMEW667EWuQYKUB7XjwEBAYbFLpIC02IFKzMLEqQOPO5mtTaq128ehbff/vy9dn1/dvGvnJk7KlJ9AFe0Mzw+tn5b74zV6522rT3sW14YBglCx8YGiLtGhLuxbQEQwG4tobvHcsuLYdt44NPH+ofjBFpwRxMRfz5M71PPyskAGSK1KECRf9TUBDQZ5AqHjvgzl/pP3dj+ZVT2w49AyOGfVZKTMQmtzSX4bWzSz89M9/17qn947/zQv+OPrDFuaW+f9FaOzO99tNT2YPDEXd8uWSnyvidF0ZNgj969ew3fvT+4szULz0+dGQCZVtWJesB0cgW3HdmczE3nyVVQsHXam5Dpoug4v/b25cFyXWd533/OecufXub6Znp6dkHM4MZ7CAWYuMGkOJukZJoiVYsOY4TlbOV7NhVrlQllSq/+UFJXHmIE1uWZceSI5oWZS5awEUwCZAASGyDHZgNs+893dPbXc45ebjdwyEIkFTk5DzN3N7vf8+5//n/b6nlR0LVkjpVnRyakQQ0lA7FT5iuWiHQWqmEQtgsMXBN8Bm/nMNL72ZPXZvwAtrSt+GZhxs6YiivwoyDK0RNbOzrSgxXVlz10rHT/oH+pF3XEYcGccWgJWMKTOggACANvgJcmcI7g8Oa0a4NDft6YEgYBsgIReWJM3GHrSxQu5i01lorcJAZ8Ei4IzW40FUVAvCgUkfY24oHdnT5jL9xYeiDOawKSGG6ZU8w5XoYKeHCdKWYX723K/UHz9XvjqGdI0PY24jPHxiA5uMLpWIZnPN8LogKZEz89mPN33zmXlZceu3c9H/+SfatKcybpDlIgpHhS73GR5WhhFSYsylU0RbwUS3bICDhV3UriQwRZsJMB0yCwAMYHgjKg/bAq8udUggQdpnBARse6bKvUeJYNfmZeXzr5cWXB7NzS6Xt3a3f/GpDTONbfzr4py/dXKlAe15a4PHN+P2vDLQ42bySL5y88e2fL90ooyLBFGPcDEBgTClFRAFwaR4/fn8CXLTE6dl98QYfSQOhRKwEBGdSktYf7o3uvJmtTf1qrYnpqmAKAeAUFAophge3810DHdmV/LdfunirgqKGGTH9YsE0UdYoBGSZbN89Dc02DELgQ0pIwFVwA0lclCuwHDJTogJEgTjw5F7n9/7pI31tjUO3pv7bX5/72xOVqyuYqsAzoAwT3IRGEO7GAb+ifU8xzkEcjEmQVJAKvoQMAAmDQ0tXKl+IsJpTowiDGCyQCc2lDLzA9xQUoRKEZWJAQ5MWRqTIMFrA3532/uivjg/enHQM45F9O371iYapUXznO6+TU9+9bSPZ4I4pJSLA7lb8x28cPLC1Oxqre+vS+J+/URpykbOw7BuS7ADcY2YgrDkPb110RxfLhld4ev9AZwwpq9qfDVX1NcDu6tn80cSBoAwQwdDVCrEmAFpBUSwWA5Ax8fm99Uuzk1emV793rPjNR6MJgMxoMUDEgCpnm1JxI4YLWRSmsaUXNsfRUbx2cVFz0ZKKRqMoAhMFuC4GGmATnACPDqAz1fqdlxYHJ/3v/vTi4HD7k/e3bLXQbkOAGfBJSxBnHAApqcqup5kmLjTnVNvYWmGjqAKTa9tQIFQ0DK0NHsIboTXBs8EU50wxVpHwPRm1OAGVEIYOa8nH8SGcGMarZ25JldzTYuzeGHvkfmdoVH/nhVf6t+7Zsrkt5WB0ER0pEIcF1AGbbPzOYy2vX2z5wenpn12bPz/DfuvpzsOtMElwDU+IJeCta3hveLlQDh7emP78DtSLGqVTa4ESNJdkyRqa8c5BWmP6QSuEiIUQoFKV8UCgNOewgAYDB7pxpSc1W6ITV6f6G3uf3MlTwhAcNrCxLTl6bvJnby+cNCk/N3HfvoH2Xucn59xr45PbWlP9LSQcXJrD99+40BC1nru3q7cpkjRQ8tFr4/d/bcdr5/DTU9cHx2dGXprb2dv88O6Wg91IkGUSlAQLtC1IGNyTpFnoW1vDqStABkoKU8NmigdFQ6Em5wIChICWkGAIpCSphOacM8F9DZ/gAy4wncN7gytvnlu5PFXQMfvgjt7f2Ofc04YAmHaXHtq7tXNb27tnlhbnF3pTkUxDNN3c0NpAm1sQBRotPLgVVNf6vbemR7Krf/KjS8Ge9LMH04zgA6dm8eMLk0srpU1tDc8ebO6yIBCCxMCrisxKarGOKXv3mbRuBZRKQ5PQ4AwBAG5YFV9ZgiUIPMCXHmi7vhr5YHjuxdPjdc0bDnfAlMhEcHBb+6U5NTJfUNr3PD5xcjp6NbqUW2mpo+f2dz6wFTfz+JtjM+emg+ZI5XO7KZRytjSiUSQ1vr4f+3oGXjk78861iX+4WbicXXqzI7a3zzrUg2aOOCev4nIQCUMRBbqKazMAwQDNLELCRDrqJHVgeYg4YFprFSgSrtbgpByQ5irwlZJaExiTvGpmeeIaTlyenp4vmlz0ddhPHuo8tNHsEohXIDmObG3ct63xahazxeDiMs1IuJNLhiMNrbd1Ze7fgsPdSEfwdB8aROv/PHZrcnL53QtL3ZnE1h772ir+9xl1fi5o45VndvVvbYWFKg7I1C4QgDiIi9BxS905daCwPVE1XgzrYUpr4gETYS8ORBIspK/6nicgC2bktTn8txfHppbKO3qafudXGnvjqOdY1Dg+j++/WR68OcIM03dLkIW2VPSJfZu+er9dcfHto6W3Lw5ZNn79iR2/vgURD6aJcgAVQPpIRFEkTBGuLOF7R1dHZhYdU8ZZpbcpemhz194e1uKAA1JWtZcEYBIcwCKEuoGLwNkrxaigzd1OyqzKKXqAC1Q0XB+MEF4ZoQH12AzePj9xcmhpwRW5wIoY7KGdPc/cT5ttxACrrA1GHhAwkIHrRfynHy1emPMbGqKB55YLrmC2V1JJ5m1Jy0f3dh7aTgq4XsEbb830J9gj9ze7wHdeW3zxYsEg/LPd9b/xYDJtwZChKQkMVYKWIEsjlDXkSt+57vBhkFBtMSlSgSbukgDB1CBoV1FITjIB0u5qhfIR88VB/Nmb41PZ3OcP9X3tcGQrR9zAOPDHP1o8+v5oe1f7rp5YQyTY2F6/vQsV4KVj5Z8MTnuVwnP3dv2TI3UtAAOu5zG3ApMw0AHTBWPwBPIa4zmML+DcpZHxudxkziNu1ycSXU1127rrN3WgzkGMIy4QE4hVKfJQgAdgjR0tYXJIoCKhOQoaBYmiwkoF4/O4Mo7JufzUzIQrA0+qZNTcvbH5gR2NW1pRBzANIWEokIY2NJHmLpsX+Osr+NPXLsUTzpef6klzXBx0/+HykrQTmkrKX9nYFv36k217ktAFRG3kNV4ZxF8dPZ9V0X0Drb/3cHR3CixMOxk4tKEqUFKRxWo3l1oWcJfljoiUJkVgVeJeDRxGYJqMkB7tQxigSiVuOybw8Ebcyna+cnLkxOBYTHR1HnEogAm08WyyMn140+7nHuQmEAXmA7x8wn37/BAgHtrd96X7os1A4KIocGMJf/nCsdYofu2Jw9s6YXFw4OKFyRNnL+zcf+hfPtMzMofzE7g2lZucy164eWtoaChhI5OqSyXttsa65vp4KoqYBcuAKWAIGALKh5LgBtwCiiVwCws5rJQxu+zemlueWsgtrvoFaSggFbc3pGJ93c3burClCRkDUQBAnpBjWFrGyqrkMd2aEp0mooS+NDa3JK/cmpuZ6jl4EChZVybErex8PBl14sm5hflTx+XBRzrTMUyV8fZl/9W3h7yADzSZXzwQ7a2HAWjlCSZq2n9cM8bIAKo70xoY8u73JCIECprAGKfQJiXMnQhKg1G1WGfZDqRvaa/bjD6/G+Xl+NFL+XcGRzNO97MHoo0aX9qzsa+5rrWdZwAJVIAPrgbHLowFFLmvO/n1B6MtJgIPpoWbU/jbt8Ync7w3ZTsGoKA0AoaG5nQiXr9wa3xXZ/0DbdjdhplccmI2OT6zODw5M7WQuzw2HdgJ/0ZRRKKCUzIedyxyjIjNDQPwKx4xFShdLrsaKLuy7MtS2S+XyzIoWfDSKae30dm8oak7HeloQDqBJGACZgBTIKtxdgGvvr86PlPJq4AlsbM9+i92JTIONmewuTk+PO0dP1fkKjp0bXIlt7C5KfL8Yz1Jgdk5KxMzobAgcXQUPzg5P5HTDRH1tfu7jrQjGgZG+pyRBCkwRdWCSrWAQCD2kZrDh7ba6yPGav0bghJK6aoNDoWtAh5+CjiRYsqLSL/HMn7j4abZ5ZWbi4WXT15T2P5bh8yOVjQ3N0U4tIJicIHzF2/MZCu7Nm16Zr/TZcICygZuLuAn70+fuTnZns4c2N/T0YqAUOR4f1QH2jz0xKE6A0mOOqAOyCSxPYncxsa5YuNEDiMLmFjxZnKVgivzJbe4ms0ulvxKoFzDsZxAlrQOJEgTLMvwfT8i7Lp4or+zsaPJam9ERwMaHTREQyfOKmTEUoFmIi9x4qb7X39yYd6PecXABwU5Vswu7mvd7mxACtjTU/fmhcLCivvzk6sNlvvAzq6n7kttSsCUYJk6xaAJ743hxbeXbswW68zg+Ud3Ht6MhIIIQWpcKK0JBmowmLUlTkNBf8Rvce3v9fskzdcBIUj7FCphawpxlCFFMADXZBEzAJYS2KDxB89v/MM/e2fMT/34/GzM6vzCHqQ5bKAgETAUA5RKnmOa3RlnIIWkxrJCluPoEF45nzUTdZv6W3Zs4RAoKrxxAd89Nlz0gv6U9bvPb2iOgADDUwQPwjAUd+JojWN/O1xpGtzUQK6MYglSoxJgtQhXolBxJelAM8YRjxmOQY0WGmJwTDCANAyNOIOWIFblNUCDMVECjo/i229dmnL19k2tTw7UlZeKJ68PNbW1NmzA9SV0AHt7sLW7+dRw1mbqi4/3HtmEZsAAIhwGsBzgzAr+5OXRS1Mlk7xnH9z41B40ADzQtiCpNGeWBpgC1VSKQtA3UcB0COTXt0UIt7UqSCuANHFoBrbOTBLhFkoSOIgrguKMa2hXxzR1x/Dvv/HAt/5ucjLrvvruVYHNX9yFeoIpoBVSAls6MxemJpeWC/PlmCIsMfz0DF48MVOiyL4u54uHY0Jg0cPRs3jx+I1JJCO2KMpCpQwVhdIgg2lthxep9mAwGApxpWylBdftESEjkIAHaMAFyrBkregfzhIHsABo31MkSRBgAxXXC2AWNW7Ne44pepvYQgVnJ7ypstnb1/H8r9TtFygVo9t27xzJ4y9eys2PXf3G/QP37q7f1GNcGStGuKqPoAkISoFhi5WKNCx+Ywn/5fs3ri94yZj5yO7+Lz9kNANBJXAsxkC+HzDLAMBVyOiiEFsa6nwpCmkP7LYIfXS5I1VrC3BNIM3DRS681ECh/HZV6i10HGCMDAMm0BfDv36y/W+O3rg8Sy+dHGPU/fRutAF+IZ+OJw7vzLwzXPhgaOG/c7Ghxz435J67MVkI7L29ma/sMnfEMFbC6+eDn16YnlFWczohl6fbEywpUQeQxpTE4EgxEot2tCJpgvkghYjFeA2XCUBpGAQFxcEsoByAB7BNSAYZSEvwQPmKBZJHcoACfCAQpgY+GMa3j15IJsz/8LWdZGOmoFYL/mNd6U4BC1gGfnCs8u5oYWwuXydSJybQug293WiN5cuForfaS0DUEYGEcviJGfz5K+PDKxRxnIMD9V89YrRbsKQ2DQiCkr4QQgEqqKlFEQBw7TFon0iCceKffk+qPQYNCvPC0C2jCm4CC6vFFGKmNSCggQhQLuHeNMzH+r/9xuzF0YXv/uT9grfryc2iN5kgYFMzvvZE39+/lz19dvAfzjBmxwnUn7aeus8ZaMJqBUfPyR++d63CzMcf66+P4u2/rzBP2xxCI+B4dxR//foV0xJH9m59fIcZBQJgtQBiqItW2Vghk90HqwAaMARMDXjQ8KOW4BLEDA/GSLb0xoVFV4lN7and/bb24DMsesmhsYWpVTTVIxa3E9HEzavD7j29ZYFiDitz4y313YnGxqErIxNL5bKsH0jjczsy09PTbQmoACWJMuHMCL7/ztTl6XzE4I/u3/jcQXTaMHxEw22ZXyFiXBiBhCFq0ji1001QFAp03ylCHw1SCHysYcYCqnrjcSAEk4YyVQC4hqGKIAJziqWK4/AWh4oB7Uzzbzye+e7R4NS495fHxxcqbc8dtLZYEApPbEBfrP58/7bzNyaWVryB3g2bNsQ2bcBiCS+8sfrDCzP1zR372uiZe3BjFEG+mNzQIRyUAvgmzi3jQikuim5kZOHQjjbDwLUpvPXOWT/Al57c3deECEdOYiKHKzNYKSMWRXsd+uvQZIEpgzS0r4WgCsfwuH7hrYkVHdvcY5RN+0g3ejvQlm6dW1g+d2H+0cPpzgZECJdG8xfG0NqL9iR+875uM2O+cROzl5bgRryV1qYYfnV/cy4fSzfCIRQFfnoWr52anCkEmXrry/f1PLING2yYIVREEbTi0iPLBlymAsYjAWcAuASFvDxwphmwxmXG+hVP63XE5ipEp4YfWoMR6OqawhQoxBFVVbe10r4XdexyJS8EiwvbBLam8K++0L78wvyl6fybg5OTM+x3v7BhYwSGxtYmtDU4T+8dKLhgDLaBWzn84M35nw1O8rr22dlpZJqnpnFmEJaTDNwKJ5CJFYVL0/6CilvMvjKXX/LaWAQ3cvjphYmExZ+2d5c5Fir42enl109emcnxiuZceA1RenTXpuceauiwID1EBCkFxhGNRiPJroWyc/x6ruDK+FOZjjakmmLxWMO5K+NPHk73ppGJsdmi879eHXce77xvC3r6zdFFnD99nVeWWuNtzQ5igGZINEUD4NYq/u7Y5NFL2RVpxG35m09vfaoPKcAClO9xbioFxkBWJFSPMQ3uKxcsEi54IFUF82vGKYQFfCRC4d9i/bQK65EEvQ7zEdqeVl0SeIhdJw5YgCYuADA74QMBJEfQpIUTwR/+WvoHxxNvXBl//1bhj3409eWH2g60Ii4RUXAY4gKrPm7O4YVjw8cuTdrJxmZr1a6PXxlZOjtcdFl9hOx03LY1ygoTK5icWdTMDAxj0S1fnQLvxfkZ5EWqvyOSqMMS8OrJ0quDMwt+pDUVG2jLTGWnZpeW/+rUjVV+z28firRzQIIxSBctjUgnomOrbn2m+8rc9F+8Nf/sU+nGNuBMMJcLZrLY1IhDmzLXV2ZH8/StH8++ci3jGBi6OZTLFTa2tzxz35Z2u7rEBQYGJ/HDn88Njsz5OuhvT37pvi1H+tAAmCHmW5iaEM4ZABo6VPvi4RTQoTJteONHldhaq3R/xOyT6EPJ6fWDoPn6zG7teDWzhyYR6kiE2UQIuuFaMuXFVGSjRc/tsxvS/W+enxiamPvuq8WRzd2fu8fsi6PgQ8tKxLJtgaaYsbO7Jd7Y+NVnUzMLOH4xODuaG85WEpFoImHZEeQI16ZRyBcbE5mAfFmmG1PlaCoylteucLrammzgzBjeuzmzpMWmHTueP2DU25gpJ98dLB0/fWV2OZ9bjXTVVTFCBsGxUBcXKlgM3Eh7Z9uF8at0PN3cDlcYFatuagVb6vH4Pmd4tfX98fzE7PL7F/Mmoxj37tvS/diuhp5mCAsljjkfb76H10/dnF0NODcPbG59fE/Lg91IAMFqkRncsO0wQQ4RcmFY8PHTCLa+OXFH4/Pqckd3661/6lAapCxeFfQLaVhM6YSmHXF0bUdftP2VU6VLU4WXLiy+P5V8Ynf0YB+aDdsPsKEeX3+wc64E00Ibw85W7G0Rrw82/I/XJnXJq8hUjmFZ4Z2LM/CK926MCQdnTk0PDud4fWRqxbct556eTge4NoHxpVxTOvXUPmNXE6JAVxwth5x7Wvb2NqCrEdpXgQ5AxISREGhrj9ePTlhO9okHkufP95y/djnlDeSjTqm0en7MO9Jtdkbwbx6zT1zW52+slAMk69N9zcm9bejJgIAF4P1J/Pi0Pzi2kF/xMrb3+N4NT+6uG2iEDQS+jMQjJoAg9Ljja7W49ZPhjsFYP9bAxuG/d8juPuvQtXaNDkAKYERMwdAA1zClikv1SI/Y1DPww1P6h++OXbhxa2WR3RxLf25faqAONpCMoikFDkQAX4IBaYLj5+L1qfr6OAPGV3FlYi4VEXu60JDBtfflyGzZvYSlXKXNiXQ2QEksLmclKJ0we5uRBKYm3dOjhZxKSGaMruoW19vSYpmcE8lQZqSpHhFeVuVyX6a757AVFOJnb40Z9c1ugPHFYqFixrRsdfhT2yNHtve7QIXgcDQAFYWRAt6+Kn92emp2VQZK9rXGf/1I54MDaNKIAr6EYJwBSnvMr8C01hgVn/WMar0GB0ctrh9JHH7RoYgkA8BFKI8LpsElh2QAIZBwhA6Uzmj65/vpnpYN3zs6fHE6+NGJ7M8H8w/v7X7uALpsuBJxDmho6QEmgpy7Os+565WEh6axebiB6qjTG5NojKOzLjqzat2Y1tpDZ1dTSx2YD6F8A56FigUUgNM3Z77/s6tFnsn6EV6e++bTva0t7XHwCKpt3KYI4kLP58o6j8OtyDzZ+cevFy4vAiyxuFpaWEV3mpc8bRFFGSqEAoMPXHdxeQovn129OrYgi5UUc+/fkvnKwy2b6pEApAcIGDwskyquNQxWY+z8AkG62/glZlLtXqU0gTg0C/UkQq0Pw2DKLRlCNhq2CxxsQ9MXev/23eDUyPLw4vzLb58bGTIe3Nm7Z0ukNQoXEDCiJg5uSf7brx5hJrZnYANz46vCX40EyEQQAwbam0/dWmEqsGRxS2tLnY0gQEdT3LquV0v+UoA6gdbe7oce7Z4q4OfnlgLlN3e0MxPSB4xqvbg+hvqIPbfoTo3CacWeZjx3MOYfW8hX/LTwqFwKAltLUhwKcBmyEtfn8c655eODw/MVK5lMbsg4T9+75amdSDPEFHjgW8LwpAIxUUV9c2gF14Vtrj9dt/H3Pj7WVrn1y13VVvuO9aJPHQRt6CqKUWsWpu5V8xoJQAnbAUH6geCinmNzHP/uUXFmIv3S+fKp4dkz4/6l+dHOqx17++OPbUGXSe0SnTa+sgeeRpRQBHoilUp3rDvBExwCGOiIJYxZqRbTfGVrS0vYm9naGXm3sWGkwN8YBO9FUxq7Yyiegyxnm+O8KQlZhiNgKvgajCPCEbVshcro2KK7p7HOwq/0ojvelM16mYjY18K4r3xGOY7pMi7P4MQ1NTg0UywWE5H6+5uNTe2Jxw7Ub4hDuDBF6OFkaPgMvkEGwKGgpGaGCce4jQ/x2cdtURSfehO7a5A0gIB0lSalQ9BKKHZEihikVyFhcMNwPaVAFiFC6v5ensx0bRxqPT64eHV85frI9Py8eeM6O9jbuKctOtAKx4JNCI34njrQtLe/qdFClENo7OrGtkZ3Zn5hc9LcmgHX4Bq7O7G7PT5+deHcudHiiKyzrYIbvzEyVydye/paOuqRMMAltKqadKUi6EyJXIo3sqJNjZaEBPalUU6bAnA1SppdXsTZcXn25uTYYilbYp4bdKWT+/qbH9ps9KdRb0EAlgACqNBqhGkjdP2RCsxiBpMagQoMup259ymn9I6ZNhGtWSPccT7dlrB/5NVaEjTWOr5V/HuIVVS1tZBLGC6hQtAapg9wXeCUk1jI4/IY3h6c+ODmeODELSfamEhsaK7f2yd2daPNggl4EhEOIQEfhgHJMZTFrYnFtGkOtCdSMSjpKmnNuXj50vIP37s8X4FktnKDOJf7+xueOdi/v50npM9hSAlN2heUDbC8ClSUw2VLvWDQAJMmioQF4INhnB523xsp5CpSl3LczbZFg/u2dT2yt21LGnENS4GzAFpChZRHBkYUdog0CwtlVYg5qbVtzG2n7uORWNcZv8Pzf4kgAZoUtGRKVxFuxKpIaqWrHpTgigyXUAEARCSY9j2iAEIRCgFmChgv4qVT0zeXK4tFZRBLopJkQVdzamNHw6HtkToHUQ7mwwCiRi3yqOrjEWRE84Awq3Ezj0sLKBEshiYHvUm026jTPq94nEcVoDhcjiIgNWKASSiGKkzA+CLODcsPhievzeVmyrrEUkr6XY46tKnpqV3JXRlEAb9YqbMtrr0qIlMzrZlmXIMR1TJtzTSYrvHRuEaV6/mx8Rmz8NuD9PF4rj/y8WjJ0FIeMFVYKAoZW6SIK2gW+hTWAHw+EWkYvgSYp3lFgxkQvIoJWQIG59TJkcVrk/mZWW8xHxS0xSzLEaqvLb2zL7apFS1JNFhIMtQqKmCoigyGnxIAEihqVDQ4gw2QhBX4JhlSQQsEAgVgFSgBnouSj9k8rk/j0sTC9NJq2fW1Lx1SUY5UIrJzY8eBAbM3hUbAUOAaBkfVukqDaaXhASE1nvGqzK4PUiAmmeER04ChwfWdg/SpAbtrkD7je1X1qYnCFo6hIHR4lpRiUEQKRJqJqrKbBqQON9luADIkDE9DEWyCAQX4Ltcus7PA9TyGJ3FjHhdnS6Pzy67rardiaZ2KsdZG3tUU2dDQlK6LGDFELEQNmAwmg2MgwWHViBgeEBrm6gA6AAIECnkXc2UslrBQ9JbypaXlYq5cWcwWKxquYMK0Y7bVknA2Z+r7mlhfC9I24oADRACmlO9VFDGybF3rBhH8kHBJxLkUTIf7RQlSmrhPpgKMO80k/TGK+d2i9eHq938dJIQzSVelZKtUH/KJaaU1wJk2wt1ulecF5VVTdAq1YIzq+ljjf4XVdyALXF3E0Dwm57JTs9nJmXy2IguCBWQIihqmJSzDibCoIM4MZsVjNhoMREhbJnEGFaBY8UseAcS1llKV/WDFDRZdXfL9oLTilQtKyQgL2utYJmH0ZNJ9Hc3tzbw+gaiABRiAE9KFJCwOAxC6zIj52gqdIlQVmxxukENVzZoRJHmkZdhS0KHgysfO4WcM0trZJinlJ9zW7pjaf3gw9NNSNb4PVOg1qHUoF75uE8aUAgIwDY2QRKzBtAhvroogGBik8pUnhc/It+ACWWCpgKUspldwc0WOzBXmlsqFsswXVoPA07qimOGyJGnEqWhz6fmKiAymOTMCbQeBZIEEBSDfB3weMblImmiui27obO5rNTY2I+OgxUE8VMytNXrIcw3GFRMaEAwEGNrjREpxEJM1NxFAMQQEXq106ir2n8Gt2hMwU98mVveLBwnAh0Giu5RgP/GtFbQmVWX+hFoFklFo8c4VeFhchAT5AOAzMK5EUKWAaR6AS4IbKE5MKDCNUDtJ6WBVKk+Y4Y0HQBnIeyhUUA4wNYPFfHEquzhXqMwXjGLFI7lKzJeEkue7AYc2bdiW4PVOJBGntiajIWlnGhpTccQtxATiNqK8qv8l1ixfAK0hpTIYSeUrzYlxAJxX1UzCArZmga7S9zSFXh4KGqyqplWbT4AGiU9o5X38lP7jB0nrMIXTtJZvMyYZFLFQnIXX3GChA/AAAJQJAlggdQBFmkgzU4c2BBoUwCAQg5KuQZ4WRkFbRGT6JS4VFxa4IQmuAmNwgbzCqoILQEBreAEUh2KQgNYQPgTB4rAYLCACmAAUhIbJIBSk7wmrWhHgJFnIXtIcWkNJKSU3LTAGWVWdkqGfcBXOGwp+6qpLVCg0Fhpche8YWivgDmvd7QG4Sy59e5Bue/bd0u7bQvXhLWp9UyOU2qrttAmgML0JtY1CETD6EPKsq6AJBlTX9PAwtK/BAnAFzZQfXgwajIiHTL2QmB4qyoRdyRDDtnbBUA3+QIBRe/RD6UZd3dkQEeNgoGoaXTPerV6doXEfq576auuBPsT8rv1MDdQYHuFQ+Oglvv58ri+Hr82QtXLqbQ/dIUh33PF+cpA+dXzGmtX6oatd/A9n9vpPX5OjuP14LcqhSe/auzFiGlopdRs5K1ys1r/bJ5+Bj3/Dz/LDP/mi/4Qnr41fqsD6Gcf66+UXeiFjLHztJ6StHy+lrG+vrT9y+zae1tqhv3CEbvvcX/R3rf2i9R/3Cb/xDrnHbeMTXvyPNe72ETXT3U+5wd75tXc5+JFN4i/T8PwlxvrofpYZ+f9jJn2WcVuc1r79Ojtk4E6z5LZ3+IQnfPza/+T4/L+L321L96d+h/8DkFGo+FjU1yIAAAAASUVORK5CYII=" alt="Logo Chi cục Thủy lợi Quảng Ngãi">
            <h1 class="welcome-title">Xin chào!</h1>
            <div class="welcome-card">
                Tôi là trợ lý hỗ trợ công việc của Chi cục Thủy lợi tỉnh Quảng Ngãi.
            </div>
            <div class="welcome-card welcome-card-large">
                Anh/chị có thể yêu cầu tôi tham mưu về thủy lợi, cấp nước sạch nông thôn, đê điều,
                phòng, chống thiên tai, tài nguyên nước, khí tượng thủy văn; quản lý chất lượng,
                thi công xây dựng và bảo trì công trình xây dựng chuyên ngành trên địa bàn tỉnh Quảng Ngãi.
            </div>
            <div class="welcome-spacer"></div>
        </section>
        """,
        unsafe_allow_html=True,
    )
else:
    visible_messages = conversation["messages"][-30:]
    hidden_count = len(conversation["messages"]) - len(visible_messages)

    if hidden_count > 0:
        st.caption(f"Đã ẩn {hidden_count} tin nhắn cũ để tăng tốc hiển thị.")

    for message_index, message in enumerate(visible_messages):
        avatar = "👨🏻" if message["role"] == "user" else "💧"
        with st.chat_message(message["role"], avatar=avatar):
            if message["role"] == "assistant":
                render_assistant_content(message["content"])
                render_export_buttons(
                    message["content"],
                    key_prefix=(
                        f"history_{conversation['id']}_{message_index}"
                    ),
                )
            else:
                st.markdown(message["content"])

# Điều khiển vị trí cuộn:
# - Có hội thoại: cuộn xuống tin nhắn mới nhất.
# - Chưa có hội thoại: luôn đưa màn hình về logo ở đầu trang.
if conversation["messages"]:
    st.markdown('<div id="chat-bottom-anchor"></div>', unsafe_allow_html=True)

    components.html(
        """
        <script>
        setTimeout(() => {
            const anchor = window.parent.document.getElementById("chat-bottom-anchor");
            if (anchor) {
                anchor.scrollIntoView({ behavior: "auto", block: "end" });
            }
        }, 80);
        </script>
        """,
        height=0,
    )
else:
    components.html(
        """
        <script>
        const forceWelcomeTop = () => {
            const parentWindow = window.parent;
            const parentDoc = parentWindow.document;

            const active = parentDoc.activeElement;
            if (active && typeof active.blur === "function") {
                active.blur();
            }

            parentDoc
                .querySelectorAll('textarea, input')
                .forEach((element) => {
                    if (typeof element.blur === "function") {
                        element.blur();
                    }
                });

            const welcomeAnchor =
                parentDoc.getElementById("welcome-top-anchor");

            if (welcomeAnchor) {
                welcomeAnchor.scrollIntoView({
                    behavior: "auto",
                    block: "start"
                });
            } else {
                parentWindow.scrollTo({
                    top: 0,
                    left: 0,
                    behavior: "auto"
                });
            }

            const main = parentDoc.querySelector(
                '[data-testid="stMain"], section.main, .main'
            );
            if (main) {
                main.scrollTop = 0;
            }
        };

        [0, 100, 300, 700, 1400, 2200].forEach((delay) => {
            setTimeout(forceWelcomeTop, delay);
        });

        window.parent.addEventListener("pageshow", forceWelcomeTop, {
            once: true
        });
        </script>
        """,
        height=0,
    )


# =========================================================
# Ô NHẬP DƯỚI CÙNG VÀ XỬ LÝ CÂU HỎI
# =========================================================
# Ô chat native của Streamlit:
# - Nút dấu cộng để đính kèm một hoặc nhiều tài liệu;
# - Shift + Enter để xuống dòng;
# - Enter để gửi.
chat_submission = st.chat_input(
    "Hỏi Trợ lý CCTL_QNG...",
    key="main_chat_input",
    accept_file="multiple",
    file_type=["pdf", "docx", "doc", "txt", "md", "csv", "xlsx"],
    max_upload_size=200,
)

question = ""
chat_files: list[Any] = []

if chat_submission:
    # Khi accept_file được bật, st.chat_input trả về đối tượng có
    # thuộc tính text và files.
    question = str(getattr(chat_submission, "text", "") or "").strip()
    chat_files = list(getattr(chat_submission, "files", []) or [])

    # Đưa ngay các file đính kèm ở khung chat vào Vector Store.
    upload_messages: list[str] = []
    if chat_files:
        try:
            client = get_client()

            for attached_file in chat_files:
                suffix = Path(attached_file.name).suffix.lower()
                if suffix in {".xlsx", ".csv"}:
                    save_table_file(database, attached_file)
                else:
                    upload_document(client, database, attached_file)
                upload_messages.append(attached_file.name)

        except Exception as error:
            st.error(f"Không thể đưa file đính kèm vào kho: {error}")
            st.stop()

    # Cho phép chỉ đính kèm file mà không cần nhập câu hỏi.
    if not question and upload_messages:
        conversation_id = conversation["id"]
        attached_text = (
            "Đã đính kèm tài liệu: "
            + ", ".join(upload_messages)
        )
        append_message(
            database,
            conversation_id,
            "user",
            attached_text,
        )

        with st.chat_message("user", avatar="👤"):
            st.markdown(attached_text)

        answer = (
            "Tôi đã đưa tài liệu vào kho dùng chung. "
            "Anh/chị có thể đặt câu hỏi về nội dung tài liệu ngay bây giờ."
        )

        with st.chat_message("assistant", avatar="💧"):
            render_assistant_content(answer)
            render_export_buttons(
                answer,
                key_prefix=f"attached_{conversation_id}",
            )

        append_message(
            database,
            conversation_id,
            "assistant",
            answer,
        )
        st.rerun()

if question:
    conversation_id = conversation["id"]

    displayed_question = question
    if chat_files:
        displayed_question += (
            "\n\n📎 **Tệp đính kèm:** "
            + ", ".join(file.name for file in chat_files)
        )

    append_message(database, conversation_id, "user", displayed_question)

    with st.chat_message("user", avatar="👤"):
        st.markdown(displayed_question)

    structured_table_answer = lookup_structured_table(
        database,
        question,
    )

    if structured_table_answer:
        answer = structured_table_answer
        with st.chat_message("assistant", avatar="💧"):
            render_assistant_content(answer)
            render_export_buttons(
                answer,
                key_prefix=(
                    f"table_{conversation_id}_"
                    f"{len(conversation.get('messages', []))}"
                ),
            )
    else:
        with st.chat_message("assistant", avatar="💧"):
            try:
                client = get_client()
                current_messages = database["conversations"][conversation_id]["messages"]

                auto_document_search = (
                    bool(chat_files)
                    or (
                        bool(database.get("uploaded_files"))
                        and question_requests_documents(question)
                    )
                )
                effective_file_search = use_file_search or auto_document_search
                effective_fast_mode = fast_mode
                effective_deep_mode = deep_mode

                spinner_text = (
                    "Đang phân tích chuyên sâu..."
                    if effective_deep_mode
                    else "Đang tra cứu nhanh..."
                    if effective_file_search
                    else "Đang trả lời..."
                )

                with st.spinner(spinner_text):
                    answer_parts: list[str] = []
                    for text_delta in stream_openai_answer(
                        client,
                        database,
                        current_messages,
                        use_file_search=effective_file_search,
                        fast_mode=effective_fast_mode,
                        deep_mode=effective_deep_mode,
                    ):
                        if text_delta:
                            answer_parts.append(str(text_delta))

                    answer = "".join(answer_parts).strip()

                if not answer:
                    answer = "Tôi chưa tạo được câu trả lời. Anh vui lòng thử lại."

                render_assistant_content(answer)
                render_export_buttons(
                    answer,
                    key_prefix=(
                        f"latest_{conversation_id}_"
                        f"{len(current_messages)}"
                    ),
                )

            except Exception as error:
                answer = (
                    "Không thể kết nối hoặc xử lý yêu cầu với OpenAI. "
                    f"Chi tiết lỗi: `{error}`"
                )
                st.error(answer)

    append_message(database, conversation_id, "assistant", answer)
    st.rerun()

